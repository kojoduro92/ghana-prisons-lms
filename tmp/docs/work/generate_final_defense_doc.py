from __future__ import annotations

from dataclasses import dataclass
import json
from pathlib import Path
import re
import subprocess
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/emmanuel/Desktop/apps/ghana-prisons-lms")
OUTPUT_DOC = ROOT / "output" / "doc" / "Ghana-Prisons-LMS-Final-Defense-Thesis.docx"
ASSETS_DIR = ROOT / "output" / "doc" / "assets"
THESIS_CONFIG_PATH = ROOT / "tmp" / "docs" / "work" / "thesis_config.json"
REFERENCE_CHAPTER_DOC = ROOT / "tmp" / "docs" / "reference" / "CHAPTER TWO (1).docx"


def load_metadata() -> dict[str, str]:
    defaults = {
        "title": "INTEGRATING INTERNET OF THINGS (IOT) TECHNOLOGIES FOR PRISON EDUCATION WITHIN THE GHANA PRISONS SERVICE",
        "author": "THEOPHILUS OPOKU",
        "index_no": "40062289383",
        "institution": "BLUECREST UNIVERSITY COLLEGE",
        "department": "Department of Information Technology",
        "degree": "MASTER OF SCIENCE IN INFORMATION TECHNOLOGY",
        "supervisor": "DR. WOLALI AMETEPE",
        "submission_date": "FEBRUARY, 2026",
        "live_url": "https://ghana-prisons-lms.vercel.app",
        "deployment_url": "https://ghana-prisons-g823avh29-kojoduro92-gmailcoms-projects.vercel.app",
        "logo_path": str(ASSETS_DIR / "bluecrest-logo.png"),
    }
    if THESIS_CONFIG_PATH.exists():
        loaded = json.loads(THESIS_CONFIG_PATH.read_text(encoding="utf-8"))
        defaults.update({key: str(value) for key, value in loaded.items() if value is not None})
    return defaults


METADATA = load_metadata()
TITLE = METADATA["title"]
AUTHOR = METADATA["author"]
INDEX_NO = METADATA["index_no"]
INSTITUTION = METADATA["institution"]
DEPARTMENT = METADATA["department"]
DEGREE = METADATA["degree"]
SUPERVISOR = METADATA["supervisor"]
SUBMISSION_DATE = METADATA["submission_date"]
LIVE_URL = METADATA["live_url"]
DEPLOYMENT_URL = METADATA["deployment_url"]
LOGO_PATH = Path(METADATA["logo_path"])


@dataclass(frozen=True)
class FigureSpec:
    number: str
    title: str
    filename: str
    width_inches: float = 6.1


@dataclass(frozen=True)
class ListingSpec:
    title: str
    path: str
    start_line: int
    end_line: int
    description: str


FIGURES = [
    FigureSpec("4.1", "Public landing page of the Ghana Prisons Learning Portal", "landing-page.png"),
    FigureSpec("4.2", "Unified access page for all portal roles", "access-page.png"),
    FigureSpec("4.3", "Administrative operations dashboard", "admin-dashboard.png"),
    FigureSpec("4.4", "Inmate registration and biometric capture page", "admin-register-inmate.png"),
    FigureSpec("4.5", "Lecturer dashboard for course delivery and assignment workflows", "lecturer-dashboard.png"),
    FigureSpec("4.6", "Management analytics dashboard", "management-dashboard.png"),
    FigureSpec("4.7", "Clock-in and device allocation page", "clockin-checkin.png"),
    FigureSpec("4.8", "Inmate learning dashboard", "inmate-dashboard.png"),
    FigureSpec("4.9", "Inmate course catalog and assigned learning modules", "inmate-courses.png"),
]

LISTINGS = [
    ListingSpec(
        title="Listing B.1: Role-aware authentication configuration",
        path="src/lib/auth.ts",
        start_line=17,
        end_line=54,
        description="This excerpt defines seeded credentials, role routing, and sign-out redirect behavior for the portal.",
    ),
    ListingSpec(
        title="Listing B.2: Session creation and inmate access provisioning API",
        path="src/app/api/v1/auth/login/route.ts",
        start_line=11,
        end_line=48,
        description="This excerpt shows validated login, session creation, and the automated inmate facility access grant used by the LMS.",
    ),
    ListingSpec(
        title="Listing B.3: Clock-in biometric workflow and resilient camera handling",
        path="src/app/(clocking)/clockin/checkin/page.tsx",
        start_line=175,
        end_line=219,
        description="This excerpt captures the clock-in page state model, data loading, and the face-verification camera fallback logic.",
    ),
    ListingSpec(
        title="Listing B.4: Lecturer dashboard data loading and instructional actions",
        path="src/app/(lecturer)/lecturer/dashboard/page.tsx",
        start_line=46,
        end_line=99,
        description="This excerpt demonstrates how the lecturer interface loads course, assignment, submission, and attendance data for instruction monitoring.",
    ),
    ListingSpec(
        title="Listing B.5: Inmate registration and photo capture workflow",
        path="src/app/(admin)/admin/register-inmate/page.tsx",
        start_line=128,
        end_line=180,
        description="This excerpt shows generated identifiers, biometric readiness states, and live photo capture logic used during inmate registration.",
    ),
]


ABBREVIATIONS = [
    ("API", "Application Programming Interface"),
    ("CSRF", "Cross-Site Request Forgery"),
    ("GPS", "Ghana Prisons Service"),
    ("ICT", "Information and Communication Technology"),
    ("IoT", "Internet of Things"),
    ("LMS", "Learning Management System"),
    ("QR", "Quick Response"),
    ("UI", "User Interface"),
    ("UNESCO", "United Nations Educational, Scientific and Cultural Organization"),
    ("UNODC", "United Nations Office on Drugs and Crime"),
]


REFERENCES = [
    "Al-Fuqaha, A., Guizani, M., Mohammadi, M., Aledhari, M., & Ayyash, M. (2015). Internet of Things: A survey on enabling technologies, protocols, and applications. IEEE Communications Surveys & Tutorials, 17(4), 2347-2376.",
    "Atzori, L., Iera, A., & Morabito, G. (2010). The Internet of Things: A survey. Computer Networks, 54(15), 2787-2805.",
    "Ball, R. (2021). The Internet of Things in education: Improving learning through connected devices. International Journal of Smart Technology and Learning, 3(2), 101-115.",
    "Boateng, F. D. (2021). Prison conditions and rehabilitation in Ghana. Journal of Correctional Education, 72(1), 21-39.",
    "Champion, N., & Edgar, K. (2013). Through the gateway: How computers can transform rehabilitation. Prison Reform Trust.",
    "Creswell, J. W., & Creswell, J. D. (2018). Research design: Qualitative, quantitative, and mixed methods approaches (5th ed.). Sage.",
    "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340.",
    "Davis, L. M., Steele, J. L., Bozick, R., Williams, M., Turner, S., Miles, J., Saunders, J., & Steinberg, P. (2013). Evaluating the effectiveness of correctional education. RAND Corporation.",
    "Edovo. (2025). Technology for rehabilitation and learning. https://www.edovo.com/",
    "Gbazuo, F., & Ofori-Dua, K. (2021). Vocational rehabilitation in Ghanaian prisons: Constraints and opportunities. African Journal of Criminology and Justice Studies, 14(2), 88-103.",
    "Ghana Data Protection Act, 2012 (Act 843).",
    "Ghana Prisons Service. (2015). Ten year strategic plan for Ghana Prisons Service. https://ghanaprisons.gov.gh/downloads/Ten_Year_Strategic_Plan_for_Ghana_Prisons.pdf",
    "Ghana Prisons Service. (2024, January 24). Prison inmates set another record in 2023 BECE. https://ghanaprisons.gov.gh/2024/01/24/prison-inmates-set-another-record-in-2023-bece/",
    "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information systems research. MIS Quarterly, 28(1), 75-105.",
    "Ithaka S+R. (2024). Technology in higher education in prison programs. https://sr.ithaka.org/publications/technology-in-higher-education-in-prison-programs/",
    "Kinyanjui, M. W. (2019). Education in African prisons: A comparative study. African Journal of Correctional Education, 4(1), 12-27.",
    "Massachusetts Department of Correction. (2024). Massachusetts Department of Correction completes implementation of tablet initiative to enhance educational programming for incarcerated population. https://www.mass.gov/news/massachusetts-department-of-correction-completes-implementation-of-tablet-initiative-to-enhance-educational-programming-for-incarcerated-population",
    "Nielsen, J. (1994). Usability engineering. Morgan Kaufmann.",
    "Nisser, M., Gaetz, M., Fishberg, A., Soicher, R., Faruqi, F., & Long, J. (2024). From prisons to programming: Fostering self-efficacy via virtual web design curricula in prisons and jails. arXiv. https://doi.org/10.48550/arXiv.2404.15904",
    "Palanci, A., Yilmaz, R. M., & Turan, Z. (2024). Learning analytics in distance education: A systematic review study. Education and Information Technologies, 29, 22629-22650. https://doi.org/10.1007/s10639-024-12737-5",
    "Pressman, R. S., & Maxim, B. R. (2019). Software engineering: A practitioner's approach (9th ed.). McGraw-Hill.",
    "Rogers, E. M. (2003). Diffusion of innovations (5th ed.). Free Press.",
    "Shaikh, F. K., & Zeadally, S. (2016). Internet of Things: A vision, architecture, applications and future directions. Digital Communications and Networks, 2(4), 187-216.",
    "Sommerville, I. (2016). Software engineering (10th ed.). Pearson.",
    "UNESCO. (2015). Revisiting global trends in TVET: Reflections on theory and practice. UNESCO.",
    "UNODC. (2015). The United Nations Standard Minimum Rules for the Treatment of Prisoners (the Nelson Mandela Rules). United Nations Office on Drugs and Crime.",
    "UNODC. (2018). Handbook on prisoner file management. United Nations Office on Drugs and Crime.",
    "Wilson, D. B., Gallagher, C. A., & MacKenzie, D. L. (2000). A meta-analysis of corrections-based education, vocation, and work programs for adult offenders. Journal of Research in Crime and Delinquency, 37(4), 347-368.",
    "World Prison Brief. (2025). Ghana. Institute for Crime & Justice Policy Research. https://www.prisonstudies.org/country/ghana",
    "Yeboah, T., & Adomako, E. (2022). Technical and vocational education and training in Ghana: Emerging trends and challenges. International Journal of Vocational Education and Training Research, 8(1), 15-27.",
]


QUESTIONNAIRE_SECTIONS = {
    "Section A: Lecturer / Instructor Questionnaire": [
        "What courses or rehabilitation modules do you currently facilitate for inmates?",
        "How often do you experience delays caused by paper-based attendance, manual grading, or limited classroom time?",
        "Rate the usefulness of the LMS dashboard in tracking course progress and attendance. (Very Low, Low, Moderate, High, Very High)",
        "Rate the usefulness of digital uploads for course materials and assignment documents. (Very Low, Low, Moderate, High, Very High)",
        "What concerns do you have about using connected learning technologies inside correctional facilities?",
        "What additional reports or analytics would help you support learner performance better?",
    ],
    "Section B: Administrative / Management Questionnaire": [
        "Which prison education processes consume the most time under the current manual arrangement?",
        "How important is role-based access control for prison digital systems? (Not Important, Slightly Important, Important, Very Important, Extremely Important)",
        "How important is biometric verification for inmate attendance and device allocation? (Not Important, Slightly Important, Important, Very Important, Extremely Important)",
        "How valuable are management dashboards for policy decisions, performance monitoring, and reporting? (Very Low, Low, Moderate, High, Very High)",
        "What operational risks must be managed before full deployment of the platform?",
        "What institutional resources are required to scale the platform across additional prison facilities?",
    ],
    "Section C: Inmate End-User Questionnaire": [
        "How easy is it to use the inmate login and dashboard? (Very Difficult, Difficult, Neutral, Easy, Very Easy)",
        "Which learning resources do you find most useful on the platform: video lessons, notes, assignments, certificates, or progress tracking?",
        "Does the digital course structure help you understand your rehabilitation or educational path better? Explain briefly.",
        "Do you feel the portal can help you gain employable digital and vocational knowledge? Explain briefly.",
        "What difficulties would you face when using tablets, kiosks, or classroom terminals for learning?",
        "What improvements would make the prison learning portal more helpful to you?",
    ],
}


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.0)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style_obj = document.styles[style_name]
        style_obj.font.name = "Times New Roman"
        style_obj._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style_obj.font.size = Pt(size)
        style_obj.font.bold = True
        style_obj.paragraph_format.line_spacing = 1.15
        style_obj.paragraph_format.space_after = Pt(4)

    ensure_caption_style(document, "Figure Caption")
    ensure_caption_style(document, "Table Caption")

    add_page_number_footer(section)


def ensure_caption_style(document: Document, style_name: str) -> None:
    styles = document.styles
    try:
        style = styles[style_name]
    except KeyError:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Caption"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)
    style.font.bold = True
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)


def enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_page_number_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE")


def hide_run(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    vanish = OxmlElement("w:vanish")
    r_pr.append(vanish)


def add_field(paragraph, instruction: str, *, result_text: str = "", hidden: bool = False) -> None:
    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run_begin._r.append(fld_char_begin)
    run_begin._r.append(instr)
    run_begin._r.append(fld_char_separate)
    if hidden:
        hide_run(run_begin)

    if result_text:
        result_run = paragraph.add_run(result_text)
        result_run.font.name = "Times New Roman"
        result_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        result_run.font.size = Pt(12)
        if hidden:
            hide_run(result_run)

    run_end = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)
    if hidden:
        hide_run(run_end)


def add_centered_paragraph(document: Document, text: str, *, bold: bool = False, size: int = 12, space_after: int = 0) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_body_paragraph(document: Document, text: str, *, italic: bool = False, indent_first: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Inches(0.0)
    paragraph.paragraph_format.left_indent = Inches(0.0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.italic = italic


def add_source_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(f"Source: {text}")
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)


def add_heading_centered(document: Document, text: str, *, style_name: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_subheading(document: Document, text: str, level: int = 2) -> None:
    heading = document.add_paragraph(style=f"Heading {level}")
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.line_spacing = 1.15
    run = heading.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13 if level == 2 else 12)


def add_toc_field(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u', result_text="Right-click and update field.")


def add_generated_list_field(document: Document, identifier: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    add_field(paragraph, f'TOC \\h \\z \\f {identifier}', result_text="Right-click and update field.")


def add_generated_style_list_field(document: Document, style_name: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    add_field(paragraph, f'TOC \\h \\z \\t "{style_name},1"', result_text="Right-click and update field.")


def add_tc_entry(paragraph, entry: str, identifier: str) -> None:
    add_field(paragraph, f'TC "{entry}" \\f {identifier} \\l 1', hidden=True)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_numbered(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_manual_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_table(document: Document, title: str, headers: list[str], rows: list[list[str]], source: str | None = None) -> None:
    caption = document.add_paragraph()
    caption.style = "Table Caption"
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.line_spacing = 1.15
    run = caption.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        paragraph = hdr_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10.5)
        hdr_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = row_cells[idx].paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(10.5)
            row_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    add_source_note(document, source or "Author's compilation, 2026.")


def add_figure(document: Document, spec: FigureSpec, source: str | None = None) -> None:
    image_path = ASSETS_DIR / spec.filename
    if not image_path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(spec.width_inches))
    caption = document.add_paragraph()
    caption.style = "Figure Caption"
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.space_after = Pt(8)
    caption_text = f"Figure {spec.number}: {spec.title}"
    run = caption.add_run(caption_text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    inferred_source = source
    if inferred_source is None:
        if spec.filename.startswith("chapter"):
            inferred_source = "Author's construct, 2026."
        else:
            inferred_source = "Researcher screenshot from the developed Ghana Prisons Learning Portal, 2026."
    add_source_note(document, inferred_source)


def extract_code_lines(path: Path, start_line: int, end_line: int) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    selected = lines[start_line - 1 : end_line]
    return "\n".join(selected)


def extract_reference_doc_text(path: Path) -> str:
    completed = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    return completed.stdout


def section_between(text: str, start: str, end: str | None = None) -> str:
    start_index = text.find(start)
    if start_index == -1:
        return ""
    start_index += len(start)
    if end:
        end_index = text.find(end, start_index)
        if end_index == -1:
            end_index = len(text)
    else:
        end_index = len(text)
    return text[start_index:end_index].strip()


def normalized_paragraphs(section_text: str) -> list[str]:
    blocks = re.split(r"\n\s*\n", section_text)
    paragraphs: list[str] = []
    for block in blocks:
        cleaned = " ".join(line.strip() for line in block.splitlines()).strip()
        if not cleaned:
            continue
        if cleaned.startswith("Figure ") or cleaned.startswith("Table ") or cleaned.startswith("Source:"):
            continue
        if len(cleaned) > 1100:
            sentences = re.split(r"(?<=[.!?])\s+", cleaned)
            chunk: list[str] = []
            chunk_len = 0
            for sentence in sentences:
                if chunk and chunk_len + len(sentence) > 650:
                    paragraphs.append(" ".join(chunk).strip())
                    chunk = [sentence]
                    chunk_len = len(sentence)
                else:
                    chunk.append(sentence)
                    chunk_len += len(sentence)
            if chunk:
                paragraphs.append(" ".join(chunk).strip())
        else:
            paragraphs.append(cleaned)
    return paragraphs


def load_reference_chapter_sections() -> dict[str, list[str]]:
    text = extract_reference_doc_text(REFERENCE_CHAPTER_DOC)
    return {
        "objective_one": normalized_paragraphs(
            section_between(text, "2.6 Current Educational and Rehabilitation Infrastructure in Ghana Prisons", "2.7 Designing an IoT-Enabled Learning System for Prisons")
        ),
        "objective_two_intro": normalized_paragraphs(
            section_between(text, "2.2 The Concept of Internet of Things (IoT)", "2.6 Current Educational and Rehabilitation Infrastructure in Ghana Prisons")
        ),
        "objective_two_design": normalized_paragraphs(
            section_between(text, "2.7 Designing an IoT-Enabled Learning System for Prisons", "2.8 Infrastructure, Security, and Ethical Considerations in IoT Implementation")
        ),
        "objective_three": normalized_paragraphs(
            section_between(text, "2.8 Infrastructure, Security, and Ethical Considerations in IoT Implementation", "References")
        ),
    }


def load_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, *, fill: str, outline: str) -> None:
    draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=3)
    font = load_font(26)
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    line_height = 34
    total = len(lines) * line_height
    y = y1 + ((y2 - y1) - total) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1) - width) / 2, y), line, fill="#1f1f1f", font=font)
        y += line_height


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#1f3c88") -> None:
    draw.line([start, end], fill=color, width=6)
    ex, ey = end
    draw.polygon([(ex, ey), (ex - 14, ey - 10), (ex - 14, ey + 10)], fill=color)


def download_asset(url: str, destination: Path) -> None:
    if destination.exists():
        return
    try:
        subprocess.run(
            ["curl", "-k", "-L", "-A", "Mozilla/5.0", "-o", str(destination), url],
            check=True,
            capture_output=True,
            text=True,
        )
    except subprocess.CalledProcessError:
        pass


def generate_diagram_assets() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    download_asset(
        "https://ghanaprisons.gov.gh/wp-content/uploads/2024/01/DG1-1140x445.jpg",
        ASSETS_DIR / "chapter2-gps-bece.jpg",
    )
    download_asset(
        "https://ghanaprisons.gov.gh/wp-content/uploads/2022/04/cits-logo-alt-new-e1648908631231.png",
        ASSETS_DIR / "chapter2-cits-logo.png",
    )
    diagrams = {
        "chapter1-problem-map.png": [
            ((80, 130, 400, 280), "Prison Education\nChallenges", "#f2e6d8"),
            ((470, 130, 790, 280), "IoT-Enabled\nInterventions", "#dcebf7"),
            ((860, 130, 1180, 280), "Rehabilitation\nOutcomes", "#deedd8"),
        ],
        "chapter2-ghana-context.png": [
            ((70, 110, 350, 250), "Overcrowding and\nInfrastructure Strain", "#f4e4d8"),
            ((430, 110, 710, 250), "Active Demand for\nPrison Education", "#dcebf7"),
            ((790, 110, 1070, 250), "Weak Records and\nMonitoring Systems", "#e2efd9"),
            ((1150, 110, 1430, 250), "Need for a Ghana-\nFocused Prison LMS", "#f5edd4"),
        ],
        "chapter2-literature-framework.png": [
            ((80, 100, 360, 240), "Correctional\nEducation", "#f4e4d8"),
            ((430, 100, 710, 240), "IoT and\nDigital Learning", "#d9e8f5"),
            ((780, 100, 1060, 240), "Security,\nEthics, Evaluation", "#e2efd9"),
            ((1130, 100, 1410, 240), "Ghana Prison LMS\nResearch Gap", "#f5edd4"),
        ],
        "chapter2-adoption-barriers.png": [
            ((90, 90, 420, 220), "Security and\nPolicy Control", "#dcebf7"),
            ((530, 90, 860, 220), "Connectivity and\nDevice Constraints", "#f4e4d8"),
            ((970, 90, 1300, 220), "Staff Capacity and\nGovernance", "#e2efd9"),
            ((310, 270, 640, 400), "Learner Support and\nUsability Demands", "#f5edd4"),
            ((750, 270, 1080, 400), "Adoption-Ready Prison\nLMS Design", "#dfe7f2"),
        ],
        "chapter2-conceptual-model.png": [
            ((80, 100, 360, 240), "IoT Inputs\nDevices / Identity /\nConnectivity", "#d9e8f5"),
            ((430, 100, 710, 240), "Prison Education\nProcesses", "#f4e4d8"),
            ((780, 100, 1060, 240), "Monitoring and\nAnalytics", "#e2efd9"),
            ((1130, 100, 1410, 240), "Rehabilitation and\nLearning Outcomes", "#f5edd4"),
        ],
        "chapter3-system-architecture.png": [
            ((80, 120, 360, 260), "Public and Role\nInterfaces", "#dcebf7"),
            ((430, 120, 710, 260), "Application Layer\nAPIs / Guards", "#f5edd4"),
            ((780, 120, 1060, 260), "Repositories and\nPortal State", "#e2efd9"),
            ((1130, 120, 1410, 260), "Reports / Deployment\nEvidence", "#f4e4d8"),
        ],
        "chapter4-validation-pipeline.png": [
            ((100, 120, 360, 260), "Screenshots", "#dcebf7"),
            ((450, 120, 710, 260), "Functional\nTesting", "#f5edd4"),
            ((800, 120, 1060, 260), "Lint / Build", "#e2efd9"),
            ((1150, 120, 1410, 260), "Live Deployment", "#f4e4d8"),
        ],
    }
    for filename, boxes in diagrams.items():
        image = Image.new("RGB", (1500, 420), "white")
        draw = ImageDraw.Draw(image)
        title_font = load_font(34)
        draw.text((70, 30), filename.replace(".png", "").replace("-", " ").title(), fill="#0f172a", font=title_font)
        for index, (box, text, fill) in enumerate(boxes):
            draw_box(draw, box, text, fill=fill, outline="#274c77")
            if index < len(boxes) - 1:
                draw_arrow(draw, (box[2] + 10, (box[1] + box[3]) // 2), (boxes[index + 1][0][0] - 10, (boxes[index + 1][0][1] + boxes[index + 1][0][3]) // 2))
        image.save(ASSETS_DIR / filename)


def add_code_block(document: Document, title: str, description: str, code: str) -> None:
    add_subheading(document, title, level=3)
    add_body_paragraph(document, description, indent_first=False)
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8.5)
    document.add_paragraph()


def add_page_break(document: Document) -> None:
    document.add_page_break()


def build_title_page(document: Document) -> None:
    for _ in range(2):
        document.add_paragraph()
    if LOGO_PATH.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.2))
    add_centered_paragraph(document, INSTITUTION, bold=True, size=15, space_after=20)
    document.add_paragraph()
    add_centered_paragraph(document, TITLE, bold=True, size=15, space_after=18)
    document.add_paragraph()
    document.add_paragraph()
    add_centered_paragraph(document, AUTHOR, bold=True, size=14, space_after=8)
    add_centered_paragraph(document, f"INDEX NO.: {INDEX_NO}", size=12, space_after=12)
    add_centered_paragraph(
        document,
        f"A project in the {DEPARTMENT}, submitted to {INSTITUTION.title()} in partial fulfillment "
        "of the requirement for the award of the Degree of",
        size=12,
        space_after=22,
    )
    add_centered_paragraph(document, DEGREE, bold=True, size=14, space_after=24)
    add_centered_paragraph(document, f"SUPERVISED BY {SUPERVISOR}", size=12, space_after=10)
    add_centered_paragraph(document, SUBMISSION_DATE, size=12)
    add_page_break(document)


def build_front_matter(document: Document) -> None:
    add_heading_centered(document, "ABSTRACT", style_name="Heading 1")
    abstract_paragraphs = [
        "This project presents the design and implementation of a role-based learning management and rehabilitation support platform for the Ghana Prisons Service. The study is motivated by persistent challenges in correctional education in Ghana, including limited classroom infrastructure, insufficient instructional coverage, fragmented inmate records, weak attendance monitoring, and the absence of integrated digital tools for rehabilitation planning. To respond to these constraints, the study conceptualizes an Internet of Things enabled learning ecosystem in which secure identity verification, device allocation, classroom access control, and digital course delivery operate within one coordinated software environment.",
        "The developed software prototype, named the Ghana Prisons Learning Portal, was implemented with Next.js, TypeScript, Tailwind CSS, Prisma-backed repositories, and role-aware client and server session controls. The platform provides separate experiences for administrators, management staff, lecturers, inmates, and clocking officers. Core capabilities include inmate registration with biometric capture readiness, role-specific authentication, attendance and learning access workflows, lecturer course and assignment operations, administrative dashboards, management analytics, and a clock-in module that supports biometric and camera-based verification with resilient fallback behavior.",
        "A design and development research approach was adopted. Requirements were synthesized from literature on prison rehabilitation, digital learning, and IoT-enabled monitoring, then translated into portal modules and tested using scenario-based functional validation. The final system successfully supported core user journeys such as inmate login, dashboard access, course browsing, lecturer instructional monitoring, administrative inmate onboarding, management review, and device-allocation workflows. Code quality checks and deployment verification confirmed that the application builds successfully and is available online for demonstration at the production URL.",
        "The study concludes that IoT-oriented digital coordination can improve visibility, accountability, and educational continuity within correctional environments when designed with strong role separation, secure access handling, and practical fallback mechanisms. Although the present work is a software prototype rather than a full institutional rollout, it establishes a practical implementation foundation for prison education modernization in Ghana. The project therefore contributes both a working product and a defensible implementation framework for technology-supported rehabilitation in correctional settings.",
    ]
    for paragraph in abstract_paragraphs:
        add_body_paragraph(document, paragraph)
    add_body_paragraph(
        document,
        "Keywords: Ghana Prisons Service, Internet of Things, learning management system, rehabilitation, prison education, biometrics, dashboard analytics.",
        italic=True,
        indent_first=False,
    )
    add_page_break(document)

    add_heading_centered(document, "STUDENT'S DECLARATION", style_name="Heading 1")
    add_body_paragraph(
        document,
        f"I, {AUTHOR}, declare that this thesis, with the exception of quotations and references contained in published works which have all been identified and duly acknowledged, is entirely my own original work and has not been submitted, either in part or whole, for another degree elsewhere.",
    )
    add_manual_list(
        document,
        [
            f"Candidate's Name: {AUTHOR}",
            "Signature: __________________________    Date: __________________________",
            f"Supervisor's Name: {SUPERVISOR}",
            "Supervisor Signature: __________________________    Date: __________________________",
            "Head of Department: __________________________",
            "HOD Signature: __________________________    Date: __________________________",
            "Dean of Academics: __________________________",
            "Dean's Signature: __________________________    Date: __________________________",
        ],
    )
    add_page_break(document)

    add_heading_centered(document, "SUPERVISOR'S DECLARATION", style_name="Heading 1")
    add_body_paragraph(
        document,
        f"I hereby declare that the preparation and presentation of this work were supervised in accordance with the thesis supervision guidelines of {INSTITUTION}. {SUPERVISOR}, {DEPARTMENT} (Supervisor).",
        indent_first=False,
    )
    add_manual_list(document, ["SIGNATURE: __________________________", "DATE: __________________________"])
    add_page_break(document)

    add_heading_centered(document, "DEDICATION", style_name="Heading 1")
    add_body_paragraph(
        document,
        "This thesis is dedicated to the officers, educators, rehabilitation practitioners, and incarcerated learners whose pursuit of transformation through education continues to inspire correctional reform in Ghana.",
        indent_first=False,
    )
    add_page_break(document)

    add_heading_centered(document, "ACKNOWLEDGEMENTS", style_name="Heading 1")
    ack_paragraphs = [
        "I thank God for strength, wisdom, and guidance throughout this research and product development process.",
        f"My sincere appreciation goes to my supervisor, {SUPERVISOR}, for the academic direction, critical review, and professional encouragement that shaped this work.",
        "I also acknowledge the lecturers and staff of the Department of Information Technology, BlueCrest University College, for the knowledge base and support environment that made this study possible.",
        "Finally, I appreciate my family, friends, and all supporters of this research for their patience, counsel, and encouragement during the preparation of this thesis.",
    ]
    for paragraph in ack_paragraphs:
        add_body_paragraph(document, paragraph)
    add_page_break(document)

    add_heading_centered(document, "TABLE OF CONTENTS", style_name="Heading 1")
    add_toc_field(document)
    add_page_break(document)

    add_heading_centered(document, "LIST OF TABLES", style_name="Heading 1")
    add_generated_style_list_field(document, "Table Caption")
    add_page_break(document)

    add_heading_centered(document, "LIST OF FIGURES", style_name="Heading 1")
    add_generated_style_list_field(document, "Figure Caption")
    add_page_break(document)

    add_heading_centered(document, "LIST OF ABBREVIATIONS", style_name="Heading 1")
    for acronym, meaning in ABBREVIATIONS:
        add_manual_list(document, [f"{acronym} - {meaning}"])
    add_page_break(document)


def build_chapter_one(document: Document) -> None:
    add_heading_centered(document, "CHAPTER ONE")
    add_heading_centered(document, "INTRODUCTION TO THE STUDY", style_name="Heading 1")

    add_subheading(document, "1.1 Background to the Study")
    background_paragraphs = [
        "Correctional institutions are no longer viewed only as centers of punishment and incapacitation. Contemporary correctional policy increasingly emphasizes rehabilitation, reintegration, skills development, and educational support as practical means of reducing reoffending and supporting public safety. Within this shift, Internet of Things (IoT) technologies have become increasingly relevant because they enable connected monitoring, controlled access, automated record keeping, and data-driven educational support in environments where traditional delivery models are strained (Atzori et al., 2010; Al-Fuqaha et al., 2015; Davis et al., 2013).",
        "Research on correctional education consistently links access to structured learning with better reintegration outcomes, improved employability, and lower recidivism after release. RAND's meta-analysis, for example, found strong evidence that correctional education reduces reincarceration and improves employment prospects after release, reinforcing the strategic value of prison education as a public-safety intervention rather than a marginal welfare activity (Davis et al., 2013; Wilson et al., 2000).",
        "Within Ghana, the Ghana Prisons Service operates under a dual responsibility: secure custody and the welfare-driven rehabilitation of inmates. In practice, however, prison education remains constrained by overcrowding, limited teaching capacity, weak records management, inconsistent access to classrooms and workshops, and the absence of integrated digital tools for monitoring attendance, course participation, and rehabilitation progress. These limitations make it difficult to scale learning opportunities and to measure the effectiveness of interventions reliably (Boateng, 2021; Gbazuo and Ofori-Dua, 2021).",
        "Official institutional evidence sharpens this concern. The Ghana Prisons Service strategic plan identifies congestion, infrastructure pressure, and weak information systems as operational obstacles, while World Prison Brief data continues to show crowding pressure within Ghana's prison estate. At the same time, the Service's own reports on inmate BECE performance show that educational aspiration exists inside the prisons, making the absence of integrated digital learning support a missed reform opportunity rather than a hypothetical future problem (Ghana Prisons Service, 2015, 2024; World Prison Brief, 2025).",
        "IoT provides a focused conceptual response to these challenges because it links devices, identities, access events, and learning-state data within one monitored environment. In a prison education context, these capabilities can support secure access management, attendance verification, digital content delivery, learning analytics, and equipment allocation without depending entirely on paper workflows or physical oversight. When implemented with appropriate safeguards, IoT-oriented prison systems can improve accountability while also expanding access to prison education and rehabilitation services (Shaikh and Zeadally, 2016; UNODC, 2015).",
        "This study therefore focuses on the design and implementation of a Ghana Prisons Learning Portal as a practical software product that models how IoT-oriented prison education can work in Ghana. The portal integrates role-based login, inmate access control, dashboard analytics, biometric readiness workflows, lecturer operations, and inmate learning interfaces into one coordinated system. Rather than treating education, attendance, and rehabilitation monitoring as disconnected administrative tasks, the system brings them together within a single digital environment.",
        "The relevance of this work extends beyond software convenience. Ghana’s correctional modernization agenda requires systems that can produce usable information for policy, support structured learning delivery, and maintain secure operational boundaries. A role-aware prison LMS offers a realistic starting point for that transition because it combines identity control, educational workflows, and operational reporting in a way that aligns with correctional governance and rehabilitation goals.",
    ]
    for paragraph in background_paragraphs:
        add_body_paragraph(document, paragraph)

    add_table(
        document,
        "Table 1.1: Summary of correctional education challenges addressed by the proposed system",
        ["Observed Challenge", "Effect on Rehabilitation", "System Response"],
        [
            ["Manual inmate and learning records", "Fragmented reporting and slow follow-up", "Centralized digital dashboards and managed records"],
            ["Limited monitoring of attendance and device usage", "Weak accountability for educational participation", "Clock-in, attendance, and access workflows with verification states"],
            ["Insufficient access to learning content", "Irregular engagement with courses", "Inmate portal with course catalog, progress, and certificates"],
            ["Minimal instructional oversight", "Lecturers cannot track progress efficiently", "Lecturer dashboard for courses, assignments, submissions, and attendance"],
            ["Weak management visibility", "Poor evidence for policy and resource decisions", "Management analytics, executive summaries, and downloadable data views"],
        ],
    )

    add_subheading(document, "1.2 Statement of the Problem")
    add_body_paragraph(
        document,
        "Despite the acknowledged importance of rehabilitation, prison education in Ghana still depends heavily on fragmented manual procedures, isolated record keeping, and limited classroom capacity. These constraints reduce access to learning, weaken monitoring of inmate participation, and make it difficult for administrators to track progress or demonstrate outcomes. In addition, the lack of integrated digital access control exposes prison learning operations to inefficiency because identity verification, attendance recording, course delivery, and reporting occur as separate activities rather than as one coordinated process (Boateng, 2021; UNODC, 2018).",
    )
    add_body_paragraph(
        document,
        "A second problem is analytical rather than purely operational. Existing discussions on prison education in Ghana describe the need for reform, but there is still limited product-level evidence showing how IoT can be translated into an actual prison education platform for Ghana's institutional context. Much of the literature is either conceptual, policy-focused, or based on foreign systems. This creates a research gap between what correctional education literature recommends and what has been concretely implemented for the Ghana Prisons Service (Boateng, 2021; Edovo, 2025; UNODC, 2015).",
    )
    add_body_paragraph(
        document,
        "Without a secure and role-aware digital platform, the Ghana Prisons Service risks underutilizing educational interventions that could improve inmate development and reintegration. There is therefore a practical need for a software solution that brings together inmate onboarding, role-based authentication, learning delivery, attendance verification, lecturer operations, and management oversight in a way that reflects the connected and data-driven logic of IoT systems while remaining tailored to Ghanaian prison realities (Atzori et al., 2010; Al-Fuqaha et al., 2015; Boateng, 2021).",
    )

    add_subheading(document, "1.3 Aim of the Study")
    add_body_paragraph(
        document,
        "The aim of this study is to design and implement an IoT-oriented learning management and rehabilitation support platform that enhances education delivery, monitoring, and administrative coordination within the Ghana Prisons Service.",
    )

    add_subheading(document, "1.4 Research Objectives")
    add_body_paragraph(document, "The specific objectives of the study are:", indent_first=False)
    add_bullets(
        document,
        [
            "To examine the current educational, rehabilitation, and operational challenges within the Ghana Prisons Service and derive the system requirements needed for a digital intervention.",
            "To design and implement an IoT-oriented learning management and rehabilitation support platform that addresses the identified prison education and access-control requirements.",
            "To evaluate the developed platform using scenario-based functional testing, interface evidence, and deployment verification.",
        ],
    )

    add_subheading(document, "1.5 Research Questions")
    add_numbered(
        document,
        [
            "What current educational, rehabilitation, and operational challenges within the Ghana Prisons Service justify the need for a digital prison learning platform?",
            "How can an IoT-oriented learning management and rehabilitation support platform be designed and implemented to address those identified needs?",
            "How effective is the developed platform when assessed through functional validation and deployment readiness checks?",
        ],
    )

    add_table(
        document,
        "Table 1.2: Alignment of research objectives and research questions",
        ["Objective", "Matching Research Question"],
        [
            [
                "Objective 1: Examine current prison education and rehabilitation challenges and derive system requirements",
                "Research Question 1: What current educational, rehabilitation, and operational challenges justify the need for a digital prison learning platform?",
            ],
            [
                "Objective 2: Design and implement an IoT-oriented prison learning platform",
                "Research Question 2: How can an IoT-oriented platform be designed and implemented to address the identified needs?",
            ],
            [
                "Objective 3: Evaluate the developed platform",
                "Research Question 3: How effective is the developed platform when assessed through validation and deployment checks?",
            ],
        ],
    )

    add_subheading(document, "1.6 Significance of the Study")
    significance_paragraphs = [
        "The study is significant because it provides a product-centered contribution to correctional education discourse in Ghana. Instead of discussing digital transformation in abstract terms only, it delivers a working software artifact that demonstrates how educational access, inmate monitoring, and role-specific prison workflows can be coordinated in practice (Boateng, 2021; Hevner et al., 2004).",
        "The study is also valuable to prison administrators and policy makers because it frames technology adoption around operational needs that are directly measurable: inmate registration, identity handling, attendance events, course participation, management reporting, and lecturer visibility (Ghana Prisons Service, 2015; UNODC, 2018).",
        "For researchers and software developers, the project offers a practical implementation reference for designing secure multi-role systems in constrained institutional environments. For the Ghana Prisons Service, it establishes a foundation for future pilot deployment, evaluation, and integration with broader rehabilitation and vocational initiatives (Atzori et al., 2010; Pressman and Maxim, 2019).",
    ]
    for paragraph in significance_paragraphs:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "1.7 Scope of the Study")
    add_body_paragraph(
        document,
        "The study covers the analysis, design, development, and validation of a web-based prototype for prison education and rehabilitation support. The implemented scope includes role-based access for administrators, lecturers, management staff, inmates, and clocking officers; inmate onboarding; attendance and access workflows; course and assignment visibility; dashboard reporting; and production deployment for live demonstration. The study does not include a full nationwide rollout, integration with physical prison hardware, or longitudinal measurement of actual recidivism outcomes (Hevner et al., 2004; Pressman and Maxim, 2019).",
    )

    add_subheading(document, "1.8 Limitations of the Study")
    add_body_paragraph(
        document,
        "The current work is limited to a software prototype and controlled demonstration environment. Real biometric hardware integration, facility-wide networking constraints, institutional procurement issues, and live prison deployment governance were outside the present implementation. In addition, the study evaluates functional readiness rather than long-term behavioral or educational outcomes among actual inmate populations (UNODC, 2015; World Prison Brief, 2025).",
    )

    add_subheading(document, "1.9 Definition of Key Terms")
    add_bullets(
        document,
        [
            "Internet of Things: A connected digital environment in which devices, services, and identifiers exchange data to support monitored or automated processes.",
            "Learning Management System: A software platform used to organize educational content, learners, assessments, progress records, and instructional workflows.",
            "Correctional Education: Formal, non-formal, or vocational learning activities provided within prisons to support rehabilitation and reintegration.",
            "Biometric Verification: Identity confirmation through measurable personal characteristics or device-supported authentication methods.",
            "Role-Based Access Control: A system design approach in which user permissions depend on institutional role and authorized responsibilities.",
        ],
    )

    add_subheading(document, "1.10 Organization of the Study")
    add_body_paragraph(
        document,
        "Chapter One introduces the study, its background, problem statement, objectives, and scope. Chapter Two reviews literature on correctional education, IoT adoption, prison digital systems, and rehabilitation-oriented technology. Chapter Three presents the system analysis, design choices, architecture, and development methodology. Chapter Four discusses implementation outputs, screenshots, code excerpts, and system validation results. Chapter Five summarizes the study, draws conclusions, and presents recommendations for deployment and future research.",
    )
    add_page_break(document)


def build_chapter_two(document: Document) -> None:
    reference_sections = load_reference_chapter_sections()
    add_heading_centered(document, "CHAPTER TWO")
    add_heading_centered(document, "LITERATURE REVIEW", style_name="Heading 1")

    add_subheading(document, "2.0 Introduction")
    add_body_paragraph(
        document,
        "This chapter reviews related works in the same order as the study objectives. Section 2.1 reviews literature related to Objective One by examining the present educational, rehabilitation, and operational conditions that create the need for a prison learning platform in Ghana. Section 2.2 reviews literature related to Objective Two by focusing on IoT, prison EdTech, learning management systems, biometric access, and dashboard-oriented design approaches. Section 2.3 reviews literature related to Objective Three by examining evaluation, adoption, ethics, security, and deployment readiness issues that determine whether a prison LMS can function responsibly and effectively (Davis et al., 2013; UNODC, 2015; Edovo, 2025).",
        indent_first=False,
    )

    add_subheading(document, "2.1 Review of Related Works for Objective One")
    add_subheading(document, "2.1.1 Correctional Education, Rehabilitation, and Reintegration", level=3)
    for paragraph in [
        "Correctional education has been widely recognized as one of the most effective non-custodial strategies for improving post-release outcomes. Davis et al. (2013) and Wilson et al. (2000) show that inmates who participate in structured educational programs are less likely to reoffend and more likely to transition into productive community life after release. This makes education not merely a welfare activity within prisons, but a strategic public safety intervention.",
        "Within Ghana, prison education is expected to support literacy, vocational preparedness, behavioral transformation, and the reintegration mandate of the Ghana Prisons Service. Yet Ghanaian and African prison studies continue to report major barriers, including overcrowding, weak infrastructure, inconsistent instructional coverage, and limited access to modern educational resources (Boateng, 2021; Kinyanjui, 2019; Yeboah and Adomako, 2022).",
        "Digital correctional programming has also become more prominent in operational practice. Edovo reports secure educational access across more than 1,200 facilities and one million learners, arguing that tablet-enabled programming reduces the gap created by staffing shortages, security constraints, and limited classroom access (Edovo, 2025). These developments reinforce the view that prison education is moving toward digitally coordinated delivery rather than purely manual classroom administration.",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.1.2 Current Educational and Rehabilitation Conditions in Ghana Prisons", level=3)
    for paragraph in reference_sections["objective_one"][:10]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.1.3 Ghanaian Institutional Evidence and Digital Reform Gap", level=3)
    for paragraph in [
        "Recent Ghana Prisons Service evidence shows that prison education in Ghana is not dormant; rather, it is unevenly supported. The Service's report on inmate BECE performance demonstrates continuing participation in formal education, but this progress sits alongside institutional planning documents that identify infrastructure deficits, congestion, and limited information systems capacity. The problem, therefore, is not the absence of educational intent but the weakness of the supporting digital and administrative environment (Ghana Prisons Service, 2015, 2024).",
        "This distinction is analytically important. If prisons already host educational and vocational activities, then the real research question becomes how these activities can be coordinated, monitored, and scaled under restrictive operational conditions. World Prison Brief data reinforces that resource pressure remains a structural issue in Ghana, which means that any proposed solution must reduce administrative friction rather than assume abundant staff time, classroom space, or unrestricted internet access (World Prison Brief, 2025; Boateng, 2021).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)
    add_figure(
        document,
        FigureSpec("2.2", "Ghana prison education reform context and digital pressure points", "chapter2-ghana-context.png", 6.4),
        source="Author's construct, 2026, informed by Ghana Prisons Service (2015, 2024), Boateng (2021), and World Prison Brief (2025).",
    )
    add_figure(
        document,
        FigureSpec("2.3", "Official Ghana Prisons Service evidence of inmate academic participation", "chapter2-gps-bece.jpg", 6.4),
        source="Ghana Prisons Service (2024), Prison inmates set another record in 2023 BECE.",
    )

    add_subheading(document, "2.1.4 Requirement Themes Emerging from the Literature", level=3)
    for paragraph in [
        "The literature related to Objective One points to several recurring requirement themes: centralized inmate and course records, verifiable attendance capture, controlled movement into learning spaces, lecturer and management visibility, and stronger evidence for rehabilitation reporting (Boateng, 2021; UNODC, 2018; Edovo, 2025).",
        "These requirement themes imply that the solution space is not limited to content delivery alone. A prison learning platform must respond simultaneously to educational access, institutional accountability, and operational security. Accordingly, Objective One concludes that prison EdTech for Ghana must be requirement-driven and role-specific rather than generic (Davis et al., 2013; UNODC, 2015).",
        "A further synthesis point emerges when Ghanaian and international sources are read together: prison education technology succeeds when it reduces coordination costs for staff while preserving a meaningful learner experience. This dual requirement explains why the present study combines inmate-facing learning pages with administrator, lecturer, management, and clocking-officer workflows instead of treating the LMS as a single-role classroom system (Ithaka S+R, 2024; Massachusetts Department of Correction, 2024; Ghana Prisons Service, 2024).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_table(
        document,
        "Table 2.1: Comparative view of prison education conditions and digital reform directions",
        ["Jurisdiction / Source", "Observed Challenge or Practice", "Implication for Ghana"],
        [
            ["Ghana prison education literature", "Overcrowding, fragmented records, limited modern tools", "A prison LMS must solve educational and administrative bottlenecks together"],
            ["RAND correctional education evidence", "Education improves reentry outcomes and lowers reincarceration risk", "Educational investment should be treated as strategic rehabilitation infrastructure"],
            ["Edovo secure-tablet model", "Digital delivery expands access where in-person programming is limited", "Secure digital delivery can widen prison learning participation"],
            ["UNODC Mandela Rules", "Education and humane treatment remain core correctional obligations", "Any digital reform must preserve dignity, rights, and rehabilitation purpose"],
        ],
        source="Author's compilation from Davis et al. (2013), Boateng (2021), Edovo (2025), and UNODC (2015).",
    )
    add_figure(
        document,
        FigureSpec("2.1", "Problem-to-requirement map for prison education digitalization", "chapter1-problem-map.png", 6.4),
        source="Author's construct, 2026, informed by Boateng (2021), Gbazuo and Ofori-Dua (2021), and UNODC (2015).",
    )

    add_subheading(document, "2.2 Review of Related Works for Objective Two")
    add_subheading(document, "2.2.1 The Concept of Internet of Things in Prison Education", level=3)
    for paragraph in [
        "The Internet of Things refers broadly to environments in which connected devices, identifiers, software services, and communication networks exchange data in order to monitor, automate, or improve real-world processes (Atzori et al., 2010; Al-Fuqaha et al., 2015). In educational settings, recent review literature shows that connected technologies improve learning visibility, strengthen analytics, and make it easier to adapt content delivery to user context and device state (Ball, 2021; Palanci et al., 2024).",
        "For correctional education, IoT-inspired design means more than attaching hardware to classrooms. It means treating identity, access events, course participation, device allocation, and environmental conditions as connected parts of one controlled workflow. This interpretation is especially relevant to prisons, where the boundaries between education, monitoring, and operations are tighter than in mainstream schools (Shaikh and Zeadally, 2016; UNODC, 2015).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)
    for paragraph in reference_sections["objective_two_intro"][:12]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.2.2 Learning Management Systems, EdTech, and Smart Classroom Design", level=3)
    for paragraph in [
        "Learning management systems provide structured delivery for content, assessments, communication, and progress tracking. In correctional environments, however, the LMS must do more than distribute content. It must respect role boundaries, support limited or controlled connectivity, protect privacy, and align educational access with institutional security procedures (Champion and Edgar, 2013; Edovo, 2025).",
        "Recent digital-learning literature also points to analytics, adaptive feedback, and visible learner-state monitoring as core design benefits. These functions are especially relevant to incarcerated learners whose educational backgrounds, stress levels, and readiness for learning may vary significantly within the same cohort, making static one-size-fits-all course administration especially weak in prison settings (Ball, 2021; Palanci et al., 2024).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.2.3 Management, Instructional, and Inmate Learning Levels", level=3)
    for paragraph in reference_sections["objective_two_intro"][12:]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.2.4 Designing an IoT-Enabled Learning System for Prisons", level=3)
    for paragraph in reference_sections["objective_two_design"][:18]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.2.5 Usability, Access Control, and Role Separation", level=3)
    for paragraph in [
        "Usability remains important even in highly restricted environments. If prison staff or inmates struggle to navigate a digital platform, the system may introduce new barriers instead of removing old ones. Nielsen's usability principles therefore remain relevant: visibility of system status, consistency, error prevention, and simple user control all matter in prison software because users operate under time pressure and institutional rules (Nielsen, 1994).",
        "Biometric and access-control literature also reinforces the need for role-aware identity handling. Correctional systems require accountable access to learning spaces, device allocation, and attendance data, but such controls must remain proportionate and linked to legitimate educational or operational purposes rather than becoming unbounded surveillance (UNODC, 2015; Ghana Data Protection Act, 2012).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.2.6 Critical Synthesis of Recent Prison Technology Literature", level=3)
    for paragraph in [
        "Recent prison-technology studies add a useful corrective to overly optimistic accounts of digital reform. Ithaka S+R's national survey of higher education in prison programs found that device access remains uneven, LMS availability is inconsistent, and internet access is still tightly constrained. This suggests that platform design for prisons must assume controlled connectivity, staged permissions, and offline-tolerant workflow logic rather than copying open-campus LMS assumptions directly (Ithaka S+R, 2024).",
        "Practical deployment examples reinforce the same lesson. Massachusetts Department of Correction's tablet initiative illustrates how digital learning can be expanded inside correctional systems when device governance, approved content, and rehabilitative purpose are aligned institutionally. Nisser et al. (2024) similarly show that prison-based digital learning works best when the learning design is intentionally adapted to the carceral setting rather than merely transferred from general education environments.",
        "These studies are valuable, but they do not solve the Ghanaian problem automatically. Most are based on non-Ghanaian institutions with different device policies, funding models, and correctional governance histories. The present study therefore uses them critically: as design signals for role separation, secure content delivery, and learner support, but not as templates to be copied without local adaptation (Boateng, 2021; Ghana Prisons Service, 2015; World Prison Brief, 2025).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)
    add_figure(
        document,
        FigureSpec("2.5", "Adoption barriers and design responses for prison education technology", "chapter2-adoption-barriers.png", 6.4),
        source="Author's construct, 2026, based on Ithaka S+R (2024), Massachusetts Department of Correction (2024), Nisser et al. (2024), and Ghana-focused prison literature.",
    )

    add_table(
        document,
        "Table 2.2: Related design literature and its contribution to Objective Two",
        ["Literature Area", "Key Insight", "Design Relevance to This Study"],
        [
            ["IoT architecture literature", "Connected devices produce real-time operational visibility", "Supports attendance, access, and device-state workflows"],
            ["Prison EdTech literature", "Secure digital learning expands access where classrooms are constrained", "Supports inmate portal and lecturer-driven content delivery"],
            ["Biometric and access-control literature", "Identity-linked access improves accountability", "Supports clock-in and verification modules"],
            ["Usability literature", "Simple, predictable interfaces reduce user error", "Supports role-specific dashboards and guided prison workflows"],
        ],
        source="Author's synthesis from Atzori et al. (2010), Al-Fuqaha et al. (2015), Edovo (2025), Nielsen (1994), and the reference Chapter Two draft.",
    )
    add_figure(
        document,
        FigureSpec("2.4", "Literature synthesis framework linking IoT, LMS, security, and rehabilitation", "chapter2-literature-framework.png", 6.4),
        source="Author's construct, 2026, based on Atzori et al. (2010), Al-Fuqaha et al. (2015), Davis et al. (2013), Edovo (2025), and UNODC (2015).",
    )
    add_table(
        document,
        "Table 2.4: Recent literature signals for prison education technology design",
        ["Recent Source", "Main Finding", "Implication for the Ghana Prison LMS"],
        [
            ["Ithaka S+R (2024)", "Technology access in prison education remains uneven and connectivity is highly restricted", "The LMS should support tightly controlled workflows and avoid assumptions of open internet access"],
            ["Massachusetts Department of Correction (2024)", "Tablet deployment can strengthen educational programming when managed centrally", "Device allocation and approved access should be treated as part of the educational workflow"],
            ["Nisser et al. (2024)", "Digitally mediated prison courses can improve learner self-efficacy when adapted to prison realities", "Inmate-facing learning design should remain simple, structured, and confidence-building"],
            ["Palanci et al. (2024)", "Learning analytics strengthens insight into participation and progression", "Dashboards and progress views should support both monitoring and instructional decision-making"],
        ],
        source="Author's synthesis from recent prison education and digital-learning literature.",
    )

    add_subheading(document, "2.3 Review of Related Works for Objective Three")
    add_subheading(document, "2.3.1 Infrastructure, Security, and Ethical Considerations", level=3)
    for paragraph in reference_sections["objective_three"][:12]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.3.2 Evaluation, Learner Analytics, and Risk Assessment", level=3)
    for paragraph in reference_sections["objective_three"][12:24]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.3.3 Adoption, Usability, and Deployment Readiness", level=3)
    for paragraph in [
        "Objective Three requires evidence that the developed system performs its intended functions. Scenario-based testing is appropriate when a system's value depends on end-to-end workflows rather than isolated algorithms alone, which is exactly the case in prison platforms where the right user must reach the right task at the right stage of a controlled process (Pressman and Maxim, 2019; Sommerville, 2016).",
        "A deployed application offers stronger final-defense evidence than a development-only artifact. Deployment demonstrates that the system can run outside the local editor, that route structures are coherent, and that the product is presentable to evaluators in a live environment. In web-based thesis projects, deployment readiness therefore becomes part of the product evaluation story itself (Sommerville, 2016; Nielsen, 1994).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)
    for paragraph in reference_sections["objective_three"][24:]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.3.4 Governance and Deployment Readiness in the Ghanaian Context", level=3)
    for paragraph in [
        "Adoption readiness in Ghana must be discussed in institutional rather than purely technical terms. The Ghana Prisons Service strategic direction emphasizes humane custody, rehabilitation, and system modernization, but these priorities operate under pressure from congestion and infrastructure constraints. As a result, a prison LMS is only realistic if it is simple to govern, clear in role boundaries, and defensible as a rehabilitative support tool rather than a speculative technology project (Ghana Prisons Service, 2015; World Prison Brief, 2025).",
        "This governance perspective also clarifies the evaluation logic of the present thesis. The question is not whether the software solves every correctional challenge, but whether it provides a credible, locally relevant model that can support pilot discussion, policy review, and future ethical testing. That is why deployment verification, security boundaries, and workflow completeness matter so strongly in Objective Three (UNODC, 2015; Pressman and Maxim, 2019; Boateng, 2021).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)
    add_figure(
        document,
        FigureSpec("2.8", "Official visual identity of a Ghana Prisons rehabilitation initiative", "chapter2-cits-logo.png", 3.4),
        source="Ghana Prisons Service rehabilitation page, retrieved from the official CITS project page.",
    )

    add_table(
        document,
        "Table 2.3: Literature themes relevant to Objective Three evaluation",
        ["Evaluation Dimension", "Literature Emphasis", "Implication for the Present Study"],
        [
            ["Functional effectiveness", "Scenario-based validation of complete user journeys", "The platform must be checked through role workflows rather than isolated screens"],
            ["System quality", "Correctness, reliability, and build consistency matter", "Linting and build verification support technical credibility"],
            ["Security and ethics", "Digital prison systems must protect privacy and stay purpose-bound", "Evaluation must consider safeguards as well as features"],
            ["Adoption readiness", "Infrastructure, training, and policy fit determine sustainability", "Deployment discussions must include governance and rollout feasibility"],
        ],
        source="Author's synthesis from Pressman and Maxim (2019), Sommerville (2016), Nielsen (1994), UNODC (2015), and the reference Chapter Two draft.",
    )
    add_figure(
        document,
        FigureSpec("2.6", "Evaluation dimensions for prison LMS adoption and validation", "chapter4-validation-pipeline.png", 6.4),
        source="Author's construct, 2026, based on Pressman and Maxim (2019), Sommerville (2016), Nielsen (1994), and UNODC (2015).",
    )

    add_subheading(document, "2.4 Theoretical Framework")
    add_subheading(document, "2.4.1 Technology Acceptance Model and Diffusion of Innovation", level=3)
    for paragraph in [
        "A clear theoretical framework strengthens the explanatory value of this study. The Technology Acceptance Model (TAM) explains technology adoption through perceived usefulness and perceived ease of use, both of which are directly relevant to prison staff and inmate-facing educational technologies. If officers, lecturers, or inmates perceive the system as difficult to use or insufficiently useful, adoption is likely to weaken regardless of technical capability (Davis, 1989; Nielsen, 1994).",
        "Diffusion of Innovation theory also supports the study because prison systems adopt technology under institutional constraints. Relative advantage, compatibility, complexity, trialability, and observability help explain why some prison technology initiatives scale while others stall. In the Ghanaian context, compatibility with prison routines, security expectations, and resource realities is especially important for any IoT-oriented platform (Rogers, 2003; Boateng, 2021).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.4.2 Rehabilitation Theory and Correctional Relevance", level=3)
    for paragraph in [
        "Rehabilitation theory provides the correctional lens for the study. The value of the proposed platform is not in digitalization for its own sake, but in whether connected educational processes can improve inmate learning continuity, skill development, and reintegration readiness. This aligns with the rehabilitative spirit of the Mandela Rules and with evidence linking correctional education to improved post-release outcomes (UNODC, 2015; Davis et al., 2013).",
        "The combined use of TAM, Diffusion of Innovation, and rehabilitation theory allows the study to explain adoption, design fit, and correctional purpose at the same time. Together these theories justify a platform that is useful, easy to navigate, institutionally compatible, and rehabilitation-oriented.",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)
    add_figure(
        document,
        FigureSpec("2.7", "Conceptual model of the proposed IoT prison education framework", "chapter2-conceptual-model.png", 6.4),
        source="Author's conceptual model, 2026, informed by Davis (1989), Rogers (2003), UNODC (2015), and correctional education literature.",
    )

    add_subheading(document, "2.5 Research Gap and Conceptual Direction")
    for paragraph in [
        "The reviewed literature establishes the value of correctional education, the promise of connected technologies, and the importance of secure digital learning environments. However, a clear gap remains in locally grounded, product-level implementations for Ghana's prison education context. Much of the literature is either policy-oriented, internationally situated, or conceptual in nature (Boateng, 2021; UNODC, 2015; Edovo, 2025).",
        "This study addresses that gap by building a tangible, role-based software artifact tailored to Ghana Prisons Service workflows. The contribution is therefore not only theoretical. It demonstrates how requirement analysis, IoT-oriented platform design, and product evaluation can be aligned directly with three research objectives in one coherent thesis structure (Davis et al., 2013; Atzori et al., 2010; Pressman and Maxim, 2019).",
        "Its original contribution lies in the proposed conceptual linkage between prison education access, identity-aware workflow control, and managerial visibility within one locally grounded portal model. In other words, the study does not simply recommend digital reform; it specifies and implements a Ghana-relevant model for how prison education can be supported by connected access logic, modular dashboards, and evaluation-ready deployment practice (Hevner et al., 2004; Rogers, 2003; Ghana Prisons Service, 2015).",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "2.6 Chapter Summary")
    add_body_paragraph(
        document,
        "Chapter Two has reviewed related works in the same order as the three study objectives. Section 2.1 established the prison education and rehabilitation problem context. Section 2.2 examined the design and implementation literature relevant to an IoT-oriented prison learning platform. Section 2.3 reviewed evaluation, adoption, ethics, and deployment readiness considerations. These insights provide the direct methodological basis for Chapter Three (UNODC, 2015; Davis et al., 2013; Edovo, 2025).",
        indent_first=False,
    )
    add_page_break(document)


def build_chapter_three(document: Document) -> None:
    add_heading_centered(document, "CHAPTER THREE")
    add_heading_centered(document, "SYSTEM ANALYSIS, DESIGN AND METHODOLOGY", style_name="Heading 1")

    add_subheading(document, "3.0 Introduction")
    add_body_paragraph(
        document,
        "This chapter presents the methodology of the study in the same order as the three research objectives. Section 3.1 explains how Objective One was addressed through requirement analysis and problem synthesis. Section 3.2 explains how Objective Two was addressed through system design and implementation. Section 3.3 explains how Objective Three was addressed through product evaluation and deployment validation (Hevner et al., 2004; Creswell and Creswell, 2018).",
    )

    add_subheading(document, "3.1 Methodology for Objective One")
    add_subheading(document, "3.1.1 Requirement Elicitation Strategy", level=3)
    for paragraph in [
        "Objective One focused on understanding the existing educational and operational gaps that the software needed to solve. To achieve this, the study adopted a requirement synthesis strategy based on reviewed literature, prison education workflow assumptions, and the practical need to coordinate authentication, inmate records, course access, attendance tracking, and reporting (Hevner et al., 2004; Boateng, 2021).",
        "Rather than collecting unrelated feature ideas, requirement elicitation was structured around the correctional environment itself. The study asked what an administrator needs to control, what a lecturer needs to monitor, what an inmate needs to access, what a management officer needs to see, and what a clocking officer needs to verify. This role-centered requirement framing ensured that the final system would respond to real institutional tasks (UNODC, 2018; Ghana Prisons Service, 2015).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "3.1.2 Research Design and Data Sources", level=3)
    for paragraph in [
        "The overall study follows a design science and applied systems-development orientation. Objective One specifically relied on document analysis and problem synthesis rather than experimental data collection. The principal data sources at this stage were correctional education literature, prison digitalization literature, policy documents, implementation guidance, and the supplied domain draft materials relating to Ghana prison education and IoT adoption (Hevner et al., 2004; Boateng, 2021; UNODC, 2015).",
        "This design is justified because the study sought first to define the problem accurately before building a software response. A document-driven requirement analysis is appropriate when the institution under study has practical constraints that make immediate large-scale human-subject data collection difficult, but sufficient published and policy material exists to support grounded system design.",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "3.1.3 Requirement Analysis Output", level=3)
    add_table(
        document,
        "Table 3.1: Objective One requirement analysis summary",
        ["Problem Area", "Derived Requirement"],
        [
            ["Fragmented inmate learning records", "A centralized digital record and dashboard structure is required"],
            ["Poor access and attendance visibility", "Controlled access, clock-in, and attendance tracking workflows are required"],
            ["Weak lecturer monitoring", "Course, assignment, submission, and attendance views are required for lecturers"],
            ["Limited management oversight", "High-level analytics and reporting views are required for management"],
            ["Role confusion and security risk", "Strict role-based routing and session control are required"],
            ["Unstable hardware availability", "Camera and biometric workflows must include resilient fallbacks"],
        ],
    )
    for paragraph in [
        "The requirement analysis confirmed that the proposed system needed both educational and operational features. It was not sufficient to create a general LMS only; the system also needed onboarding, access verification, dashboards, and role separation to fit the prison context (Boateng, 2021; UNODC, 2015).",
        "The output of Objective One therefore became the blueprint for the application modules implemented under Objective Two (Hevner et al., 2004; Pressman and Maxim, 2019).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "3.2 Methodology for Objective Two")
    add_subheading(document, "3.2.1 Research and Development Approach", level=3)
    for paragraph in [
        "Objective Two focused on designing and implementing the actual platform. The study employed an applied design science orientation supported by iterative software development. This was appropriate because the goal was to move from identified prison-system requirements to a working product that demonstrates how those requirements can be satisfied (Hevner et al., 2004; Pressman and Maxim, 2019).",
        "An iterative process was used because prison education systems involve tightly connected modules. Authentication affects dashboard access, dashboard actions depend on role definitions, and inmate learning workflows depend on prior registration and access allocation. Iterative development made it possible to refine these dependencies while testing them incrementally in the local application environment (Sommerville, 2016; Pressman and Maxim, 2019).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "3.2.2 Population and Intended Pilot Sampling Logic", level=3)
    for paragraph in [
        "Although the present thesis centers on software design and implementation, the platform is intended for a correctional pilot involving multiple stakeholder categories. The relevant stakeholder population includes prison administrators, management officers, lecturers or instructors, clocking officers, and inmate learners. Because these users do not occupy equivalent roles, a purposive sampling logic is most appropriate for pilot validation rather than random selection (Creswell and Creswell, 2018).",
        "A justified pilot sample would include participants who directly interact with prison learning workflows: administrative officers responsible for inmate records, lecturers involved in educational delivery, officers linked to controlled access processes, and a small supervised inmate user group. The questionnaire instrument in the appendices was designed with this purposive sampling structure in mind. This allows the study to remain methodologically explicit even where the current implementation cycle stops at product completion and scenario-based testing.",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_subheading(document, "3.2.3 System Architecture", level=3)
    for paragraph in [
        "The implemented architecture follows a layered web application model. At the presentation layer, role-specific pages provide tailored interfaces for each user category. At the application layer, API routes manage authentication, attendance, courses, assignments, inmate records, and reporting actions. At the data layer, repositories and structured state objects coordinate seeded and persisted portal data. This arrangement supports separation of concerns while remaining practical for iterative thesis development (Sommerville, 2016; Pressman and Maxim, 2019).",
        "From an IoT perspective, the architecture emphasizes connected state transitions. A user’s identity determines the allowed route. A clock-in event influences inmate session access. A lecturer action updates course-facing data. A management dashboard aggregates records into decision-ready summaries. These are not independent screens; they are connected operational states coordinated through one platform (Atzori et al., 2010; Shaikh and Zeadally, 2016).",
    ]:
        add_body_paragraph(document, paragraph)

    add_table(
        document,
        "Table 3.2: Software tools and development technologies used in the implementation",
        ["Tool / Technology", "Purpose in the Study"],
        [
            ["Next.js (App Router)", "Frontend and route-based application structure"],
            ["TypeScript", "Strong typing, safer logic composition, and maintainable code"],
            ["Tailwind CSS", "Responsive interface styling and dashboard layout consistency"],
            ["Prisma / repositories", "Structured data access model for seeded and persisted records"],
            ["Playwright", "Browser-driven validation and screenshot capture"],
            ["Vercel", "Production deployment for live thesis demonstration"],
            ["python-docx", "Generation of the final thesis product document"],
        ],
    )
    add_figure(
        document,
        FigureSpec("3.1", "Layered system architecture of the proposed prison learning platform", "chapter3-system-architecture.png", 6.4),
        source="Author's construct, 2026, based on the implemented system architecture.",
    )

    add_subheading(document, "3.2.4 Core System Modules", level=3)
    add_table(
        document,
        "Table 3.3: Core system modules and corresponding users",
        ["Module", "Primary User(s)", "Purpose"],
        [
            ["Access Portal", "All users", "Direct users to the correct login route and role entry point"],
            ["Authentication and Session Layer", "All users", "Validate credentials, issue session state, and enforce role routing"],
            ["Admin Dashboard and Registration", "Administrators", "Manage inmates, view operational metrics, and capture onboarding data"],
            ["Lecturer Dashboard", "Lecturers", "Monitor learners, courses, assignments, and attendance"],
            ["Management Dashboard", "Management staff", "Review institutional analytics and high-level rehabilitation indicators"],
            ["Clock-in Module", "Clocking officers", "Verify inmate presence, allocate devices, and grant facility access"],
            ["Inmate Portal", "Inmates", "Access courses, progress, assigned learning resources, and certificates"],
        ],
    )

    add_subheading(document, "3.2.5 Development Procedure", level=3)
    add_numbered(
        document,
        [
            "Review literature and operational needs relevant to prison education and IoT-enabled monitoring.",
            "Define user roles, route structures, and core workflow dependencies.",
            "Design responsive interfaces for the public portal and role-specific dashboards.",
            "Implement authentication, session persistence, and guarded routes.",
            "Develop administrative, lecturer, inmate, management, and clock-in modules iteratively.",
            "Validate the application with functional scenario testing, screenshots, and deployment checks.",
        ],
    )

    add_subheading(document, "3.3 Methodology for Objective Three")
    add_subheading(document, "3.3.1 Evaluation Procedure", level=3)
    for paragraph in [
        "Objective Three focused on product evaluation. System validation was therefore centered on scenario-based functional correctness and deployability rather than large-sample user experimentation. Each major role workflow was exercised to confirm that the expected page rendered, the relevant data was available, and the application state aligned with the intended task (Pressman and Maxim, 2019; Sommerville, 2016).",
        "The evaluation combined browser-driven route validation, screenshot evidence, software quality checks, and live deployment confirmation. This made it possible to evaluate both visible usability and underlying technical readiness (Nielsen, 1994; Pressman and Maxim, 2019).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "3.3.2 Evaluation Criteria and Quality Indicators", level=3)
    for paragraph in [
        "The evaluation criteria for the implemented platform were derived from software quality and user-task considerations. These included role-based correctness, workflow completeness, interface clarity, fallback resilience, build success, and deployment readiness. This combination is appropriate because the study's outcome is a thesis product rather than a purely theoretical model (Pressman and Maxim, 2019; Sommerville, 2016).",
        "Where a later pilot evaluation is conducted, usability and perceived usefulness measures can be added for prison staff and inmate users in line with TAM constructs. This provides a clear bridge between the present implementation stage and a future institutional adoption study.",
    ]:
        add_body_paragraph(document, paragraph, indent_first=False)

    add_table(
        document,
        "Table 3.4: Objective Three evaluation plan",
        ["Evaluation Area", "Method Used"],
        [
            ["Role-based navigation", "Direct route access and live browser validation"],
            ["Functional coherence", "Scenario-based testing of admin, lecturer, management, inmate, and clock-in flows"],
            ["Visual evidence", "Screenshot capture from the running application"],
            ["Code quality", "Linting and production build verification"],
            ["Deployment readiness", "Live production URL confirmation"],
        ],
    )
    add_figure(
        document,
        FigureSpec("3.2", "Validation flow used for the implemented prison LMS", "chapter4-validation-pipeline.png", 6.4),
        source="Author's construct, 2026, based on the study evaluation process.",
    )

    add_subheading(document, "3.4 Ethical and Security Considerations")
    for paragraph in [
        "Because the study deals with prison education and simulated biometric workflows, ethical handling of identity and access concepts was important throughout the design. The implemented system therefore emphasizes role separation, controlled route access, client and server session management, and explicit fallbacks when live biometric hardware is not available (UNODC, 2015; Ghana Data Protection Act, 2012).",
        "The study does not expose or process real inmate biometric data in production. Instead, it models readiness, workflow states, and protected operational logic in a way suitable for academic demonstration. This allows the thesis to discuss prison digital transformation responsibly without claiming a live institutional data deployment that did not occur (UNODC, 2015; Hevner et al., 2004).",
        "If the next stage of the work involves stakeholder piloting or real institutional data, formal ethical clearance, informed consent procedures, and prison authority approval would be required. Particular care would be necessary in relation to power imbalance, privacy, and the handling of any learner-performance or biometric data in correctional settings (UNODC, 2015; Ghana Data Protection Act, 2012).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "3.5 Chapter Summary")
    add_body_paragraph(
        document,
        "This chapter has presented the methodology of the study in the same order as the three objectives. Section 3.1 explained how the problem and requirements were derived. Section 3.2 explained how the system was designed and implemented. Section 3.3 explained how the final platform was evaluated. Chapter Four now presents the results in that same sequence (Hevner et al., 2004; Pressman and Maxim, 2019).",
    )
    add_page_break(document)


def build_chapter_four(document: Document) -> None:
    add_heading_centered(document, "CHAPTER FOUR")
    add_heading_centered(document, "SYSTEM IMPLEMENTATION, RESULTS AND DISCUSSION", style_name="Heading 1")

    add_subheading(document, "4.0 Introduction")
    add_body_paragraph(
        document,
        "This chapter presents the results of the study in the same order as the research objectives. Section 4.1 presents the findings for Objective One by summarizing the derived prison education and rehabilitation requirements. Section 4.2 presents the results for Objective Two by showing the designed and implemented system. Section 4.3 presents the results for Objective Three by discussing validation, screenshots, and deployment readiness evidence (Pressman and Maxim, 2019; Sommerville, 2016).",
    )

    add_subheading(document, "4.1 Results for Objective One")
    for paragraph in [
        "Objective One sought to examine the current educational, rehabilitation, and operational challenges within the Ghana Prisons Service and derive the system requirements needed for a digital intervention. The main result of this objective was the identification of a requirement set built around role separation, centralized records, controlled access, lecturer visibility, management analytics, and resilient attendance workflows (Boateng, 2021; UNODC, 2018).",
        "The findings confirm that prison education in Ghana cannot be improved meaningfully through a course display interface alone. The environment requires a coordinated digital system in which educational access, inmate onboarding, verification, device allocation, and oversight reporting are linked. This requirement profile justified the multi-role structure of the developed portal (Ghana Prisons Service, 2015; World Prison Brief, 2025).",
    ]:
        add_body_paragraph(document, paragraph)

    add_table(
        document,
        "Table 4.1: Objective One results - identified problem areas and system requirements",
        ["Identified Finding", "Resulting Requirement in the Study"],
        [
            ["Manual and fragmented prison education records", "Centralized digital record and dashboard environment"],
            ["Weak visibility into inmate participation and attendance", "Attendance-aware portal and access event tracking"],
            ["Poor linkage between operational access and learning opportunity", "Clock-in and facility access logic connected to learning workflows"],
            ["Different users require different decisions and privileges", "Strict role-based access and interface separation"],
            ["Hardware uncertainty in constrained environments", "Fallback-aware camera and biometric workflow design"],
        ],
    )
    for paragraph in [
        "A further result of Objective One is conceptual clarity. The platform had to be positioned not as a generic school portal, but as a correctional education and access-management system. This result shaped the later design choices in Objective Two (Hevner et al., 2004; Boateng, 2021).",
        "Objective One therefore did not end with a narrative description only. It produced the implementation blueprint that guided module design, page hierarchy, and interface responsibilities across the application (Pressman and Maxim, 2019; Sommerville, 2016).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "4.2 Results for Objective Two")
    for paragraph in [
        "Objective Two focused on the design and implementation of the Ghana Prisons Learning Portal itself. The main result is a responsive prison learning and rehabilitation platform with role-specific user journeys. Public users begin from a landing page and a unified access interface. Administrators can register inmates and monitor institutional activity. Lecturers can review courses, attendance, and assignments. Management staff can access analytics summaries. Clocking officers can grant inmate learning access through controlled device allocation and biometric-oriented workflows. Inmates can enter their own portal to view learning content, assigned courses, progress, and certificates (Atzori et al., 2010; Pressman and Maxim, 2019).",
        "These results show that the identified prison requirements were translated into actual software modules rather than remaining at a conceptual level. The platform operationalizes IoT-oriented coordination through session-aware routing, connected access workflows, and dashboard visibility across roles (Shaikh and Zeadally, 2016; Hevner et al., 2004).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "4.2.1 Implemented Interface Screens", level=3)
    add_body_paragraph(
        document,
        "Figures 4.1 to 4.9 present key screens captured directly from the implemented application during live browser execution on 18 March 2026. Together they demonstrate that Objective Two resulted in a coherent software product from public entry to specialized prison workflows (Pressman and Maxim, 2019; Nielsen, 1994).",
        indent_first=False,
    )
    for spec in FIGURES:
        add_figure(document, spec)

    add_subheading(document, "4.2.2 Interpretation of the Implemented Modules", level=3)
    for paragraph in [
        "The public landing page and access portal establish a controlled entry pattern for the system. This is important because prison platforms cannot assume a one-size-fits-all interface. Role separation begins at the point of entry and continues through every downstream dashboard (Nielsen, 1994; UNODC, 2015).",
        "The administrative modules demonstrate how inmate onboarding can be digitized using generated identifiers, photo readiness, and role-aware workflows. This addresses one of the institutional weaknesses identified in the literature: fragmented and paper-heavy prison records (UNODC, 2018; Ghana Prisons Service, 2015).",
        "The lecturer dashboard shows the value of merging instructional data into one view. Course listings, pending assignments, submissions, and recent attendance records are displayed in a form that can support teaching and follow-up decisions more effectively than isolated records (Davis et al., 2013; Palanci et al., 2024).",
        "The clock-in page is especially important from an IoT perspective because it links verification state, room allocation, device type, and inmate access to learning spaces. This module exemplifies the central claim of the thesis: that rehabilitation systems become more useful when connected operational events are treated as part of one digital workflow (Atzori et al., 2010; Shaikh and Zeadally, 2016).",
        "The inmate dashboard and course pages translate institutional control into learner-facing value. Instead of exposing only administrative features, the system also presents a structured educational experience for the inmate, which is necessary if technology is to serve rehabilitation rather than administration alone (Davis et al., 2013; Nisser et al., 2024).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "4.3 Results for Objective Three")
    for paragraph in [
        "Objective Three focused on evaluating the developed platform. The results show that the system supported the major user scenarios defined for the study and that the application was ready for live final-defense demonstration (Pressman and Maxim, 2019; Sommerville, 2016).",
        f"The software was also prepared for defense demonstration beyond the local environment. At the time of this thesis assembly, the live production URL was {LIVE_URL}, with the current deployment accessible via {DEPLOYMENT_URL}. This strengthens the practical contribution of the study because the product exists as a runnable service rather than as a static design mockup only (Hevner et al., 2004; Nielsen, 1994).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "4.3.1 Scenario-Based Functional Validation", level=3)
    add_table(
        document,
        "Table 4.2: Scenario-based test results for Objective Three",
        ["Scenario", "Expected Outcome", "Observed Status"],
        [
            ["Open landing page", "Public portal loads with prison learning overview and course links", "Passed"],
            ["Open access page", "Role entry points for admin, management, lecturer, inmate, and clocking officer are visible", "Passed"],
            ["Admin session to dashboard", "Administrative dashboard loads after successful login", "Passed"],
            ["Admin registration page", "Inmate registration interface loads with camera and biometric readiness controls", "Passed"],
            ["Lecturer session", "Lecturer dashboard loads courses, assignments, submissions, and attendance summaries", "Passed"],
            ["Management session", "Management analytics dashboard loads without role collision", "Passed"],
            ["Clock-in session", "Clocking officer check-in page loads verification and device allocation workflow", "Passed"],
            ["Inmate session", "Inmate dashboard and course pages load assigned learning views", "Passed"],
            ["Live deployment", "Application accessible from production URL for defense use", "Passed"],
        ],
    )

    add_table(
        document,
        "Table 4.3: Deployment and quality verification summary",
        ["Verification Item", "Observation"],
        [
            ["Linting", "Project lint checks completed successfully during the final implementation cycle"],
            ["Production build", "Application build completed successfully during the final implementation cycle"],
            ["Role smoke tests", "Admin, lecturer, management, clocking, and inmate views were exercised during screenshot capture"],
            ["Deployment state", f"Production alias available at {LIVE_URL}"],
        ],
    )

    add_subheading(document, "4.3.2 Quality and Deployment Verification", level=3)
    add_body_paragraph(
        document,
        "The evaluation results indicate that the final platform was not merely visually complete but also technically coherent. Linting and production build verification supported the internal consistency of the codebase, while the live deployment link supported the external demonstrability of the product (Pressman and Maxim, 2019; Sommerville, 2016).",
    )

    add_subheading(document, "4.3.3 Selected Source Code Evidence", level=3)
    add_body_paragraph(
        document,
        "To complement the interface evidence, Appendix B includes selected code excerpts from the running codebase. These excerpts were chosen because they represent the backbone of the implemented solution: role-aware authentication, inmate session provisioning, clock-in verification logic, lecturer activity loading, and inmate registration readiness handling (Hevner et al., 2004; Pressman and Maxim, 2019).",
    )

    add_subheading(document, "4.4 Discussion of Findings")
    for paragraph in [
        "The implemented product demonstrates that a prison learning platform can unify educational and operational concerns without collapsing role boundaries. The system allows different actors to work from a shared platform while still protecting their responsibilities and screens through route restrictions and session-aware logic (UNODC, 2015; Pressman and Maxim, 2019).",
        "The product also shows the value of resilient design in constrained environments. For example, the clock-in and registration modules do not fail completely when camera or biometric support is limited. Instead, they preserve process continuity through structured fallbacks. This is especially important in correctional environments, where hardware readiness and connectivity may vary across facilities (Ghana Prisons Service, 2015; Sommerville, 2016).",
        "From a research perspective, the software confirms that IoT concepts can be interpreted pragmatically for prison education. The contribution is not in claiming a sensor-heavy prison deployment, but in showing how connected identity, access, event tracking, and role-aware dashboards can support rehabilitation operations within one coherent digital system (Atzori et al., 2010; Shaikh and Zeadally, 2016; Hevner et al., 2004).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "4.5 Chapter Summary")
    add_body_paragraph(
        document,
        "This chapter has presented the results of the study in the same order as the objectives. Section 4.1 presented the requirement findings, Section 4.2 presented the implemented platform, and Section 4.3 presented the evaluation results. The findings show that the study moved coherently from problem analysis to product design and then to product validation (Hevner et al., 2004; Pressman and Maxim, 2019).",
    )
    add_page_break(document)


def build_chapter_five(document: Document) -> None:
    add_heading_centered(document, "CHAPTER FIVE")
    add_heading_centered(document, "SUMMARY, CONCLUSION AND RECOMMENDATIONS", style_name="Heading 1")

    add_subheading(document, "5.0 Introduction")
    add_body_paragraph(
        document,
        "This chapter summarizes the study, presents the major conclusions, and outlines recommendations for deployment and future work (Boateng, 2021; Hevner et al., 2004).",
    )

    add_subheading(document, "5.1 Summary of the Study")
    for paragraph in [
        "The study set out to examine how IoT-oriented digital systems could improve education and rehabilitation delivery within the Ghana Prisons Service. Literature on correctional education, digital transformation, and secure institutional systems showed a clear need for integrated tools that move beyond manual records and fragmented access control (Boateng, 2021; UNODC, 2015).",
        "In response, the study designed and implemented the Ghana Prisons Learning Portal, a web-based platform that provides role-specific access for inmates, lecturers, administrators, management staff, and clocking officers. The system combines educational workflows with connected operational logic, including session control, access readiness, device allocation, and dashboard reporting (Atzori et al., 2010; Hevner et al., 2004).",
        "Validation through scenario-based testing, screenshot evidence, and deployment confirmation showed that the product supports the intended end-to-end flows required for final defense demonstration and future refinement (Pressman and Maxim, 2019; Sommerville, 2016).",
    ]:
        add_body_paragraph(document, paragraph)

    add_subheading(document, "5.2 Summary of Findings by Objective")
    add_numbered(
        document,
        [
            "The study established that IoT concepts are applicable to prison education when interpreted as coordinated digital identity, access, monitoring, and reporting rather than as unrestricted consumer connectivity.",
            "The study identified the major functional requirements of a prison learning platform and translated them into a working system architecture.",
            "The implemented product successfully delivered role-based modules for administrative control, lecturer monitoring, inmate learning, management oversight, and clocking workflows.",
            "The product was successfully deployed online and validated through live role-based scenario testing and interface capture.",
        ],
    )

    add_subheading(document, "5.3 Conclusion")
    add_body_paragraph(
        document,
        "The study concludes that integrating IoT-oriented digital workflows into correctional education can strengthen educational access, accountability, and operational coordination within the Ghana Prisons Service. The developed portal demonstrates that prison learning systems can be secure, multi-role, and practically deployable while still remaining aligned with rehabilitation objectives. Although the present implementation remains a prototype, it provides a defendable foundation for institutional modernization and future pilot deployment (Atzori et al., 2010; Davis et al., 2013; Rogers, 2003).",
    )

    add_subheading(document, "5.4 Recommendations")
    add_bullets(
        document,
        [
            "Pilot the platform within one controlled prison education environment before wider rollout.",
            "Integrate real institutional user management and controlled biometric hardware after policy approval.",
            "Provide structured staff training for administrators, lecturers, and clocking officers before deployment.",
            "Establish governance policies for digital identity, data protection, and audit review within prison education operations.",
            "Expand the inmate learning layer with additional course content, certification pathways, and vocational tracking features.",
        ],
    )

    add_subheading(document, "5.5 Suggestions for Future Work")
    add_bullets(
        document,
        [
            "Integration with real prison device inventories, biometric hardware, and facility network controls.",
            "Longitudinal evaluation of educational participation and reintegration outcomes after platform adoption.",
            "Offline-first synchronization across multiple prison facilities.",
            "Advanced analytics for educational risk detection, completion forecasting, and rehabilitation reporting.",
            "Vocational product tracking and e-commerce extensions for prison industries and reintegration support.",
        ],
    )

    add_subheading(document, "5.6 Final Remark")
    add_body_paragraph(
        document,
        "Digital transformation in correctional education must be both practical and humane. The Ghana Prisons Learning Portal contributes to that goal by demonstrating a secure and workable product that places education, monitoring, and rehabilitation support within one integrated platform (UNODC, 2015; Ghana Prisons Service, 2015).",
    )
    add_page_break(document)


def build_references(document: Document) -> None:
    add_heading_centered(document, "REFERENCES", style_name="Heading 1")
    for reference in REFERENCES:
        add_body_paragraph(document, reference, indent_first=False)
    add_page_break(document)


def build_appendices(document: Document) -> None:
    add_heading_centered(document, "APPENDICES", style_name="Heading 1")

    add_subheading(document, "APPENDIX A: SCREENSHOT EVIDENCE")
    add_body_paragraph(
        document,
        "This appendix contains selected interface captures from the live application used in the final defense document.",
        indent_first=False,
    )
    for spec in FIGURES:
        add_figure(document, spec)

    add_page_break(document)
    add_subheading(document, "APPENDIX B: KEY IMPLEMENTATION EXCERPTS")
    for listing in LISTINGS:
        path = ROOT / listing.path
        code = extract_code_lines(path, listing.start_line, listing.end_line)
        add_code_block(document, listing.title, f"{listing.description} Source file: {listing.path}", code)

    add_page_break(document)
    add_subheading(document, "APPENDIX C: STAKEHOLDER QUESTIONNAIRE")
    add_body_paragraph(
        document,
        "The following questionnaire can be used for stakeholder validation during pilot deployment or institutional review.",
        indent_first=False,
    )
    for heading, questions in QUESTIONNAIRE_SECTIONS.items():
        add_subheading(document, heading, level=3)
        add_numbered(document, questions)

    add_page_break(document)
    add_subheading(document, "APPENDIX D: DEPLOYMENT NOTES AND DEMONSTRATION GUIDE")
    add_body_paragraph(document, f"Production URL: {LIVE_URL}", indent_first=False)
    add_body_paragraph(document, f"Deployment URL: {DEPLOYMENT_URL}", indent_first=False)
    add_body_paragraph(
        document,
        "Recommended defense demonstration order: open the landing page, move to the access page, show admin dashboard and inmate registration, show lecturer dashboard, show management dashboard, show clock-in workflow, and conclude with inmate dashboard and course pages.",
        indent_first=False,
    )
    add_body_paragraph(
        document,
        "The screenshots embedded in this thesis were captured from the local deployment on 18 March 2026 using live browser automation and reflect the implemented state of the software at submission time.",
        indent_first=False,
    )

    add_page_break(document)
    add_subheading(document, "APPENDIX E: DETAILED TEST CASE MATRIX")
    add_table(
        document,
        "Table E.1: Extended functional test matrix",
        ["Test ID", "Action", "Expected Result", "Status"],
        [
            ["TC-01", "Open landing page", "Landing page renders with categories and course cards", "Passed"],
            ["TC-02", "Open access page", "Unified access options for all roles are visible", "Passed"],
            ["TC-03", "Admin login", "Admin session redirects to dashboard successfully", "Passed"],
            ["TC-04", "Admin open inmate registration", "Registration form loads with biometric capture controls", "Passed"],
            ["TC-05", "Admin sign-out", "User returns to access page", "Passed"],
            ["TC-06", "Lecturer login", "Lecturer dashboard loads correctly", "Passed"],
            ["TC-07", "Lecturer course actions", "Course and assignment action buttons are visible", "Passed"],
            ["TC-08", "Lecturer sign-out", "User returns to access page", "Passed"],
            ["TC-09", "Management login", "Management dashboard renders with analytics widgets", "Passed"],
            ["TC-10", "Management CSV export action", "Export logic is available for analytics data", "Passed"],
            ["TC-11", "Management sign-out", "User returns to access page", "Passed"],
            ["TC-12", "Clocking officer login", "Clock-in page loads with verification controls", "Passed"],
            ["TC-13", "Clocking officer face method", "Face-verification workflow and fallback state are available", "Passed"],
            ["TC-14", "Clocking officer sign-out", "User returns to access page", "Passed"],
            ["TC-15", "Inmate login", "Inmate dashboard loads successfully", "Passed"],
            ["TC-16", "Inmate open courses", "Course listing page renders successfully", "Passed"],
            ["TC-17", "Inmate course detail route", "Course access route is available from landing and portal entry points", "Passed"],
            ["TC-18", "Production alias", "Live URL opens the deployed application", "Passed"],
            ["TC-19", "Lint verification", "Codebase passes linting", "Passed"],
            ["TC-20", "Build verification", "Application completes production build", "Passed"],
        ],
    )

    add_page_break(document)
    add_subheading(document, "APPENDIX F: ROLE-BASED USER GUIDE")
    add_subheading(document, "F.1 Administrator Demonstration Flow", level=3)
    add_numbered(
        document,
        [
            "Open the access page and choose the administrator login route.",
            "Authenticate with the admin credentials and proceed to the dashboard.",
            "Review the dashboard metrics and operational panels.",
            "Open the inmate registration page and demonstrate the biometric/photo capture workflow.",
            "Return to the dashboard or related management pages and sign out to the access page.",
        ],
    )
    add_subheading(document, "F.2 Lecturer Demonstration Flow", level=3)
    add_numbered(
        document,
        [
            "Open the lecturer login route from the access page.",
            "Authenticate and load the lecturer dashboard.",
            "Show the course list, pending assignments, attendance view, and upload action buttons.",
            "Explain how the lecturer dashboard supports educational delivery and monitoring.",
            "Sign out to the access page.",
        ],
    )
    add_subheading(document, "F.3 Management Demonstration Flow", level=3)
    add_numbered(
        document,
        [
            "Open the management login route and authenticate.",
            "Load the analytics dashboard and review progress, course statistics, and attendance summaries.",
            "Show the filter and export-oriented management controls.",
            "Explain how the dashboard supports prison oversight and reporting.",
            "Sign out to the access page.",
        ],
    )
    add_subheading(document, "F.4 Clocking Officer Demonstration Flow", level=3)
    add_numbered(
        document,
        [
            "Open the lab access and clock-in route from the access page.",
            "Authenticate as the clocking officer and load the check-in interface.",
            "Demonstrate inmate selection, room allocation, device assignment, and verification method controls.",
            "Explain the fallback behavior when live camera or biometric hardware is unavailable.",
            "Sign out to the access page.",
        ],
    )
    add_subheading(document, "F.5 Inmate Demonstration Flow", level=3)
    add_numbered(
        document,
        [
            "Open the inmate login route or begin from the public landing page.",
            "Authenticate and load the inmate dashboard.",
            "Show available courses, learner progress, and certificate-oriented views.",
            "Open the course catalog and discuss how the inmate-facing flow supports rehabilitation through education.",
            "Sign out or return to the public experience.",
        ],
    )

    add_page_break(document)
    add_subheading(document, "APPENDIX G: API AND ROUTE SUMMARY")
    add_table(
        document,
        "Table G.1: Key application routes and API endpoints",
        ["Path", "Type", "Purpose"],
        [
            ["/landing", "Public route", "Public inmate-facing landing experience and course discovery"],
            ["/access", "Public route", "Unified role-based access page"],
            ["/auth/login", "Auth route", "Shared login page for admin, lecturer, and management users"],
            ["/auth/inmate-login", "Auth route", "Dedicated inmate login route"],
            ["/auth/clockin-login", "Auth route", "Clocking officer login route"],
            ["/admin/dashboard", "Protected route", "Administrative operations dashboard"],
            ["/admin/register-inmate", "Protected route", "Inmate onboarding and biometric capture workflow"],
            ["/lecturer/dashboard", "Protected route", "Lecturer monitoring and instructional dashboard"],
            ["/management/dashboard", "Protected route", "Management analytics and reporting dashboard"],
            ["/clockin/checkin", "Protected route", "Clock-in, access verification, and device allocation workflow"],
            ["/inmate/dashboard", "Protected route", "Inmate-facing learning dashboard"],
            ["/inmate/courses", "Protected route", "Inmate course listing and access page"],
            ["/api/v1/auth/login", "API endpoint", "Credential validation and session creation"],
            ["/api/v1/auth/logout", "API endpoint", "Session termination"],
            ["/api/v1/inmates", "API endpoint", "Inmate records retrieval"],
            ["/api/v1/access/sessions?status=active", "API endpoint", "Active access session retrieval for clock-in workflows"],
            ["/api/v1/courses", "API endpoint", "Course list retrieval"],
            ["/api/v1/assignments", "API endpoint", "Assignment list retrieval"],
            ["/api/v1/submissions", "API endpoint", "Submission retrieval for lecturer monitoring"],
            ["/api/v1/attendance", "API endpoint", "Attendance event retrieval"],
        ],
    )


def build_document() -> Document:
    document = Document()
    set_document_defaults(document)
    enable_update_fields_on_open(document)
    build_title_page(document)
    build_front_matter(document)
    build_chapter_one(document)
    build_chapter_two(document)
    build_chapter_three(document)
    build_chapter_four(document)
    build_chapter_five(document)
    build_references(document)
    build_appendices(document)
    return document


def main() -> None:
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    generate_diagram_assets()
    document = build_document()
    document.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
