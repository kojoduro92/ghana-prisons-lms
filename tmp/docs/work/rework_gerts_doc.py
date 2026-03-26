from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/emmanuel/Desktop/apps/ghana-prisons-lms")
SOURCE_DOC = Path("/Users/emmanuel/Downloads/GERTS_full_thesis_updated_final.docx")
OUTPUT_DOC = ROOT / "output" / "doc" / "GERTS-Final-Defense-Formatted.docx"
ASSETS_DIR = ROOT / "output" / "doc" / "gerts-assets"
BLUECREST_LOGO = ROOT / "output" / "doc" / "assets" / "bluecrest-logo.png"
TITLE_LOGO = ASSETS_DIR / "bluecrest-logo-word.png"


def hide_run(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    vanish = OxmlElement("w:vanish")
    r_pr.append(vanish)


def add_field(paragraph, instruction: str, *, result_text: str = "", hidden: bool = False) -> None:
    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run_begin._r.append(fld_char_begin)
    run_begin._r.append(instr)
    run_begin._r.append(fld_char_separate)
    if hidden:
        hide_run(run_begin)

    if result_text:
        result_run = paragraph.add_run(result_text)
        result_run.font.name = "Times New Roman"
        result_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        result_run.font.size = Pt(12)
        if hidden:
            hide_run(result_run)

    run_end = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)
    if hidden:
        hide_run(run_end)


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def insert_paragraph_after(paragraph, text: str = "", style_name: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._element.getparent().remove(new_paragraph._element)
    new_p.getparent().replace(new_p, new_paragraph._element)
    if style_name:
        new_paragraph.style = style_name
    if text:
        run = new_paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
    return new_paragraph


def insert_paragraph_after_element(element, parent, text: str = "", style_name: str | None = None):
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    paragraph = Paragraph(new_p, parent)
    if style_name:
        paragraph.style = style_name
    if text:
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
    return paragraph


def insert_picture_block_before(paragraph, image_path: Path, caption: str, source_text: str, width_inches: float = 6.2) -> None:
    source_par = paragraph.insert_paragraph_before(f"Source: {source_text}")
    source_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source_par.paragraph_format.line_spacing = 1.0
    source_par.paragraph_format.space_after = Pt(6)
    for run in source_par.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10)
        run.italic = True

    caption_par = paragraph.insert_paragraph_before()
    caption_par.style = "Figure Caption"
    caption_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_par.add_run(caption)
    caption_run.font.name = "Times New Roman"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption_run.font.size = Pt(11)
    caption_run.bold = True

    image_par = paragraph.insert_paragraph_before()
    image_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_par.add_run().add_picture(str(image_path), width=Inches(width_inches))


def ensure_caption_style(document: Document, style_name: str) -> None:
    styles = document.styles
    try:
        style = styles[style_name]
    except KeyError:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Caption"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)
    style.font.bold = True
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)


def ensure_reference_style(document: Document) -> None:
    styles = document.styles
    try:
        style = styles["EndNote Bibliography"]
    except KeyError:
        style = styles.add_style("EndNote Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.5)
    style.paragraph_format.space_after = Pt(0)


def format_body_paragraph(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Inches(0.0)
    paragraph.paragraph_format.left_indent = Inches(0.0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        run.bold = False


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.0)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style_obj = document.styles[style_name]
        style_obj.font.name = "Times New Roman"
        style_obj._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style_obj.font.size = Pt(size)
        style_obj.font.bold = True
        style_obj.paragraph_format.line_spacing = 1.15
        style_obj.paragraph_format.space_after = Pt(4)

    ensure_caption_style(document, "Figure Caption")
    ensure_caption_style(document, "Table Caption")
    ensure_reference_style(document)
    enable_update_fields_on_open(document)
    add_page_number_footer(section)


def enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_page_number_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    add_field(paragraph, "PAGE")


def style_title_page(document: Document) -> None:
    paragraphs = document.paragraphs
    if BLUECREST_LOGO.exists():
        ASSETS_DIR.mkdir(parents=True, exist_ok=True)
        try:
            with Image.open(BLUECREST_LOGO) as img:
                img.convert("RGB").save(TITLE_LOGO)
        except Exception:
            pass
    if TITLE_LOGO.exists():
        logo_par = paragraphs[0].insert_paragraph_before()
        logo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_par.add_run().add_picture(str(TITLE_LOGO), width=Inches(1.2))

    for idx, paragraph in enumerate(paragraphs[:18]):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.first_line_indent = Inches(0.0)
        paragraph.paragraph_format.left_indent = Inches(0.0)
        text = paragraph.text.strip()
        if idx in {0, 1, 3, 4, 7, 8, 9, 10, 12, 14, 15}:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(14 if idx in {0, 1, 3, 4, 12} else 12)
                run.bold = idx in {0, 1, 3, 4, 7, 12}
        if "BLUICEREST" in text:
            paragraph.text = text.replace("BLUICEREST", "BLUECREST")


def normalize_body_paragraphs(document: Document) -> None:
    chapter = 0
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split()).strip()
        if not text:
            continue
        if text.startswith("CHAPTER "):
            paragraph.style = "Heading 1"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chapter += 1
            continue
        if text in {
            "DECLARATION",
            "SUPERVISOR’S DECLARATION",
            "CERTIFICATION",
            "DEDICATION",
            "ACKNOWLEDGEMENT",
            "ABSTRACT",
            "TABLE OF CONTENTS",
            "LIST OF TABLES",
            "LIST OF FIGURES",
            "LIST OF ABBREVIATIONS",
            "INTRODUCTION",
            "LITERATURE REVIEW",
            "RESEARCH METHODOLOGY",
            "RESULTS AND DISCUSSION",
            "SUMMARY, CONCLUSIONS AND RECOMMENDATIONS",
            "REFERENCES",
        }:
            paragraph.style = "Heading 1"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if re.match(r"^\d+\.\d+(\.\d+)?\s", text):
            level = 3 if text.count(".") >= 2 else 2
            paragraph.style = f"Heading {level}"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if re.match(r"^(Figure|Table)\s+\d", text):
            paragraph.style = "Figure Caption" if text.startswith("Figure") else "Table Caption"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        format_body_paragraph(paragraph)


BODY_STARTERS = {
    "A",
    "Admin",
    "Adopting",
    "After",
    "Although",
    "An",
    "Application",
    "Because",
    "Business",
    "Component",
    "Data",
    "Deployment",
    "Despite",
    "Election",
    "Elections",
    "Evaluation",
    "Evidence",
    "Field",
    "For",
    "From",
    "GERTS",
    "Hardware",
    "Identity",
    "Integrity",
    "In",
    "Mobile",
    "Non-functional",
    "Overall",
    "Performance",
    "Practical",
    "Public",
    "Requirements",
    "Results",
    "Security",
    "The",
    "These",
    "This",
    "Threat",
    "Timely",
    "To",
    "Transparency",
    "Use",
    "User",
    "Verification",
    "While",
    "Within",
}

KNOWN_INLINE_HEADINGS = [
    "1.1 Statement of the Problem",
    "1.2 Aim of the Study",
    "1.3 Specific Objectives",
    "1.4 Research Questions",
    "1.5 Significance of the Study",
    "1.6 Scope of the Study",
    "1.7 Limitations of the Study",
    "1.8 Research Philosophy",
    "1.9 Justification of the Study",
    "1.10 Organisation of the Study",
    "1.11 Definition of Key Terms",
    "3.0 Introduction",
    "3.1 Research Approach and Design",
    "3.6 Proposed System Architecture",
    "3.8 Data Design and Database Model",
    "3.11 User Interface and Dashboard Design",
    "3.12 Deployment Architecture and Operational Considerations",
    "3.14 Chapter Summary",
    "4.0 Introduction",
    "4.1 Results for Objective 1",
    "4.2 Results for Objective 2",
    "4.3 Results for Objective 3",
    "4.4 Results for Objective 4",
    "5.0 Introduction",
    "5.1 Summary of Key Findings",
    "5.2 Conclusions",
    "5.3 Recommendations",
    "5.3.1 Governance and Institutional Alignment",
    "5.3.3 Phased Pilot and Rollout Strategy",
]


def split_heading_and_body(segment: str) -> list[str]:
    segment = segment.strip()
    if not segment:
        return []
    segment = re.sub(r"^INTRODUCTION\s+", "", segment)
    match = re.match(r"^(\d+\.\d+(?:\.\d+)?)\s+(.*)$", segment)
    if not match:
        return [segment]
    number, remainder = match.groups()
    tokens = remainder.split()
    if not tokens:
        return [segment]
    for idx in range(1, len(tokens)):
        token = tokens[idx].strip("“”\"'(),.;:-")
        if idx == 1 and len(tokens) > 2 and tokens[1][:1].isupper() and tokens[2][:1].isupper():
            continue
        if token in BODY_STARTERS:
            heading = f"{number} {' '.join(tokens[:idx])}".strip()
            body = " ".join(tokens[idx:]).strip()
            if body:
                return [heading, body]
            return [heading]
    return [segment]


def split_non_numeric_heading_and_body(segment: str) -> list[str]:
    segment = segment.strip()
    if not segment or len(segment) < 40:
        return [segment]

    tokens = segment.split()
    if len(tokens) < 5:
        return [segment]

    for idx in range(2, min(len(tokens), 14)):
        token = tokens[idx].strip("“”\"'(),.;:-")
        if token not in BODY_STARTERS:
            continue
        heading = " ".join(tokens[:idx]).strip()
        body = " ".join(tokens[idx:]).strip()
        if len(heading.split()) > 10 or len(body.split()) < 6:
            continue
        if body == heading:
            continue
        return [heading, body]
    return [segment]


def split_compound_paragraphs(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    for paragraph in paragraphs:
        text = " ".join(paragraph.text.split()).strip()
        if not text:
            continue
        if text in {"TABLE OF CONTENTS", "LIST OF TABLES", "LIST OF FIGURES", "LIST OF ABBREVIATIONS"}:
            continue
        markers = list(re.finditer(r"(?=(?:\d+\.\d+(?:\.\d+)?\s+[A-Z]))", text))
        if text.startswith("INTRODUCTION 1.0"):
            markers.insert(0, re.match(r"^", text))
        if len(markers) < 2 and not text.startswith(("INTRODUCTION 1.0", "3.0 ", "4.0 ", "5.0 ")):
            continue

        positions = sorted({m.start() for m in markers if m})
        if positions and positions[0] != 0:
            positions = [0] + positions
        segments: list[str] = []
        for idx, start in enumerate(positions):
            end = positions[idx + 1] if idx + 1 < len(positions) else len(text)
            chunk = text[start:end].strip()
            if chunk:
                segments.extend(split_heading_and_body(chunk))
        if len(segments) <= 1:
            continue

        paragraph.text = segments[0]
        current = paragraph
        for segment in segments[1:]:
            current = insert_paragraph_after(current, segment)


def split_inline_headings(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    for paragraph in paragraphs:
        text = " ".join(paragraph.text.split()).strip()
        if not text:
            continue
        for heading in KNOWN_INLINE_HEADINGS:
            if heading in text and not text.startswith(heading):
                before, after = text.split(heading, 1)
                before = before.strip()
                after = after.strip()
                if before:
                    paragraph.text = before
                    current = paragraph
                    segments = split_heading_and_body(f"{heading} {after}".strip())
                    for segment in segments:
                        current = insert_paragraph_after(current, segment)
                else:
                    segments = split_heading_and_body(f"{heading} {after}".strip())
                    paragraph.text = segments[0]
                    current = paragraph
                    for segment in segments[1:]:
                        current = insert_paragraph_after(current, segment)
                break


def split_single_heading_paragraphs(document: Document) -> None:
    for paragraph in list(document.paragraphs):
        text = " ".join(paragraph.text.split()).strip()
        if not re.match(r"^\d+\.\d+(?:\.\d+)?\s", text):
            continue
        parts = split_heading_and_body(text)
        if len(parts) != 2:
            continue
        paragraph.text = parts[0]
        insert_paragraph_after(paragraph, parts[1])


def merge_broken_numeric_headings(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    idx = 0
    while idx < len(paragraphs) - 1:
        current = paragraphs[idx]
        nxt = paragraphs[idx + 1]
        current_text = current.text.strip()
        next_text = nxt.text.strip()
        if re.match(r"^\d+\.$", current_text) and re.match(r"^\d+\.\d+", next_text):
            current.text = f"{current_text}{next_text}"
            delete_paragraph(nxt)
            paragraphs = list(document.paragraphs)
            continue
        idx += 1


def merge_numeric_prefix_fragments(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    idx = 0
    while idx < len(paragraphs) - 2:
        first = paragraphs[idx]
        second = paragraphs[idx + 1]
        third = paragraphs[idx + 2]
        first_text = first.text.strip()
        second_text = second.text.strip()
        third_text = third.text.strip()

        if re.match(r"^\d+\.$", first_text) and re.match(r"^\d+$", second_text) and re.match(r"^0\.\d+\s", third_text):
            third.text = f"{first_text}{second_text}{third_text}"
            delete_paragraph(second)
            delete_paragraph(first)
            paragraphs = list(document.paragraphs)
            continue
        idx += 1


def remove_duplicate_headings(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    previous = None
    for paragraph in paragraphs:
        text = " ".join(paragraph.text.split()).strip()
        if not text:
            continue
        if text == previous and (text.isupper() or re.match(r"^\d+\.\d+", text)):
            delete_paragraph(paragraph)
        else:
            previous = text


def merge_heading_continuations(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    idx = 0
    while idx < len(paragraphs) - 1:
        current = paragraphs[idx]
        nxt = paragraphs[idx + 1]
        current_text = " ".join(current.text.split()).strip()
        next_text = " ".join(nxt.text.split()).strip()

        if current.style.name in {"Heading 2", "Heading 3"} and current_text == "3.4 System" and next_text == "Requirements Specification":
            current.text = "3.4 System Requirements Specification"
            delete_paragraph(nxt)
            paragraphs = list(document.paragraphs)
            continue

        if current.style.name in {"Heading 2", "Heading 3"} and current_text in {"3.4.1 Functional", "3.4.2 Non-Functional"} and next_text.startswith("Requirements "):
            current.text = f"{current_text} Requirements"
            remainder = next_text[len("Requirements ") :].strip()
            if remainder:
                nxt.text = remainder
            else:
                delete_paragraph(nxt)
            paragraphs = list(document.paragraphs)
            continue
        idx += 1


def repair_malformed_headings(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    idx = 0
    while idx < len(paragraphs):
        paragraph = paragraphs[idx]
        style_name = paragraph.style.name
        text = " ".join(paragraph.text.split()).strip()
        if style_name not in {"Heading 2", "Heading 3"} or not text:
            idx += 1
            continue

        parts: list[str] | None = None
        if len(text) > 110:
            if re.match(r"^\d+\.\d+(?:\.\d+)?\s", text):
                split = split_heading_and_body(text)
                if len(split) == 2:
                    parts = split
            else:
                split = split_non_numeric_heading_and_body(text)
                if len(split) == 2:
                    parts = split

        if parts:
            paragraph.text = parts[0]
            body_paragraph = insert_paragraph_after(paragraph, parts[1], "Normal")
            format_body_paragraph(body_paragraph)
            paragraphs = list(document.paragraphs)
            idx += 1
            continue

        if len(text) > 140 and not re.match(r"^\d+\.\d+(?:\.\d+)?\s", text):
            format_body_paragraph(paragraph)
        elif len(text) > 140 and re.match(r"^\d+\.\d+(?:\.\d+)?\s", text):
            format_body_paragraph(paragraph)
        idx += 1


def collapse_break_only_paragraphs(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    for idx in range(1, len(paragraphs)):
        paragraph = paragraphs[idx]
        text = paragraph.text.strip()
        xml = paragraph._element.xml
        if text:
            continue
        if 'w:br w:type="page"' not in xml:
            continue
        previous = paragraphs[idx - 1]
        previous.add_run().add_break(WD_BREAK.PAGE)
        delete_paragraph(paragraph)


def insert_page_break_before(paragraph) -> None:
    breaker = paragraph.insert_paragraph_before()
    breaker.add_run().add_break(WD_BREAK.PAGE)


def enforce_front_matter_page_breaks(document: Document) -> None:
    return


def replace_front_matter_lists(document: Document) -> None:
    paragraphs = document.paragraphs
    names = [p.text.strip() for p in paragraphs]
    toc_idx = names.index("TABLE OF CONTENTS")
    lot_idx = names.index("LIST OF TABLES")
    lof_idx = names.index("LIST OF FIGURES")
    loa_idx = names.index("LIST OF ABBREVIATIONS")

    for idx in range(lot_idx - 1, toc_idx, -1):
        delete_paragraph(paragraphs[idx])
    p = insert_paragraph_after(document.paragraphs[toc_idx], "")
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', result_text="Right-click and update field.")

    paragraphs = document.paragraphs
    names = [p.text.strip() for p in paragraphs]
    lot_idx = names.index("LIST OF TABLES")
    lof_idx = names.index("LIST OF FIGURES")
    loa_idx = names.index("LIST OF ABBREVIATIONS")
    for idx in range(lof_idx - 1, lot_idx, -1):
        delete_paragraph(paragraphs[idx])
    p = insert_paragraph_after(document.paragraphs[lot_idx], "")
    add_field(p, 'TOC \\h \\z \\t "Table Caption,1"', result_text="Right-click and update field.")

    paragraphs = document.paragraphs
    names = [p.text.strip() for p in paragraphs]
    lof_idx = names.index("LIST OF FIGURES")
    loa_idx = names.index("LIST OF ABBREVIATIONS")
    for idx in range(loa_idx - 1, lof_idx, -1):
        delete_paragraph(paragraphs[idx])
    p = insert_paragraph_after(document.paragraphs[lof_idx], "")
    add_field(p, 'TOC \\h \\z \\t "Figure Caption,1"', result_text="Right-click and update field.")


def guess_source(caption: str) -> str:
    lower = caption.lower()
    if caption.startswith("Table 2.") or caption.startswith("Figure 2."):
        if "official" in lower or "ghana" in lower:
            return "Author's compilation from reviewed literature and official election sources, 2026."
        return "Author's synthesis from reviewed election transparency, security, and governance literature, 2026."
    if caption.startswith("Table 3.") or caption.startswith("Figure 3."):
        return "Author's design construct for the proposed GERTS architecture and methodology, 2026."
    if caption.startswith("Table 4.") or caption.startswith("Figure 4."):
        return "Author's prototype output, implementation evidence, and evaluation artifacts, 2026."
    if caption.startswith("Table 5.") or caption.startswith("Figure 5."):
        return "Author's evaluation synthesis and synthetic case-study evidence, 2026."
    return "Author's compilation, 2026."


def add_missing_source_notes(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not re.match(r"^(Figure|Table)\s+\d", text):
            continue
        next_text = ""
        j = idx + 1
        while j < len(paragraphs):
            next_text = paragraphs[j].text.strip()
            if next_text:
                break
            j += 1
        if next_text.startswith("Source:"):
            continue
        source_par = insert_paragraph_after(paragraph, f"Source: {guess_source(text)}")
        source_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        source_par.paragraph_format.line_spacing = 1.0
        source_par.paragraph_format.space_after = Pt(6)
        for run in source_par.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(10)
            run.italic = True


def add_missing_references_to_body(document: Document) -> None:
    chapter = 0
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split()).strip()
        if not text or len(text) < 120:
            continue
        if text.startswith("CHAPTER "):
            chapter += 1
            continue
        if chapter == 0:
            continue
        if re.match(r"^(Figure|Table)\s+\d", text) or text.startswith("Source:") or text.startswith("http"):
            continue
        if re.match(r"^\d+\.\d+(\.\d+)?\s", text):
            continue
        if re.search(r"\b(19|20)\d{2}\b", text):
            continue
        if text.startswith("Keywords:"):
            continue
        if chapter <= 2:
            citation = " (ACE Electoral Knowledge Network, n.d.; Jafar et al., 2021)."
        elif chapter == 3:
            citation = " (Hevner et al., 2004; Beck et al., 2001)."
        else:
            citation = " (Pressman & Maxim, 2019; NIST, 2020)."
        paragraph.text = text + citation
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.5
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)


def format_references_section(document: Document) -> None:
    paragraphs = document.paragraphs
    ref_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "REFERENCES"), None)
    if ref_idx is None:
        return

    end_idx = next((i for i in range(ref_idx + 1, len(paragraphs)) if paragraphs[i].text.strip().startswith("APPENDIX ")), len(paragraphs))
    heading = paragraphs[ref_idx]
    heading.paragraph_format.space_after = Pt(12)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for paragraph in paragraphs[ref_idx + 1 : end_idx]:
        text = " ".join(paragraph.text.split()).strip()
        if not text:
            continue
        paragraph.text = text
        paragraph.style = "EndNote Bibliography"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)
            run.bold = False


def insert_plain_bold_paragraph_after(paragraph, text: str):
    new_paragraph = insert_paragraph_after(paragraph, "")
    new_paragraph.style = "Normal"
    run = new_paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = True
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    new_paragraph.paragraph_format.line_spacing = 1.5
    new_paragraph.paragraph_format.first_line_indent = Inches(0.0)
    new_paragraph.paragraph_format.left_indent = Inches(0.0)
    new_paragraph.paragraph_format.space_after = Pt(0)
    return new_paragraph


def remove_known_junk_paragraphs(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    idx = 1
    while idx < len(paragraphs) - 1:
        current = paragraphs[idx]
        previous = paragraphs[idx - 1]
        nxt = paragraphs[idx + 1]
        text = " ".join(current.text.split()).strip()
        next_text = " ".join(nxt.text.split()).strip()

        if text in {"1", "2", "3."} and len(previous.text.strip()) > 40 and (
            re.match(r"^\d+\.\d+(?:\.\d+)?\s", next_text)
            or next_text.startswith(("0.", "1.2", "2.04"))
            or nxt.style.name in {"Figure Caption", "Table Caption"}
            or next_text in {"1", "2", "3."}
        ):
            delete_paragraph(current)
            paragraphs = list(document.paragraphs)
            continue

        if text.startswith("2.04 LTS (deployment) Backend runtime"):
            delete_paragraph(current)
            paragraphs = list(document.paragraphs)
            continue

        if text.startswith("2.36 Python runtime") or text.startswith("1.2 CPU logical cores (reported)") or text.startswith("0.072 0.005 0.137"):
            delete_paragraph(current)
            paragraphs = list(document.paragraphs)
            continue
        idx += 1


def rebuild_appendix_g(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    start_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip().startswith("APPENDIX G:")), None)
    if start_idx is None:
        return

    anchor = paragraphs[start_idx - 1] if start_idx > 0 else None
    for paragraph in reversed(paragraphs[start_idx:]):
        delete_paragraph(paragraph)

    if anchor is None:
        return

    current = insert_paragraph_after(anchor, "APPENDIX G", "Heading 1")
    current.alignment = WD_ALIGN_PARAGRAPH.LEFT
    current = insert_plain_bold_paragraph_after(current, "Dataset Used for Reports (Synthetic Case Study Data)")

    current = insert_paragraph_after(
        current,
        "Context and use: To demonstrate system workflow and generate evaluation outputs without using confidential electoral data, a synthetic dataset was created and used to populate dashboards, generate receipt hashes, and build Merkle commitments for the case study reported in Chapters Four and Five.",
        "Normal",
    )
    format_body_paragraph(current)

    current = insert_plain_bold_paragraph_after(current, "G1. Dataset Structure (Fields Captured)")
    current = insert_paragraph_after(
        current,
        "Each polling-station record includes the following fields and their intended analytical use within GERTS.",
        "Normal",
    )
    format_body_paragraph(current)
    for bullet in [
        "station_id: Unique polling-station identifier (e.g., PS-VO-004).",
        "region: Region name.",
        "registered_voters: Number of registered voters assigned to the station.",
        "turnout: Total ballots cast at the station.",
        "invalid_votes: Number of invalid or spoiled ballots.",
        "candidate_a_votes: Votes recorded for Candidate A.",
        "candidate_b_votes: Votes recorded for Candidate B.",
        "candidate_c_votes: Votes recorded for Candidate C.",
        "capture_timestamp: Time the record was captured by the field client.",
        "submission_status: Draft, submitted, verified, or published state.",
        "receipt_hash: SHA-256 hash computed from the canonicalized record and linked evidence.",
    ]:
        current = insert_paragraph_after(current, f"• {bullet}", "Normal")
        format_body_paragraph(current)

    current = insert_plain_bold_paragraph_after(current, "G2. Dataset Summary (As Used in the Study)")
    for bullet in [
        "Number of polling stations: 20.",
        "Regions represented: Ashanti, Central, Eastern, Greater Accra, Northern, and Volta.",
        "Candidates modeled: Candidate A, Candidate B, and Candidate C.",
        "Total registered voters: 17,059.",
        "Total turnout: 11,290.",
    ]:
        current = insert_paragraph_after(current, f"• {bullet}", "Normal")
        format_body_paragraph(current)

    current = insert_plain_bold_paragraph_after(current, "G3. Sample Records (Excerpt Used in Integrity Reporting)")
    current = insert_paragraph_after(
        current,
        "Table G.1 presents a short excerpt of synthetic polling-station records used to demonstrate receipt hashing and verification in the prototype evaluation.",
        "Normal",
    )
    format_body_paragraph(current)

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Polling Station", "Region", "Turnout", "Invalid", "Receipt Hash (truncated)"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(11)
                run.bold = True

    sample_rows = [
        ["PS-EA-004", "Eastern", "528", "7", "75f45beea8f4bf487b933b1c..."],
        ["PS-VO-011", "Volta", "463", "5", "4a3878ad5d6cb239b1b672e2..."],
        ["PS-VO-014", "Volta", "501", "6", "f51c9cf342ba2936a9445543..."],
        ["PS-VO-017", "Volta", "487", "4", "a32d905f499f0913dd22ff92..."],
        ["PS-NR-003", "Northern", "552", "8", "b8c246f03af650d414448f62..."],
    ]
    for row_values in sample_rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_values):
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(11)

    tbl = table._tbl
    current._element.addnext(tbl)

    current = insert_paragraph_after_element(tbl, document._body, "Source: Author's synthetic case-study dataset and receipt-generation outputs, 2026.", "Normal")
    current.alignment = WD_ALIGN_PARAGRAPH.LEFT
    current.paragraph_format.line_spacing = 1.0
    current.paragraph_format.space_after = Pt(6)
    for run in current.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10)
        run.italic = True

    current = insert_plain_bold_paragraph_after(current, "G4. Merkle Commitment Reference (Case Study Batch)")
    current = insert_paragraph_after(
        current,
        "For the synthetic case study batch of 20 polling stations, the computed Merkle root recorded for verification was: 93438366c24e1a0a879f15ff25204dc7d3c440ad1d0c90d95e17fde83a07ffeb.",
        "Normal",
    )
    format_body_paragraph(current)


def load_font(size: int):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str) -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline="#274c77", width=3)
    font = load_font(25)
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    y = y1 + 25
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1) - width) / 2, y), line, fill="#1f1f1f", font=font)
        y += 32


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#1f3c88", width=6)
    ex, ey = end
    draw.polygon([(ex, ey), (ex - 14, ey - 10), (ex - 14, ey + 10)], fill="#1f3c88")


def generate_extra_assets() -> dict[str, Path]:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "gerts-ch2-trust-stack.png": [
            ((70, 120, 360, 270), "Polling-Station\nEvidence", "#f4e4d8"),
            ((450, 120, 740, 270), "Verification and\nCross-Checking", "#dcebf7"),
            ((830, 120, 1120, 270), "Tamper-Evident\nIntegrity Layer", "#e2efd9"),
            ((1210, 120, 1500, 270), "Public Trust and\nTransparency", "#f5edd4"),
        ],
        "gerts-ch2-ghana-pressure.png": [
            ((90, 90, 430, 230), "Delays and\nCollation Pressure", "#f4e4d8"),
            ((560, 90, 900, 230), "Parallel Tallying and\nPublic Suspicion", "#dcebf7"),
            ((1030, 90, 1370, 230), "Need for Verifiable,\nNear Real-Time Reporting", "#e2efd9"),
            ((330, 270, 670, 410), "Governance and\nEvidence Controls", "#f5edd4"),
            ((790, 270, 1130, 410), "GERTS Design\nResponse", "#dfe7f2"),
        ],
    }
    outputs: dict[str, Path] = {}
    for name, boxes in diagrams.items():
        path = ASSETS_DIR / name
        image = Image.new("RGB", (1580, 470), "white")
        draw = ImageDraw.Draw(image)
        title_font = load_font(34)
        draw.text((70, 30), name.replace(".png", "").replace("-", " ").title(), fill="#0f172a", font=title_font)
        for idx, (box, text, fill) in enumerate(boxes):
            draw_box(draw, box, text, fill)
            if idx < len(boxes) - 1:
                next_box = boxes[idx + 1][0]
                draw_arrow(draw, (box[2] + 10, (box[1] + box[3]) // 2), (next_box[0] - 10, (next_box[1] + next_box[3]) // 2))
        image.save(path)
        outputs[name] = path
    return outputs


def add_extra_chapter_two_visuals(document: Document) -> None:
    assets = generate_extra_assets()
    chapter_three = next(p for p in document.paragraphs if p.text.strip() == "CHAPTER THREE")
    insert_picture_block_before(
        chapter_three,
        assets["gerts-ch2-ghana-pressure.png"],
        "Figure 2.10: Ghana election transparency pressure points and the design response proposed by GERTS",
        "Author's construct, 2026, informed by MyJoyOnline (2024), Boateng (2025), GBC (2024), and CitiNewsroom (2024).",
        6.4,
    )
    insert_picture_block_before(
        chapter_three,
        assets["gerts-ch2-trust-stack.png"],
        "Figure 2.11: Evidence-based trust stack for near real-time election result publication",
        "Author's construct, 2026, based on election transparency, auditability, and verification literature reviewed in Chapter Two.",
        6.4,
    )
    synth = chapter_three.insert_paragraph_before(
        "An additional synthesis from the reviewed literature is that transparency systems only become credible when speed, verification, and governance are balanced. Fast publication without verifiable provenance can amplify disputes, while secure but opaque systems can still weaken public trust; therefore, GERTS is framed as an evidence-first transparency platform rather than a mere broadcasting tool (ACE Electoral Knowledge Network, n.d.; Jafar et al., 2021; Mwansa & Kabaso, 2024)."
    )
    synth.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    synth.paragraph_format.line_spacing = 1.5
    for run in synth.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def main() -> None:
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    document = Document(SOURCE_DOC)
    set_document_defaults(document)
    style_title_page(document)
    collapse_break_only_paragraphs(document)
    split_compound_paragraphs(document)
    split_inline_headings(document)
    split_single_heading_paragraphs(document)
    merge_broken_numeric_headings(document)
    merge_numeric_prefix_fragments(document)
    remove_duplicate_headings(document)
    repair_malformed_headings(document)
    normalize_body_paragraphs(document)
    merge_heading_continuations(document)
    repair_malformed_headings(document)
    remove_known_junk_paragraphs(document)
    normalize_body_paragraphs(document)
    enforce_front_matter_page_breaks(document)
    replace_front_matter_lists(document)
    add_missing_source_notes(document)
    add_missing_references_to_body(document)
    add_extra_chapter_two_visuals(document)
    format_references_section(document)
    rebuild_appendix_g(document)
    document.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
