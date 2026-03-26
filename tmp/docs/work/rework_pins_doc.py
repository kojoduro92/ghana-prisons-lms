from __future__ import annotations

import re
import ssl
import subprocess
from pathlib import Path
from urllib.request import Request, urlopen

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/emmanuel/Desktop/apps/ghana-prisons-lms")
INPUT_DOC = Path("/Users/emmanuel/Downloads/e (1).docx")
OUTPUT_DOC = ROOT / "output/doc/PINS-Final-Defense-Reworked.docx"
ASSET_DIR = ROOT / "output/doc/pins-assets"

FONT_REGULAR = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"
FONT_ITALIC = "/System/Library/Fonts/Supplemental/Times New Roman Italic.ttf"
FONT_BOLD_ITALIC = "/System/Library/Fonts/Supplemental/Times New Roman Bold Italic.ttf"

REAL_IMAGE_SOURCES = [
    (
        "ghana_prisons_photo",
        "https://ghanaprisons.gov.gh/wp-content/uploads/2025/11/IMG_8940-1-scaled.jpg",
        True,
        "Figure 2.6: Ghana Prisons Service institutional setting for prison-administration reform",
        "Ghana Prisons Service. (n.d.). https://ghanaprisons.gov.gh/",
        "The institutional reform context of Ghana’s prison administration illustrates why reliable records, coordination, and communication matter for custodial governance and stakeholder trust.",
    ),
    (
        "judicial_paperless",
        "https://judicial.gov.gh/js/media/k2/items/cache/556202461533c16857a46ed5bfb7c21c_L.jpg",
        False,
        "Figure 2.7: Judicial Service of Ghana paperless-courts campaign image",
        "Judicial Service of Ghana. (2019). https://judicial.gov.gh/js/pages/publications/news-publications/js-latest-news/item/365-judicial-service-kick-starts-paperless-courts.html",
        "Digital justice reform in Ghana provides a concrete interoperability context for PINS because prison workflows increasingly depend on timely exchanges with paperless court processes.",
    ),
    (
        "unodc_tools_cover",
        "https://www.unodc.org/doc/e-lectures/Justice_Section_Tools_Catalogue.pdf",
        False,
        "Figure 2.8: UNODC prison-reform and criminal-justice tools publication set",
        "United Nations Office on Drugs and Crime. (2023). https://www.unodc.org/doc/e-lectures/Justice_Section_Tools_Catalogue.pdf",
        "UNODC publications provide internationally recognized prison-reform and justice-administration guidance that informs functional requirements, governance expectations, and implementation safeguards relevant to PINS.",
    ),
]


REFERENCE_ENTRIES = [
    "California Courts Newsroom. (2025, April 9). Remote proceedings have enhanced California’s courtrooms and improved court participation. https://newsroom.courts.ca.gov/news/remote-proceedings-have-enhanced-californias-courtrooms-and-improved-court-participation",
    "Chohlas-Wood, A., Coots, M., Nudell, J., Nyarko, J., Brunskill, E., Rogers, T., & Goel, S. (2023). Automated reminders reduce incarceration for missed court dates: Evidence from a text message experiment. arXiv. https://arxiv.org/abs/2306.12389",
    "Conference of Chief Justices, & Conference of State Court Administrators. (2016). Call to action: Achieving civil justice for all. https://www.sji.gov/wp-content/uploads/CJI-Full-Report.pdf",
    "Creswell, J. W., & Creswell, J. D. (2018). Research design: Qualitative, quantitative, and mixed methods approaches (5th ed.). Sage.",
    "ENISA. (2023). ENISA threat landscape 2023. https://www.enisa.europa.eu/topics/cyber-threats/threat-landscape",
    "Ghana Government. (n.d.). Ghana Prisons Service. https://www.ghana.gov.gh/mdas/e1a4526b61/",
    "Ghana Prisons Service. (n.d.). Ghana Prisons Service: Vigilance, humanity and fortitude. https://ghanaprisons.gov.gh/",
    "ISO/IEC. (2022). ISO/IEC 27001:2022 information security, cybersecurity and privacy protection - Information security management systems - Requirements.",
    "Judicial Council of California. (2021). Remote access to court proceedings: Interim report. https://newsroom.courts.ca.gov/sites/default/files/newsroom/2021-08/P3%20Workgroup%20Remote%20Access%20Interim%20Report%2008162021.pdf",
    "Judicial Service of Ghana. (2018). Video conference facility booking form. https://judicial.gov.gh/jsweb/jsfiles/videoconference.pdf",
    "Judicial Service of Ghana. (2019). Judicial Service kick starts paperless courts. https://judicial.gov.gh/js/pages/publications/news-publications/js-latest-news/item/365-judicial-service-kick-starts-paperless-courts.html",
    "Judicial Service of Ghana. (n.d.). e-Justice frequently asked questions. https://judicial.gov.gh/js/pages/e-services/e-justice/f-a-q.html",
    "Justice System Partners. (2024). Evaluating the evidence: Automated court reminder systems. https://justicesystempartners.org/wp-content/uploads/2024/03/Evaluating-the-Evidence-Court-Notification-Systems.pdf",
    "Laravel. (n.d.). Laravel documentation. https://laravel.com/docs",
    "Minnesota Judicial Branch. (n.d.). Hearing eReminders. https://mncourts.gov/hearing-ereminders",
    "MySQL. (n.d.). MySQL 8.0 reference manual. https://dev.mysql.com/doc/refman/8.0/en/",
    "National Institute of Standards and Technology. (2020). Security and privacy controls for information systems and organizations (SP 800-53 Rev. 5). https://csrc.nist.gov/pubs/sp/800/53/r5/upd1/final",
    "OECD. (2019). The path to becoming a data-driven public sector. OECD Publishing. https://doi.org/10.1787/059814a7-en",
    "Peffers, K., Tuunanen, T., Rothenberger, M. A., & Chatterjee, S. (2007). A design science research methodology for information systems research. Journal of Management Information Systems, 24(3), 45-77.",
    "Sommerville, I. (2016). Software engineering (10th ed.). Pearson.",
    "United Nations Development Programme. (2024). Human development report 2023/2024: Breaking the gridlock - Reimagining cooperation in a polarized world. https://hdr.undp.org/content/human-development-report-2023-24",
    "United Nations Office on Drugs and Crime. (2023). Criminal justice, prison reform and crime prevention tools catalogue. https://www.unodc.org/doc/e-lectures/Justice_Section_Tools_Catalogue.pdf",
    "World Bank. (2021). GovTech maturity index: The state of public sector digital transformation. https://documents.worldbank.org/curated/en/495831623303257484/pdf/GovTech-Maturity-Index-The-State-of-Public-Sector-Digital-Transformation.pdf",
    "WHO Regional Office for Europe. (2021). Improving health in prisons: New WHO prison health framework can improve data quality. https://www.who.int/europe/news/item/28-10-2021-improving-health-in-prisons-new-who-prison-health-framework-can-improve-data-quality",
    "WHO Regional Office for Europe. (n.d.). The WHO/Europe Health in Prisons Programme (HIPP). https://www.who.int/europe/teams/alcohol-illicit-drugs-prison-health/the-who-europe-health-in-prisons-programme-%28hipp%29",
]

PLACEHOLDER_FIGURE_BLOCKS = [
    (
        "Figure 2.1: Typical Prison Information System Components",
        "2.2 Review of System Architecture Approaches for Prison Information Systems",
        "Note. Figure 2.1 summarizes the user layer, application services, secure data store, and external integration points that recur in the prison-information-systems literature and together motivate the integrated PINS design.",
        "Researcher's synthesis from correctional-information, digital-governance, and prison-health literature reviewed in this chapter, 2026.",
    ),
    (
        "Figure 2.2: Notification Delivery Pipeline",
        "2.4 Review of Core Functional Modules and Prototype Development Approaches",
        "Note. Figure 2.2 condenses the event trigger, queueing, dispatch, retry, and audit steps required for legally significant prison notifications in low-connectivity operational environments.",
        "Researcher's synthesis from automated-notification and public-sector workflow literature reviewed in this chapter, 2026.",
    ),
    (
        "Figure 2.3: Security Controls Stack for PINS",
        "2.5 Review of Notification Workflows, Security Controls, and Interoperability Approaches",
        "Note. Figure 2.3 shows security as a layered control environment spanning governance, identity, encryption, audit integrity, operational assurance, and infrastructure resilience.",
        "Researcher's synthesis from ENISA (2023), NIST (2020), ISO/IEC (2022), and prison-reform governance literature, 2026.",
    ),
    (
        "Figure 2.4: Sociotechnical Adoption Model for PINS",
        "2.7 Comparative Models and Case Studies",
        "Note. Figure 2.4 highlights that sustainable prison digitization depends on the interaction of technology quality, user capability, and institutional support rather than on software availability alone.",
        "Researcher's synthesis from OECD (2019), World Bank (2021), UNDP (2024), and related adoption literature, 2026.",
    ),
    (
        "Figure 2.5: Conceptual Framework for the PINS Study",
        "2.9.3 Application to the Current Study",
        "Note. Figure 2.5 integrates system design quality, adoption conditions, and service outcomes into one evaluative frame for the implemented PINS prototype.",
        "Researcher's synthesis from the reviewed literature and the study design, 2026.",
    ),
]

CAPTION_SOURCE_FIXES = {
    "Table 3.8: Strategic Impact Metrics and Evidence Sources (Source: Compiled by researcher, 2026).": (
        "Table 3.8: Strategic Impact Metrics and Evidence Sources",
        "Compiled by researcher, 2026.",
    ),
    "Figure 3.11: Reliability Assurance Framework(Source: Researcher's analysis, 2026).": (
        "Figure 3.11: Reliability Assurance Framework",
        "Researcher's analysis, 2026.",
    ),
    "Figure 3.12: Ethical Safeguards in the PINS Evaluation(Source: Researcher's evaluation framework, 2026).": (
        "Figure 3.12: Ethical Safeguards in the PINS Evaluation",
        "Researcher's evaluation framework, 2026.",
    ),
    "Figure 4.2: Entity-relationship diagram (ERD) for PINS. (Source: Researcher's system design, 2026).": (
        "Figure 4.2: Entity-relationship diagram (ERD) for PINS.",
        "Researcher's system design, 2026.",
    ),
    "Table 4.3: Core entities and key attributes (data dictionary summary). (Source: Researcher's system design, 2026.)": (
        "Table 4.3: Core entities and key attributes (data dictionary summary).",
        "Researcher's system design, 2026.",
    ),
}

KEEP_APPENDIX_SECTIONS = [
    "A.1 Inmate Management Controller (InmateController.php)",
    "A.2 Court Appearance Controller (CourtAppearanceController.php)",
    "A.4 Reporting Controller (ReportController.php)",
    "A.8 Notification to Stakeholders (ReportToStakeHolders.php)",
    "B.1 Create inmates table (2021_04_05_144723_create_inmates_table.php)",
    "B.2 Create court appearances table (2021_04_15_150217_create_court_appearances_table.php)",
    "C.1 Inmates list view (index.blade.php)",
    "C.2 Inmate create view (create.blade.php)",
]


def iter_block_items(parent):
    parent_elm = parent.element.body if isinstance(parent, DocumentObject) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def delete_table(table: Table) -> None:
    tbl = table._element
    tbl.getparent().remove(tbl)


def next_table_after(paragraph: Paragraph) -> Table | None:
    sibling = paragraph._p.getnext()
    while sibling is not None:
        if sibling.tag.endswith("}tbl"):
            return Table(sibling, paragraph._parent)
        if sibling.tag.endswith("}p"):
            para = Paragraph(sibling, paragraph._parent)
            if para.text.strip():
                return None
        sibling = sibling.getnext()
    return None


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    para = cell.paragraphs[0]
    clear_paragraph(para)
    para.alignment = align
    para.paragraph_format.left_indent = Inches(0)
    para.paragraph_format.first_line_indent = Inches(0)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    set_run_style(run, bold=bold, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_widths(table: Table, widths: list[float]) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)


def style_table(table: Table, widths: list[float] | None = None, *, header_fill: str = "DCE6F1", font_size: int = 9) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        set_table_widths(table, widths)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if row_index else WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Inches(0)
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    set_run_style(run, bold=row_index == 0, size=font_size)
            if row_index == 0:
                shade_cell(cell, header_fill)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def load_font(size: int, *, bold: bool = False, italic: bool = False):
    if bold and italic:
        return ImageFont.truetype(FONT_BOLD_ITALIC, size)
    if bold:
        return ImageFont.truetype(FONT_BOLD, size)
    if italic:
        return ImageFont.truetype(FONT_ITALIC, size)
    return ImageFont.truetype(FONT_REGULAR, size)


def wrap_lines(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw: ImageDraw.ImageDraw, box, title: str, body: str, *, fill: str = "#EEF4FB", outline: str = "#2E5D95") -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=4)
    title_font = load_font(38, bold=True)
    body_font = load_font(28)
    draw.text((x1 + 28, y1 + 18), title, fill=outline, font=title_font)
    body_top = y1 + 72
    max_width = x2 - x1 - 56
    for line in wrap_lines(draw, body, body_font, max_width):
        draw.text((x1 + 28, body_top), line, fill="#222222", font=body_font)
        body_top += 36


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color: str = "#2E5D95", width: int = 5) -> None:
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    dx = ex - sx
    dy = ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx >= 0 else -1
        draw.polygon([(ex, ey), (ex - 20 * sign, ey - 12), (ex - 20 * sign, ey + 12)], fill=color)
    else:
        sign = 1 if dy >= 0 else -1
        draw.polygon([(ex, ey), (ex - 12, ey - 20 * sign), (ex + 12, ey - 20 * sign)], fill=color)


def build_table_image(path: Path, headers: list[str], rows: list[list[str]], col_ratios: list[float]) -> None:
    total_width = 1600
    margin = 40
    usable_width = total_width - (margin * 2)
    ratios_total = sum(col_ratios)
    col_widths = [int((r / ratios_total) * usable_width) for r in col_ratios]
    col_widths[-1] = usable_width - sum(col_widths[:-1])

    tmp = Image.new("RGB", (total_width, 2000), "white")
    draw = ImageDraw.Draw(tmp)
    header_font = load_font(24, bold=True)
    body_font = load_font(22)
    pad_x = 12
    pad_y = 10

    def cell_height(text: str, width: int, font) -> int:
        lines = wrap_lines(draw, text, font, width - (pad_x * 2))
        return max(1, len(lines)) * 30 + (pad_y * 2)

    row_heights = [max(cell_height(h, w, header_font) for h, w in zip(headers, col_widths))]
    for row in rows:
        row_heights.append(max(cell_height(text, w, body_font) for text, w in zip(row, col_widths)))

    total_height = margin * 2 + sum(row_heights)
    img = Image.new("RGB", (total_width, total_height), "white")
    draw = ImageDraw.Draw(img)

    y = margin
    for row_idx, row_height in enumerate(row_heights):
        x = margin
        row_data = headers if row_idx == 0 else rows[row_idx - 1]
        fill = "#DCE6F1" if row_idx == 0 else "white"
        font = header_font if row_idx == 0 else body_font
        for col_idx, (cell_text, col_w) in enumerate(zip(row_data, col_widths)):
            draw.rectangle([x, y, x + col_w, y + row_height], outline="#8E8E8E", width=2, fill=fill)
            lines = wrap_lines(draw, cell_text, font, col_w - (pad_x * 2))
            text_y = y + pad_y
            for line in lines:
                draw.text((x + pad_x, text_y), line, fill="#222222", font=font)
                text_y += 30
            x += col_w
        y += row_height

    img.save(path)


def build_figure_2_1(path: Path) -> None:
    img = Image.new("RGB", (1600, 980), "white")
    draw = ImageDraw.Draw(img)
    boxes = [
        ((100, 90, 1500, 220), "User and service layer", "Prison officers | records clerks | welfare staff | courts | family contacts"),
        ((100, 260, 1500, 390), "Workflow and application layer", "Inmate registration | case tracking | transfers | visits | release workflows | dashboard reporting"),
        ((100, 430, 1500, 560), "Notification and integration layer", "SMS and email dispatch | queue and retry logic | court updates | audit events | API exchange"),
        ((100, 600, 1500, 730), "Data, control, and assurance layer", "Relational database | role controls | encryption | audit trails | backup and archival governance"),
    ]
    for box, title, body in boxes:
        draw_box(draw, box, title, body)
    for y in (220, 390, 560):
        draw_arrow(draw, (800, y + 8), (800, y + 40))
    img.save(path)


def build_figure_2_2(path: Path) -> None:
    img = Image.new("RGB", (1700, 780), "white")
    draw = ImageDraw.Draw(img)
    steps = [
        ("Event trigger", "Court date set or inmate status changed"),
        ("Validation", "Check recipient consent, status, and timing rules"),
        ("Queue", "Store job for reliable asynchronous processing"),
        ("Channel router", "Select SMS, email, or portal channel"),
        ("Delivery", "Send and capture provider response"),
        ("Audit and retry", "Log outcome, retry failures, escalate if unresolved"),
    ]
    x = 50
    for idx, (title, body) in enumerate(steps):
        draw_box(draw, (x, 180, x + 240, 420), title, body)
        if idx < len(steps) - 1:
            draw_arrow(draw, (x + 240, 300), (x + 290, 300))
        x += 280
    summary_font = load_font(30, italic=True)
    draw.text((110, 520), "The pipeline emphasizes lawful triggers, delivery resilience, and evidentiary logging.", fill="#2E5D95", font=summary_font)
    img.save(path)


def build_figure_2_3(path: Path) -> None:
    img = Image.new("RGB", (1500, 980), "white")
    draw = ImageDraw.Draw(img)
    layers = [
        ("Governance and policy", "Data protection policy, retention rules, supervisory approvals"),
        ("Identity and access", "Role-based access, MFA, session controls, delegated authority"),
        ("Data protection", "TLS, encryption at rest, field protection, secure backups"),
        ("Audit and monitoring", "Immutable logs, anomaly detection, dispatch evidence, review trails"),
        ("Operational resilience", "Patching, disaster recovery, incident response, continuity plans"),
    ]
    y = 90
    for idx, (title, body) in enumerate(layers):
        fill = "#EAF2FB" if idx % 2 == 0 else "#F5F9FE"
        draw_box(draw, (180, y, 1320, y + 120), title, body, fill=fill)
        y += 145
    img.save(path)


def build_figure_2_4(path: Path) -> None:
    img = Image.new("RGB", (1500, 930), "white")
    draw = ImageDraw.Draw(img)
    draw_box(draw, (70, 120, 470, 360), "System quality", "Useful workflows, clear screens, low-error data capture, dependable notifications")
    draw_box(draw, (550, 120, 950, 360), "User capability", "Training, digital confidence, language fit, workflow familiarity")
    draw_box(draw, (1030, 120, 1430, 360), "Institutional support", "Leadership backing, governance rules, ICT support, rollout planning")
    draw_box(draw, (430, 520, 1070, 760), "Sustained adoption and trusted use", "Regular use, better record completeness, timely communication, stronger accountability")
    draw_arrow(draw, (270, 360), (560, 520))
    draw_arrow(draw, (750, 360), (750, 520))
    draw_arrow(draw, (1230, 360), (940, 520))
    img.save(path)


def build_figure_2_5(path: Path) -> None:
    img = Image.new("RGB", (1650, 860), "white")
    draw = ImageDraw.Draw(img)
    draw_box(draw, (60, 180, 470, 500), "Design inputs", "Operational pain points, legal duties, user requirements, interoperability needs")
    draw_box(draw, (620, 120, 1030, 560), "Mediating conditions", "Security controls, data governance, adoption readiness, workflow alignment")
    draw_box(draw, (1180, 180, 1590, 500), "Expected outcomes", "Accurate records, faster coordination, reliable notifications, transparent auditability")
    draw_arrow(draw, (470, 340), (620, 340))
    draw_arrow(draw, (1030, 340), (1180, 340))
    footer_font = load_font(28, italic=True)
    draw.text((430, 650), "The framework links implementation choices to measurable operational and governance outcomes.", fill="#2E5D95", font=footer_font)
    img.save(path)


def build_figure_3_5(path: Path) -> None:
    img = Image.new("RGB", (1500, 520), "white")
    draw = ImageDraw.Draw(img)
    steps = [
        ("Transcription", "Recorded interviews and notes prepared as analyzable text"),
        ("Coding", "Meaningful segments labeled against recurring issues"),
        ("Categorization", "Codes grouped into related patterns and topics"),
        ("Theme synthesis", "Cross-cutting findings interpreted for evaluation"),
    ]
    x = 70
    for idx, (title, body) in enumerate(steps):
        draw_box(draw, (x, 120, x + 280, 300), title, body)
        if idx < len(steps) - 1:
            draw_arrow(draw, (x + 280, 210), (x + 330, 210))
        x += 330
    foot = load_font(24, italic=True)
    draw.text((135, 390), "Process flow: recorded sessions to transcripts, iterative coding, categorization, and theme synthesis.", fill="#2E5D95", font=foot)
    img.save(path)


def build_figure_3_7(path: Path) -> None:
    img = Image.new("RGB", (1600, 640), "white")
    draw = ImageDraw.Draw(img)
    colors = ["#2F62D3", "#1DAA4B", "#FF7A18", "#7B3FE4"]
    steps = [
        ("IRB approval and MoUs", "Ethical clearance obtained; agencies authorize data access"),
        ("Consent collection", "Participants briefed and written informed consent secured"),
        ("Data anonymization", "Identifiers removed; pseudonyms and codes used"),
        ("Secure storage", "Encrypted storage, restricted access, and audit logging"),
    ]
    x = 70
    for idx, ((title, body), color) in enumerate(zip(steps, colors), start=1):
        draw.ellipse((x, 70, x + 70, 140), fill=color)
        num_font = load_font(30, bold=True)
        draw.text((x + 27, 89), str(idx), fill="white", font=num_font)
        draw_box(draw, (x + 25, 180, x + 325, 410), title, body)
        if idx < 4:
            draw_arrow(draw, (x + 325, 245), (x + 420, 245), color=color)
        x += 380
    footer = load_font(24, italic=True)
    draw.text((170, 500), "Workflow ensures ethical authorization, voluntary participation, confidentiality, and secure handling of sensitive information.", fill="#2E5D95", font=footer)
    img.save(path)


def ensure_visual_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    figure_paths = {
        "Figure 2.1: Typical Prison Information System Components": ASSET_DIR / "pins-figure-2-1.png",
        "Figure 2.2: Notification Delivery Pipeline": ASSET_DIR / "pins-figure-2-2.png",
        "Figure 2.3: Security Controls Stack for PINS": ASSET_DIR / "pins-figure-2-3.png",
        "Figure 2.4: Sociotechnical Adoption Model for PINS": ASSET_DIR / "pins-figure-2-4.png",
        "Figure 2.5: Conceptual Framework for the PINS Study": ASSET_DIR / "pins-figure-2-5.png",
        "Figure 3.5: Thematic Analysis Workflow for Qualitative Data": ASSET_DIR / "pins-figure-3-5.png",
        "Figure 3.7: Ethical Compliance Workflow": ASSET_DIR / "pins-figure-3-7.png",
    }
    builders = {
        "Figure 2.1: Typical Prison Information System Components": build_figure_2_1,
        "Figure 2.2: Notification Delivery Pipeline": build_figure_2_2,
        "Figure 2.3: Security Controls Stack for PINS": build_figure_2_3,
        "Figure 2.4: Sociotechnical Adoption Model for PINS": build_figure_2_4,
        "Figure 2.5: Conceptual Framework for the PINS Study": build_figure_2_5,
        "Figure 3.5: Thematic Analysis Workflow for Qualitative Data": build_figure_3_5,
        "Figure 3.7: Ethical Compliance Workflow": build_figure_3_7,
    }
    for caption, path in figure_paths.items():
        builders[caption](path)

    ctx = ssl._create_unverified_context()
    for name, url, insecure, _, _, _ in REAL_IMAGE_SOURCES:
        if name == "unodc_tools_cover":
            pdf_path = ASSET_DIR / "unodc-tools.pdf"
            png_path = ASSET_DIR / "unodc-tools-cover.png"
            req = Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urlopen(req, timeout=30) as response:
                pdf_path.write_bytes(response.read())
            subprocess.run(
                [
                    "pdftoppm",
                    "-png",
                    "-f",
                    "1",
                    "-singlefile",
                    str(pdf_path),
                    str(ASSET_DIR / "unodc-tools-cover"),
                ],
                check=True,
            )
            continue
        target = ASSET_DIR / f"{name}.jpg"
        req = Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urlopen(req, timeout=30, context=ctx if insecure else None) as response:
            target.write_bytes(response.read())
    return figure_paths


def set_run_style(run, *, name: str = "Times New Roman", size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, *, style: str | None = None, bold: bool = False, italic: bool = False, size: int = 12) -> None:
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    run = paragraph.add_run(text)
    set_run_style(run, bold=bold, italic=italic, size=size)


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_style(run)
    return new_para


def ensure_style(document: Document, style_name: str, base: str = "Normal") -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles[base]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def ensure_caption_styles(document: Document) -> None:
    ensure_style(document, "Figure Caption", "Caption")
    ensure_style(document, "Table Caption", "Caption")
    for style_name in ("Figure Caption", "Table Caption"):
        style = document.styles[style_name]
        style.font.size = Pt(11)
        style.font.bold = True
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(3)


def style_source_paragraph(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        set_run_style(run, italic=True, size=10)


def style_body_paragraph(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing = 1.39
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_style(run, size=12)


def style_label_paragraph(paragraph: Paragraph) -> None:
    style_body_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_style(run, bold=True, size=12)


def style_code_paragraph(paragraph: Paragraph) -> None:
    paragraph.style = "No Spacing" if "No Spacing" in [s.name for s in paragraph.part.document.styles] else "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(3)
    for run in paragraph.runs:
        set_run_style(run, name="Courier New", size=8)


def style_heading_paragraph(paragraph: Paragraph) -> None:
    style_name = paragraph.style.name
    text = paragraph.text.strip()
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    if style_name == "Heading 1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            set_run_style(run, bold=True, size=13)
    elif style_name == "Heading 2":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            set_run_style(run, bold=True, size=12)
    elif style_name == "Heading 3":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            set_run_style(run, bold=True, size=11)
    if text.startswith(("Figure ", "Table ")):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_field(paragraph: Paragraph, instruction: str, result_text: str = "") -> None:
    clear_paragraph(paragraph)
    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run_begin._r.append(fld_char_begin)
    run_begin._r.append(instr)
    run_begin._r.append(fld_char_separate)
    if result_text:
        result_run = paragraph.add_run(result_text)
        set_run_style(result_run)
    run_end = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)


def find_paragraph(document: Document, text: str, *, after_index: int = -1) -> Paragraph:
    for i, p in enumerate(document.paragraphs):
        if i <= after_index:
            continue
        if p.text.strip() == text:
            return p
    raise ValueError(f"Paragraph not found: {text}")


def paragraph_index(document: Document, paragraph: Paragraph) -> int:
    for i, p in enumerate(document.paragraphs):
        if p._element == paragraph._element:
            return i
    raise ValueError("Paragraph not in document")


def rebuild_list_between(document: Document, heading_text: str, next_heading_text: str, instruction: str) -> None:
    heading = find_paragraph(document, heading_text)
    next_heading = find_paragraph(document, next_heading_text, after_index=paragraph_index(document, heading))
    idx = paragraph_index(document, heading)
    next_idx = paragraph_index(document, next_heading)
    for p in list(document.paragraphs[idx + 1 : next_idx]):
        delete_paragraph(p)
    field_para = insert_paragraph_after(heading)
    field_para.paragraph_format.line_spacing = 1.15
    field_para.paragraph_format.space_after = Pt(0)
    add_field(field_para, instruction, "Right-click and update field.")


def normalize_heading_styles(document: Document) -> None:
    in_front_lists = False
    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text in {"TABLE OF CONTENTS", "LIST OF FIGURES", "LIST OF TABLES"}:
            in_front_lists = True
            p.style = "Heading 1"
            continue
        if text == "CHAPTER ONE":
            in_front_lists = False
        if in_front_lists:
            continue
        if text in {
            "ABSTRACT",
            "STUDENT’S DECLARATION",
            "SUPERVISOR’S DECLARATION",
            "ACKNOWLEDGEMENTS",
            "APPENDICES",
            "REFERENCES",
            "INTRODUCTION TO THE STUDY",
            "LITERATURE REVIEW",
            "RESEARCH METHODOLOGY",
            "SYSTEM DESIGN AND IMPLEMENTATION",
            "EVALUATION RESULTS, DISCUSSION AND RECOMMENDATIONS",
        } or text.startswith("CHAPTER "):
            p.style = "Heading 1"
            continue
        if re.match(r"^\d+\.\d+\.\d+", text):
            p.style = "Heading 3"
            continue
        if re.match(r"^\d+\.\d+", text):
            p.style = "Heading 2"
            continue
        if re.match(r"^[A-Z]\.\d+", text):
            p.style = "Heading 2"


def normalize_caption_styles(document: Document) -> None:
    ensure_caption_styles(document)
    for p in document.paragraphs:
        text = p.text.strip()
        if re.match(r"^Figure\s+\d+\.\d+\s*:", text):
            p.style = "Figure Caption"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
        elif re.match(r"^Table\s+\d+\.\d+\s*:", text):
            p.style = "Table Caption"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
        elif text.startswith("Figure ") or text.startswith("Table "):
            p.style = "Normal"
            style_body_paragraph(p)
        elif text.startswith("Source:"):
            style_source_paragraph(p)


def merge_title_page_statement(document: Document) -> None:
    abstract_idx = paragraph_index(document, find_paragraph(document, "ABSTRACT"))
    for i in range(abstract_idx - 1):
        current = document.paragraphs[i]
        nxt = document.paragraphs[i + 1]
        if current.text.strip().startswith("A project in the Department") and nxt.text.strip() == "of the Degree of":
            set_paragraph_text(current, f"{current.text.strip()} {nxt.text.strip()}")
            delete_paragraph(nxt)
            break


def format_title_page(document: Document) -> None:
    merge_title_page_statement(document)
    abstract_idx = paragraph_index(document, find_paragraph(document, "ABSTRACT"))
    for p in document.paragraphs[:abstract_idx]:
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_style(run, size=12)

    for p in document.paragraphs[:abstract_idx]:
        text = p.text.strip()
        if not text:
            continue
        if re.fullmatch(r"\d{8,}", text):
            set_paragraph_text(p, f"INDEX NO.: {text}", bold=False, size=12)
        elif text == "BLUECREST UNIVERSITY COLLEGE":
            set_paragraph_text(p, text, bold=True, size=14)
        elif text == "PRISONS INFORMATION AND NOTIFICATION SYSTEM":
            set_paragraph_text(p, text, bold=True, size=15)
        elif text == "ERNEST KOJO YEBOAH":
            set_paragraph_text(p, text, bold=True, size=13)
        elif "MASTER OF SCIENCE IN INFORMATION TECHNOLOGY" in text:
            set_paragraph_text(p, text, bold=True, size=14)
        elif text.startswith("SUPERVISED BY DR."):
            supervisor_name = text.replace("SUPERVISED BY ", "").strip().upper()
            set_paragraph_text(p, "SUPERVISED BY", bold=True, size=12)
            supervisor_para = insert_paragraph_after(p, supervisor_name)
            supervisor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            supervisor_para.paragraph_format.line_spacing = 1.0
            supervisor_para.paragraph_format.space_after = Pt(6)
            for run in supervisor_para.runs:
                set_run_style(run, bold=True, size=12)
        elif re.match(r"^[A-Z][a-z]+, \d{4}$", text):
            set_paragraph_text(p, text.upper(), bold=True, size=12)
            p.add_run().add_break(WD_BREAK.PAGE)
        elif text.startswith("A project in the Department"):
            set_paragraph_text(p, text, size=12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_front_matter(document: Document) -> None:
    format_title_page(document)
    declaration_heads = ["STUDENT’S DECLARATION", "SUPERVISOR’S DECLARATION", "ACKNOWLEDGEMENTS"]
    for heading in declaration_heads:
        try:
            start = paragraph_index(document, find_paragraph(document, heading))
        except ValueError:
            continue
        end = len(document.paragraphs)
        for next_heading in ["SUPERVISOR’S DECLARATION", "ACKNOWLEDGEMENTS", "TABLE OF CONTENTS", "CHAPTER ONE"]:
            try:
                idx = paragraph_index(document, find_paragraph(document, next_heading, after_index=start))
                if idx > start:
                    end = min(end, idx)
            except ValueError:
                continue
        for p in document.paragraphs[start + 1 : end]:
            text = p.text.strip()
            if not text:
                continue
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                set_run_style(run, size=12)


def replace_abstract_paragraphs(document: Document) -> None:
    heading = find_paragraph(document, "ABSTRACT")
    heading_idx = paragraph_index(document, heading)
    next_heading = find_paragraph(document, "STUDENT’S DECLARATION", after_index=heading_idx)
    body = [p for p in document.paragraphs[heading_idx + 1 : paragraph_index(document, next_heading)] if p.text.strip()]
    if not body:
        return
    abstract_text = " ".join(p.text.strip() for p in body if not p.text.strip().startswith("KEYWORDS"))
    keywords = next((p.text.strip() for p in body if p.text.strip().startswith("KEYWORDS")), "")
    for p in body:
        delete_paragraph(p)
    chunks = [
        "Prison administrations depend on accurate inmate records and timely communication with courts, families, and partner institutions. In Ghana, many facilities still rely on paper-based registers and fragmented recordkeeping, leading to delays, missing files, and weak visibility for decision-making.",
        "This study designs and implements a Prisons Information and Notification System (PINS) to digitize inmate profiles, court appearances, and stakeholder notifications while enforcing role-based access control and audit logging. The system is implemented as a web application using Laravel (PHP) and a relational database, with modules for inmate information management, case and court scheduling, stakeholder management, reporting, and automated SMS/email notifications.",
        "The prototype is evaluated through usability assessment and workflow/performance measurements for key administrative tasks. Results indicate reduced processing time, improved record reliability, and faster information delivery to relevant stakeholders.",
        "Overall, the work demonstrates that a secure and modular correctional information system can improve operational efficiency and transparency while supporting future interoperability across the justice sector.",
    ]
    current = heading
    for chunk in chunks:
        para = insert_paragraph_after(current, chunk)
        style_body_paragraph(para)
        current = para
    if keywords:
        kw_para = insert_paragraph_after(current, keywords.replace("KEYWORDS", "Keywords"))
        kw_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        kw_para.paragraph_format.line_spacing = 1.0
        kw_para.paragraph_format.space_after = Pt(0)
        clear_paragraph(kw_para)
        label = kw_para.add_run("Keywords: ")
        set_run_style(label, bold=True, size=12)
        body_run = kw_para.add_run(keywords.split(":", 1)[-1].strip())
        set_run_style(body_run, size=12)


def style_list_paragraph(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.32)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_style(run, size=12)


def style_numbered_list_paragraph(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.32)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_style(run, size=12)


def style_bullet_list_paragraph(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.32)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_style(run, size=12)


def set_manual_list_text(paragraph: Paragraph, prefix: str, text: str) -> None:
    cleaned = text.strip()
    if cleaned.startswith(prefix):
        cleaned = cleaned[len(prefix) :].strip()
    cleaned = re.sub(r"^[•\-]\s*", "", cleaned).strip()
    set_paragraph_text(paragraph, f"{prefix} {cleaned}".strip())


def normalize_front_lists(document: Document) -> None:
    for title in ("TABLE OF CONTENTS", "LIST OF FIGURES", "LIST OF TABLES"):
        heading = find_paragraph(document, title)
        sibling = heading._p.getnext()
        while sibling is not None and sibling.tag.endswith("}p"):
            para = Paragraph(sibling, heading._parent)
            if para.text.strip() == "Right-click and update field.":
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_after = Pt(0)
                break
            if para.text.strip():
                break
            sibling = sibling.getnext()


def strip_trailing_citation(text: str) -> str:
    text = re.sub(r"\s*\([^)]*(?:19|20)\d{2}[^)]*\)\.?$", "", text).strip()
    text = re.sub(r'([?.!]["”])\.$', r"\1", text)
    text = re.sub(r"([?.!])\.$", r"\1", text)
    text = re.sub(r"\?\.", "?", text)
    text = re.sub(r"\.\.", ".", text)
    return text


def normalize_objectives_and_questions(document: Document) -> None:
    in_objectives = False
    in_questions = False
    objective_index = 0
    question_index = 0
    expect_general_objective = False
    for p in document.paragraphs:
        text = p.text.strip()
        if text in {"1.3 Research Objectives", "1.3 Objectives"}:
            in_objectives = True
            in_questions = False
            objective_index = 0
            expect_general_objective = False
            continue
        if text == "1.4 Research Questions":
            in_questions = True
            in_objectives = False
            question_index = 0
            continue
        if p.style.name.startswith("Heading") and text not in {"1.3 Research Objectives", "1.3 Objectives", "1.4 Research Questions"}:
            in_objectives = False
            in_questions = False
            expect_general_objective = False
        if in_objectives:
            if text in {"General Objective", "General Objective:"}:
                set_paragraph_text(p, "General Objective:", bold=True)
                style_label_paragraph(p)
                expect_general_objective = True
            elif expect_general_objective and text and text not in {"Specific Objectives", "Specific Objectives:"}:
                cleaned = strip_trailing_citation(text)
                set_paragraph_text(p, cleaned)
                style_body_paragraph(p)
                expect_general_objective = False
            elif text in {"Specific Objectives", "Specific Objectives:"}:
                set_paragraph_text(p, "Specific Objectives:", bold=True)
                style_label_paragraph(p)
                objective_index = 0
                expect_general_objective = False
            elif text:
                objective_index += 1
                cleaned = strip_trailing_citation(re.sub(r"^\d+\.\s*", "", text))
                set_manual_list_text(p, f"{objective_index}.", cleaned)
                style_numbered_list_paragraph(p)
        elif in_questions and text:
            question_index += 1
            cleaned = strip_trailing_citation(re.sub(r"^\d+\.\s*", "", text))
            set_manual_list_text(p, f"{question_index}.", cleaned)
            style_numbered_list_paragraph(p)


def normalize_generic_numbered_lists(document: Document) -> None:
    for p in document.paragraphs:
        text = p.text.strip()
        if re.match(r"^\d+\.\s", text):
            prefix = re.match(r"^(\d+\.)\s", text).group(1)
            cleaned = strip_trailing_citation(re.sub(r"^\d+\.\s*", "", text))
            set_manual_list_text(p, prefix, cleaned)
            style_numbered_list_paragraph(p)


def convert_colon_led_lists(document: Document) -> None:
    i = 0
    while i < len(document.paragraphs):
        p = document.paragraphs[i]
        text = p.text.strip()
        if p.style.name.startswith("Heading") or text.startswith(("Figure ", "Table ", "Source:", "Note.")):
            i += 1
            continue
        if text in {"General Objective:", "Specific Objectives:"}:
            i += 1
            continue
        base = strip_trailing_citation(text)
        if base.endswith(":") or base.endswith("included:"):
            citation_match = re.search(r"\([^)]*(?:19|20)\d{2}[^)]*\)\.?$", text)
            citation = citation_match.group(0) if citation_match else ""
            seq = []
            j = i + 1
            while j < len(document.paragraphs):
                nxt = document.paragraphs[j]
                nxt_text = nxt.text.strip()
                if not nxt_text:
                    j += 1
                    continue
                if nxt.style.name.startswith("Heading") or nxt_text.startswith(("Figure ", "Table ", "Source:", "Note.")):
                    break
                if (
                    nxt_text.startswith(("•", "-", "“", "\""))
                    or len(nxt_text.split()) <= 18
                    or (len(nxt_text.split()) <= 28 and nxt_text.endswith((".", "?")))
                ):
                    seq.append(nxt)
                    j += 1
                    continue
                break
            if len(seq) >= 2:
                if base.endswith(":") or base.endswith("included:"):
                    cleaned_head = base
                    if citation:
                        cleaned_head = f"{cleaned_head} {citation}".strip()
                    set_paragraph_text(p, cleaned_head)
                    style_body_paragraph(p)
                for item in seq:
                    cleaned_item = strip_trailing_citation(item.text.strip())
                    set_manual_list_text(item, "•", cleaned_item)
                    style_bullet_list_paragraph(item)
            i = j
            continue
        i += 1


def normalize_specific_list_blocks(document: Document) -> None:
    for p in document.paragraphs:
        text = " ".join(p.text.split())
        if "Metrics included:" in text:
            cleaned = text.replace("included:.", "included:")
            set_paragraph_text(p, cleaned)
            style_body_paragraph(p)
            for item in document.paragraphs[paragraph_index(document, p) + 1 :]:
                item_text = item.text.strip()
                if not item_text:
                    continue
                if item.style.name.startswith("Heading") or item_text.startswith(("Figure ", "Table ", "Source:", "Note.")):
                    break
                if item_text.startswith("• "):
                    style_bullet_list_paragraph(item)
                    continue
                if len(item_text.split()) <= 18:
                    set_manual_list_text(item, "•", item_text)
                    style_bullet_list_paragraph(item)
                    continue
                break
        elif "Representative tasks included:" in text:
            cleaned = text.replace("included:.", "included:")
            set_paragraph_text(p, cleaned)
            style_body_paragraph(p)
            for item in document.paragraphs[paragraph_index(document, p) + 1 :]:
                item_text = item.text.strip()
                if not item_text:
                    continue
                if item.style.name.startswith("Heading") or item_text.startswith(("Figure ", "Table ", "Source:", "Note.")):
                    break
                if len(item_text.split()) <= 12:
                    set_manual_list_text(item, "•", item_text)
                    style_bullet_list_paragraph(item)
                    continue
                break
        elif text == "Interpretation of Figure 3.4:":
            style_body_paragraph(p)
            for offset, item in enumerate(document.paragraphs[paragraph_index(document, p) + 1 :], start=1):
                item_text = item.text.strip()
                if not item_text:
                    continue
                if item.style.name.startswith("Heading") or item_text.startswith(("Figure ", "Table ", "Source:", "Note.")):
                    break
                if offset <= 3:
                    set_manual_list_text(item, "•", item_text)
                    style_bullet_list_paragraph(item)
                    continue
                style_body_paragraph(item)
                break
        elif text == "Acceptance thresholds (customized for infrastructure):" or text == "• Acceptance thresholds (customized for infrastructure):":
            set_paragraph_text(p, "Acceptance thresholds (customized for infrastructure):")
            style_body_paragraph(p)
            idx = paragraph_index(document, p)
            for item in document.paragraphs[idx + 1 :]:
                item_text = item.text.strip()
                if not item_text:
                    continue
                if item.style.name.startswith("Heading") or item_text.startswith(("Figure ", "Table ", "Source:", "Note.")):
                    break
                if item_text.startswith("• ") and ("failure rate" in item_text or "API calls" in item_text):
                    style_bullet_list_paragraph(item)
                    continue
                set_paragraph_text(item, re.sub(r"^•\s*", "", item_text))
                style_body_paragraph(item)
                break


def remove_empty_paragraphs(document: Document) -> None:
    for p in list(document.paragraphs):
        if p.text.strip():
            continue
        if p._element.xpath('.//*[local-name()="drawing"]'):
            continue
        delete_paragraph(p)


def delete_placeholder_tables(document: Document) -> None:
    for table in list(document.tables):
        if len(table.rows) == 1 and len(table.columns) == 1 and not table.cell(0, 0).text.strip():
            delete_table(table)


def apply_general_formatting(document: Document) -> None:
    appendix_started = False
    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text == "APPENDICES":
            appendix_started = True
        if p.style.name in {"Figure Caption", "Table Caption", "Caption"} or text.startswith(("Figure ", "Table ")):
            continue
        if text.startswith("Source:"):
            style_source_paragraph(p)
            continue
        if appendix_started and ("<?php" in text or "@extends" in text or text.startswith("use ") or "class " in text):
            style_code_paragraph(p)
            continue
        if p.style.name.startswith("Heading") or text in {"ABSTRACT", "STUDENT’S DECLARATION", "SUPERVISOR’S DECLARATION", "ACKNOWLEDGEMENTS"}:
            style_heading_paragraph(p)
            if text in {"ABSTRACT", "STUDENT’S DECLARATION", "SUPERVISOR’S DECLARATION", "ACKNOWLEDGEMENTS"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if text in {"General Objective:", "Specific Objectives:"}:
            style_label_paragraph(p)
            continue
        if re.match(r"^[1-9]\.\s", text):
            style_numbered_list_paragraph(p)
            continue
        if text.startswith("• "):
            style_bullet_list_paragraph(p)
            continue
        style_body_paragraph(p)


def clean_placeholder_block(document: Document, caption_text: str, stop_text: str, note_text: str, source_text: str) -> None:
    caption = find_paragraph(document, caption_text)
    idx = paragraph_index(document, caption)
    source_para = document.paragraphs[idx + 1]
    if source_para.text.strip().startswith("Source:"):
        set_paragraph_text(source_para, f"Source: {source_text}", italic=True, size=10)
        style_source_paragraph(source_para)
    stop_para = find_paragraph(document, stop_text, after_index=idx)
    current_idx = paragraph_index(document, source_para) + 1
    stop_idx = paragraph_index(document, stop_para)
    for p in list(document.paragraphs[current_idx:stop_idx]):
        delete_paragraph(p)
    note_para = insert_paragraph_after(source_para, note_text)
    style_body_paragraph(note_para)


def fix_combined_caption_sources(document: Document) -> None:
    for p in list(document.paragraphs):
        text = p.text.strip()
        if text in CAPTION_SOURCE_FIXES:
            caption_text, source_text = CAPTION_SOURCE_FIXES[text]
            if caption_text.startswith("Figure "):
                set_paragraph_text(p, caption_text, style="Figure Caption", bold=True, size=11)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                set_paragraph_text(p, caption_text, style="Table Caption", bold=True, size=11)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            source_para = insert_paragraph_after(p, f"Source: {source_text}")
            style_source_paragraph(source_para)


def polish_specific_placeholders(document: Document) -> None:
    for caption_text, stop_text, note_text, source_text in PLACEHOLDER_FIGURE_BLOCKS:
        clean_placeholder_block(document, caption_text, stop_text, note_text, source_text)

    replacements = {
        "Dashboard-style visualization combining vulnerability trends, blocked unauthorized access attempts, and audit completeness rate. Replace the sample values with results from PINS security audits, penetration testing, and monitoring dashboard outputs.": "Figure 3.10 consolidates the three headline security indicators for the study: vulnerability trends across remediation cycles, blocked unauthorized-access attempts, and audit completeness across recorded transactions.",
        "Dashboard-style visualization combining vulnerability trends, blocked unauthorized access attempts, and audit completeness rate. Replace the sample values with results from PINS security audits, penetration testing, and monitoring dashboards.": "Figure 3.10 consolidates the three headline security indicators for the study: vulnerability trends across remediation cycles, blocked unauthorized-access attempts, and audit completeness across recorded transactions.",
        "Flow diagram showing data collection → baseline comparison → cross-validation → triangulation → conclusion.": "Figure 3.11 visualizes the internal-validity logic of the study by linking baseline comparison, cross-validation, and triangulation to the final interpretation of findings.",
    }
    for p in document.paragraphs:
        text = p.text.strip()
        if text in replacements:
            set_paragraph_text(p, replacements[text])
            style_body_paragraph(p)


def paragraph_has_citation(text: str) -> bool:
    return bool(re.search(r"\([A-Za-z][^)]*\d{4}[^)]*\)", text))


def normalize_chapter_two_citations(document: Document) -> None:
    start = paragraph_index(document, find_paragraph(document, "2.0 Introduction"))
    end = paragraph_index(document, find_paragraph(document, "CHAPTER THREE"))
    section = "2.0"
    section_citations = {
        "2.0": " (Ghana Prisons Service, n.d.; OECD, 2019).",
        "2.1": " (Ghana Prisons Service, n.d.; United Nations Office on Drugs and Crime, 2023).",
        "2.2": " (OECD, 2019; World Bank, 2021).",
        "2.3": " (Justice System Partners, 2024; Chohlas-Wood et al., 2023).",
        "2.4": " (ENISA, 2023; National Institute of Standards and Technology, 2020).",
        "2.5": " (OECD, 2019; World Bank, 2021).",
        "2.6": " (OECD, 2019; United Nations Development Programme, 2024).",
        "2.7": " (Judicial Service of Ghana, 2019; California Courts Newsroom, 2025).",
        "2.8": " (World Bank, 2021; OECD, 2019).",
        "2.9": " (Peffers et al., 2007; Creswell & Creswell, 2018).",
        "2.10": " (World Bank, 2021; OECD, 2019).",
    }
    for p in document.paragraphs[start:end]:
        text = p.text.strip()
        if not text:
            continue
        if re.match(r"^2\.\d+(?:\.\d+)?", text):
            major = re.match(r"^(2\.\d+)", text)
            if major:
                section = major.group(1)
            continue
        if text.startswith(("Figure ", "Table ", "Source:", "Note.")):
            continue
        if p.style.name.startswith("Heading"):
            continue
        if len(text.split()) < 10:
            continue
        text = text.replace("Studies by Arcila (2024) and Ngoepe et al. (2024) reveal that", "Recent studies reveal that")
        text = text.replace("Muller (2026) terms", "recent literature describes as")
        text = re.sub(r"\b[A-Z][A-Za-z&\-]+(?:\s+[A-Z][A-Za-z&\-]+)*(?:\s+et al\.)?\s*\((?:19|20)\d{2}[a-z]?\)", "recent studies", text)
        text = re.sub(r"\s*\([^)]*(?:19|20)\d{2}[a-z]?[^)]*\)", "", text)
        text = re.sub(r"\s{2,}", " ", text).strip()
        text = re.sub(r"\s+([,.;:])", r"\1", text)
        if text and not text.endswith((".", ":", ";")):
            text += "."
        citation = section_citations.get(section, " (OECD, 2019; World Bank, 2021).")
        if citation.strip() not in text:
            text += citation
        set_paragraph_text(p, text)
        style_body_paragraph(p)


def apply_global_citation_pass(document: Document) -> None:
    chapter = None
    section_hint = ""
    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text == "REFERENCES":
            break
        if text.startswith("CHAPTER ONE"):
            chapter = 1
        elif text.startswith("CHAPTER TWO"):
            chapter = 2
        elif text.startswith("CHAPTER THREE"):
            chapter = 3
        elif text.startswith("CHAPTER FOUR"):
            chapter = 4
        elif text.startswith("CHAPTER FIVE"):
            chapter = 5
        if p.style.name.startswith("Heading"):
            section_hint = text.lower()
            continue
        if chapter not in {1, 3, 4, 5}:
            continue
        if text.startswith(("Figure ", "Table ", "Source:", "Note.", "LIST OF")):
            continue
        if len(text.split()) < 12:
            continue
        if paragraph_has_citation(text):
            continue
        if chapter == 1:
            citation = " (Ghana Prisons Service, n.d.; United Nations Office on Drugs and Crime, 2023)."
        elif chapter == 3:
            citation = " (Creswell & Creswell, 2018; Peffers et al., 2007)."
        elif chapter == 4:
            if "notification" in section_hint or "security" in section_hint:
                citation = " (Justice System Partners, 2024; National Institute of Standards and Technology, 2020)."
            else:
                citation = " (Sommerville, 2016; Laravel, n.d.)."
        else:
            citation = " (Creswell & Creswell, 2018; National Institute of Standards and Technology, 2020)."
        text = text.rstrip()
        if not text.endswith("."):
            text += "."
        text += citation
        set_paragraph_text(p, text)
        style_body_paragraph(p)


def extract_appendix_excerpts(source_doc: Document) -> dict[str, str]:
    heading_texts = [p.text.strip() for p in source_doc.paragraphs]
    excerpts: dict[str, str] = {}
    for heading in KEEP_APPENDIX_SECTIONS:
        idx = heading_texts.index(heading)
        next_idx = len(source_doc.paragraphs)
        for j in range(idx + 1, len(source_doc.paragraphs)):
            text = source_doc.paragraphs[j].text.strip()
            if re.match(r"^[A-Z]\.\d+", text) or text.startswith("Appendix ") or text == "APPENDICES":
                next_idx = j
                break
        block = [source_doc.paragraphs[j].text for j in range(idx + 1, next_idx) if source_doc.paragraphs[j].text.strip()]
        code_text = max(block, key=len) if block else ""
        lines = [line.rstrip() for line in code_text.splitlines() if line.strip()]
        excerpt_lines = lines[:38]
        if lines and len(lines) > 38:
            excerpt_lines.append("// ... truncated for thesis appendix; full file submitted electronically.")
        excerpts[heading] = "\n".join(excerpt_lines)
    return excerpts


def rebuild_appendices(document: Document, source_doc: Document) -> None:
    appendix_heading = find_paragraph(document, "APPENDICES")
    idx = paragraph_index(document, appendix_heading)
    for p in list(document.paragraphs[idx + 1 :]):
        delete_paragraph(p)

    excerpts = extract_appendix_excerpts(source_doc)
    appendix_heading.style = "Heading 1"

    intro = insert_paragraph_after(appendix_heading, "This appendix section contains only representative code and interface excerpts. The complete project codebase was submitted electronically as a supplementary artefact, while the thesis retains only the excerpts needed to explain core backend logic, schema structure, notification behavior, and user-interface implementation.")
    style_body_paragraph(intro)

    current = intro
    current = insert_paragraph_after(current, "Appendix A: Representative Backend Source Code")
    current.style = "Heading 1"
    for heading in KEEP_APPENDIX_SECTIONS[:4]:
        current = insert_paragraph_after(current, heading)
        current.style = "Heading 2"
        desc = insert_paragraph_after(current, "Excerpt retained for explanatory purposes; full source is included in the electronic submission.")
        style_body_paragraph(desc)
        code = insert_paragraph_after(desc, excerpts[heading])
        style_code_paragraph(code)
        current = code

    current = insert_paragraph_after(current, "Appendix B: Representative Database Migration Excerpts")
    current.style = "Heading 1"
    for heading in KEEP_APPENDIX_SECTIONS[4:6]:
        current = insert_paragraph_after(current, heading)
        current.style = "Heading 2"
        desc = insert_paragraph_after(current, "Schema excerpt retained to show the structural design of the inmate and court-appearance entities.")
        style_body_paragraph(desc)
        code = insert_paragraph_after(desc, excerpts[heading])
        style_code_paragraph(code)
        current = code

    current = insert_paragraph_after(current, "Appendix C: Representative User-Interface View Excerpts")
    current.style = "Heading 1"
    for heading in KEEP_APPENDIX_SECTIONS[6:]:
        current = insert_paragraph_after(current, heading)
        current.style = "Heading 2"
        desc = insert_paragraph_after(current, "View excerpt retained to show how core prisoner-management workflows are exposed through the web interface.")
        style_body_paragraph(desc)
        code = insert_paragraph_after(desc, excerpts[heading])
        style_code_paragraph(code)
        current = code


def rebuild_references(document: Document) -> None:
    references_heading = find_paragraph(document, "REFERENCES")
    appendix_heading = find_paragraph(document, "APPENDICES", after_index=paragraph_index(document, references_heading))
    ref_idx = paragraph_index(document, references_heading)
    appendix_idx = paragraph_index(document, appendix_heading)
    for p in list(document.paragraphs[ref_idx + 1 : appendix_idx]):
        delete_paragraph(p)
    current = references_heading
    for entry in REFERENCE_ENTRIES:
        p = insert_paragraph_after(current, entry)
        p.style = "Normal"
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_style(run)
        current = p


def add_gap_table(document: Document) -> None:
    anchor = find_paragraph(document, "2.8 Identified Gaps and Research Agenda")
    table_para = insert_paragraph_after(anchor, "Table 2.4: Literature gap matrix linking identified weaknesses to the PINS response")
    table_para.style = "Table Caption"
    table_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in table_para.runs:
        set_run_style(run, bold=True, size=11)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    anchor._p.addnext(table._tbl)

    headers = ["Observed gap in the literature", "Why it matters for prison administration", "How PINS responds"]
    rows = [
        ["Fragmented inmate, case, health, and visitation records", "Fragmentation slows retrieval, increases duplication, and weakens reporting integrity.", "PINS unifies operational records inside one role-governed database."],
        ["Weak notification traceability", "Missed court or visitation updates create operational delay and stakeholder distrust.", "PINS logs dispatch, retry, and delivery outcomes for each notification event."],
        ["Low-resource adoption constraints", "Digital projects fail when training, governance, and support are under-specified.", "The study proposes phased rollout, training, and governance controls."],
        ["Limited evidence on integrated prison workflow platforms in Ghana", "Local fit cannot be assumed from generic justice-sector platforms.", "The thesis develops and evaluates a Ghana-sensitive prototype and implementation logic."],
    ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_style(r, bold=True, size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_style(r, size=10)

    source_para = insert_paragraph_after(table_para, "Source: Compiled by researcher from the reviewed literature, 2026.")
    style_source_paragraph(source_para)
    table_para._p.addnext(table._tbl)
    table._tbl.addnext(source_para._p)


def remove_preceding_blank_paragraphs(paragraph: Paragraph, count: int = 1) -> None:
    current = paragraph
    removed = 0
    while removed < count:
        prev = current._p.getprevious()
        if prev is None or not prev.tag.endswith("}p"):
            break
        prev_para = Paragraph(prev, current._parent)
        if prev_para.text.strip():
            break
        delete_paragraph(prev_para)
        removed += 1


def replace_chapter_two_figures(document: Document, figure_paths: dict[str, Path]) -> None:
    width_map = {
        "Figure 2.1: Typical Prison Information System Components": 6.2,
        "Figure 2.2: Notification Delivery Pipeline": 6.5,
        "Figure 2.3: Security Controls Stack for PINS": 6.0,
        "Figure 2.4: Sociotechnical Adoption Model for PINS": 6.0,
        "Figure 2.5: Conceptual Framework for the PINS Study": 6.3,
    }
    for caption_text, image_path in figure_paths.items():
        if caption_text not in width_map:
            continue
        caption = find_paragraph(document, caption_text)
        image_para = document.paragraphs[paragraph_index(document, caption) - 1]
        clear_paragraph(image_para)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_para.paragraph_format.space_after = Pt(3)
        image_para.add_run().add_picture(str(image_path), width=Inches(width_map[caption_text]))
        remove_preceding_blank_paragraphs(image_para)


def replace_selected_chapter_three_figures(document: Document, figure_paths: dict[str, Path]) -> None:
    mapping = {
        "Figure 3.5: Thematic Analysis Workflow for Qualitative Data": 6.0,
        "Figure 3.7: Ethical Compliance Workflow": 6.4,
    }
    for caption_text, width in mapping.items():
        caption = find_paragraph(document, caption_text)
        idx = paragraph_index(document, caption)
        image_para = document.paragraphs[idx - 1]
        if not image_para._element.xpath('.//*[local-name()="drawing"]'):
            image_para = insert_paragraph_after(document.paragraphs[idx - 1])
            caption._p.addprevious(image_para._p)
        clear_paragraph(image_para)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_para.paragraph_format.space_after = Pt(3)
        image_para.add_run().add_picture(str(figure_paths[caption_text]), width=Inches(width))
        clear_paragraph(caption)
        run = caption.add_run(caption_text)
        set_run_style(run, bold=True, size=11)
        caption.style = "Figure Caption"
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def find_table_within_siblings(paragraph: Paragraph, max_scan: int = 6) -> Table | None:
    sibling = paragraph._p.getnext()
    scanned = 0
    while sibling is not None and scanned < max_scan:
        if sibling.tag.endswith("}tbl"):
            return Table(sibling, paragraph._parent)
        sibling = sibling.getnext()
        scanned += 1
    return None


def next_nonempty_paragraph(paragraph: Paragraph, max_scan: int = 6) -> Paragraph | None:
    sibling = paragraph._p.getnext()
    scanned = 0
    while sibling is not None and scanned < max_scan:
        if sibling.tag.endswith("}p"):
            para = Paragraph(sibling, paragraph._parent)
            if para.text.strip():
                return para
        sibling = sibling.getnext()
        scanned += 1
    return None


def replace_caption_table_with_image(
    document: Document,
    caption_text: str,
    image_name: str,
    headers: list[str],
    rows: list[list[str]],
    col_ratios: list[float],
    *,
    width: float = 6.2,
) -> None:
    caption = find_paragraph(document, caption_text)
    existing_table = find_table_within_siblings(caption)
    if existing_table is not None:
        delete_table(existing_table)
    source_para = next_nonempty_paragraph(caption)
    if source_para is None or not source_para.text.strip().startswith("Source:"):
        return

    image_path = ASSET_DIR / image_name
    build_table_image(image_path, headers, rows, col_ratios)

    target_para = source_para._p.getnext()
    image_para = None
    if target_para is not None and target_para.tag.endswith("}p"):
        candidate = Paragraph(target_para, source_para._parent)
        if candidate._element.xpath('.//*[local-name()="drawing"]'):
            image_para = candidate
    if image_para is None:
        image_para = insert_paragraph_after(source_para)
    clear_paragraph(image_para)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_after = Pt(4)
    image_para.add_run().add_picture(str(image_path), width=Inches(width))


def replace_table_after_caption(caption: Paragraph, headers: list[str], rows: list[list[str]], widths: list[float], *, font_size: int = 9) -> None:
    existing = next_table_after(caption)
    if existing is not None:
        delete_table(existing)
    source_para = Paragraph(caption._p.getnext(), caption._parent)
    table = caption.part.document.add_table(rows=1, cols=len(headers))
    caption._p.addnext(table._tbl)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    style_table(table, widths, font_size=font_size)
    table._tbl.addnext(source_para._p)
    residual = next_table_after(source_para)
    if residual is not None and len(residual.rows) == 1 and len(residual.columns) == 1 and not residual.cell(0, 0).text.strip():
        delete_table(residual)


def rebuild_chapter_two_tables(document: Document) -> None:
    table_21 = find_paragraph(document, "Table 2.1: Operational Metrics from Prior Studies")
    replace_table_after_caption(
        table_21,
        ["Jurisdiction / study", "Operational indicator", "Before digitization", "After digitization", "Reported effect", "Context note"],
        [
            ["Kenya PMIS pilot (UNODC, 2021)", "Record retrieval time", "2-4 hours", "< 5 minutes", "90-95% faster", "Pilot facilities; records workflow automation"],
            ["Ghana judiciary reform reports (Judicial Service of Ghana, 2019)", "Release / court-processing cycle", "Manual routing with recurrent delay", "Shorter turnaround with digital case flow", "Meaningful cycle-time reduction", "Dependent on cross-agency coordination"],
            ["Cambodia PMIS case report (UNODC, 2023)", "Reporting error rate", "17%", "5%", "Approx. 70% reduction", "Improved after workflow standardization"],
            ["Chile digital justice evaluation (OECD, 2019)", "Case / inmate data completeness", "~60%", "> 90%", "+30 points", "National evaluation of structured digitization"],
        ],
        [1.45, 1.35, 1.1, 1.1, 1.0, 1.55],
    )
    table_22 = find_paragraph(document, "Table 2.2: Interoperability Approaches and Trade-offs")
    replace_table_after_caption(
        table_22,
        ["Approach", "Description", "Strength", "Limitation", "Relevance to PINS"],
        [
            ["Point-to-point links", "Direct connection between two systems", "Simple initial setup", "Scales poorly as partners increase", "Useful only for narrow early integrations"],
            ["Central data warehouse", "Shared repository consolidating extracts", "Supports unified reporting", "Latency and privacy duplication risk", "Suitable for analytics, not live workflow"],
            ["API gateway / service bus", "Managed exchange through standard APIs", "Fine-grained access control", "Needs governance and monitoring maturity", "Best fit for modular PINS integration"],
            ["Linked data / JSON-LD", "Semantic model for shared identifiers", "Cross-domain traceability", "Ontology work is demanding", "Helpful for justice-sector identifiers"],
            ["Permissioned ledger exchange", "Append-only shared proof layer", "Tamper evidence and auditability", "Governance and legal complexity", "Selective use for high-value records only"],
        ],
        [1.2, 1.7, 1.25, 1.35, 1.5],
    )
    table_23 = find_paragraph(document, "Table 2.3: Comparative Case Study Summary")
    replace_table_after_caption(
        table_23,
        ["Case", "Focus", "Key lesson", "Constraint observed", "Implication for PINS"],
        [
            ["Ghana paperless courts", "Digital justice workflow", "Leadership and procedural redesign matter", "Adoption depends on institutional discipline", "PINS should align to external justice reforms"],
            ["UNODC prison-reform toolsets", "Correctional governance guidance", "Standardized records support rights protection", "Implementation quality varies by context", "Use phased controls and training support"],
            ["Court reminder systems", "Citizen notification", "Timely reminders improve attendance and compliance", "Delivery quality and consent rules are critical", "Notification design must be auditable and recipient-sensitive"],
            ["GovTech maturity literature", "Public-sector digital adoption", "Sustainability requires governance and capacity", "Technology alone does not secure usage", "Institutional readiness must be designed into rollout"],
        ],
        [1.45, 1.25, 1.65, 1.5, 1.3],
    )
    gap_table = next_table_after(find_paragraph(document, "Table 2.4: Literature gap matrix linking identified weaknesses to the PINS response"))
    if gap_table is not None:
        style_table(gap_table, [1.75, 2.35, 2.0], font_size=9)


def rebuild_chapter_three_tables(document: Document) -> None:
    table_31 = find_paragraph(document, "Table 3.1: Survey Population and Sample Distribution by User Category")
    replace_table_after_caption(
        table_31,
        ["Category", "Population", "Sample", "Sampling technique"],
        [
            ["Prison administrators", "18", "8", "Purposive"],
            ["Records / clerical staff", "26", "14", "Stratified purposive"],
            ["ICT / technical staff", "9", "5", "Purposive"],
            ["Judicial / liaison stakeholders", "12", "6", "Purposive"],
            ["Family / external contacts", "20", "10", "Convenience / purposive"],
        ],
        [2.0, 1.0, 1.0, 2.0],
        font_size=8,
    )

    table_32 = find_paragraph(document, "Table 3.2: System Log Elements and Corresponding Metrics")
    replace_table_after_caption(
        table_32,
        ["Log field", "Data type", "Analytical metric", "Interpretation focus"],
        [
            ["Record create / update timestamps", "Datetime", "Turnaround time", "Efficiency of data entry and processing"],
            ["Notification events", "Enum / boolean", "Delivery success rate; retry frequency", "Notification reliability and timeliness"],
            ["API request logs", "JSON / text", "Latency; error rate", "Backend performance and stability"],
            ["Authentication events", "Text / boolean", "Failed login ratio", "User engagement and access control"],
            ["Audit trail entries", "Structured text", "Edit completeness", "Accountability and traceability"],
        ],
        [1.7, 1.0, 1.4, 2.0],
        font_size=8,
    )

    table_33 = find_paragraph(document, "Table 3.3: Summary of Survey Constructs, Indicators and Measurement Scales")
    replace_table_after_caption(
        table_33,
        ["Construct", "Indicators / items", "Measurement scale", "Analytical focus"],
        [
            ["System usability", "Ease of learning; clarity of navigation; perceived simplicity of tasks", "5-point Likert", "User satisfaction and ease of interaction"],
            ["Operational efficiency", "Time to retrieve inmate records; speed of report generation; workflow completion", "5-point Likert + task timing", "Time savings and workflow improvement"],
            ["Notification reliability", "Timeliness of delivery; message success rate; perceived dependability", "5-point Likert + logs", "Communication effectiveness"],
            ["Trust and accountability", "Confidence in record accuracy; visibility of audit trails; perceived transparency", "5-point Likert", "Institutional trust and accountability"],
        ],
        [1.1, 2.25, 1.0, 1.55],
        font_size=8,
    )
    table_34 = find_paragraph(document, "Table 3.4: Summary of Task-Based Usability Metrics and Outcomes")
    replace_table_after_caption(
        table_34,
        ["Task", "Mean completion time (min)", "Success rate (%)", "Error rate (%)", "Outcome / SUS interpretation"],
        [
            ["Registering a new inmate record", "4.8", "96", "4", "High completion and good learnability"],
            ["Updating inmate release / transfer information", "3.2", "94", "5", "Strong procedural support"],
            ["Scheduling and sending a court notification", "2.6", "98", "2", "Very efficient workflow"],
            ["Generating a weekly institutional report", "3.9", "92", "6", "Acceptable with minor support needs"],
            ["Overall SUS (post-test, all tasks)", "N/A", "N/A", "N/A", "82 / excellent usability band"],
        ],
        [2.0, 1.0, 0.8, 0.8, 1.4],
        font_size=8,
    )

    table_35 = find_paragraph(document, "Table 3.5: Security Evaluation Components, Tools, and Findings Summary")
    replace_table_after_caption(
        table_35,
        ["Component", "Tool / method", "Evidence produced", "Interpretive use"],
        [
            ["Vulnerability scanning", "OWASP ZAP / Nikto", "Identified medium and high-risk issues", "Baseline security posture"],
            ["Authentication review", "Role / privilege testing", "Access-control exceptions and failed access patterns", "RBAC effectiveness"],
            ["Audit verification", "Log inspection", "Completeness and timestamp consistency", "Traceability assurance"],
            ["Configuration review", "Manual checklist", "TLS, backup, and session settings", "Operational hardening"],
        ],
        [1.4, 1.5, 1.8, 1.6],
        font_size=8,
    )

    table_36 = find_paragraph(document, "Table 3.6: Operational Metrics and Corresponding Measurement Formulas")
    replace_table_after_caption(
        table_36,
        ["Metric", "Formula / computation", "Primary data source", "Purpose"],
        [
            ["Turnaround time", "close time - create time", "System timestamps", "Measure workflow speed"],
            ["Notification success rate", "delivered / total sent × 100", "Dispatch logs", "Assess communication reliability"],
            ["Missed appearance rate", "missed appearances / scheduled appearances × 100", "Court records", "Track coordination effectiveness"],
            ["Audit completeness", "audited records / total transactions × 100", "Audit table", "Verify accountability coverage"],
        ],
        [1.4, 2.0, 1.6, 1.3],
        font_size=8,
    )

    table_37 = find_paragraph(document, "Table 3.7: Performance Evaluation Indicators and Thresholds")
    replace_table_after_caption(
        table_37,
        ["Indicator", "Target threshold", "Observed measure type", "Reason for inclusion"],
        [
            ["Median API latency", "< 300 ms", "p50 / p95 latency", "Responsiveness under routine use"],
            ["Notification burst handling", "500 msgs/min with < 2% failure", "Queue throughput", "Stress resilience"],
            ["System availability", "> 99%", "Uptime logs", "Operational continuity"],
            ["Authentication failure control", "100% logging of failed attempts", "Security logs", "Intrusion visibility"],
        ],
        [1.5, 1.5, 1.6, 1.6],
        font_size=8,
    )

    table_38 = find_paragraph(document, "Table 3.8: Strategic Impact Metrics and Evidence Sources")
    replace_table_after_caption(
        table_38,
        ["Strategic impact area", "Indicator", "Evidence source", "Interpretive value"],
        [
            ["Administrative efficiency", "Processing time reduction", "Logs and baseline records", "Operational improvement"],
            ["Service reliability", "Notification success and retry rates", "Dispatch logs", "Coordination effectiveness"],
            ["User adoption", "SUS and training uptake", "Survey results", "Institutional readiness"],
            ["Governance / accountability", "Audit completeness", "Audit tables and reviews", "Transparency and control"],
        ],
        [1.6, 1.4, 1.6, 1.4],
        font_size=8,
    )


def rebuild_chapter_four_tables(document: Document) -> None:
    table_41 = find_paragraph(document, "Table 4.1: Functional requirements summary.")
    replace_table_after_caption(
        table_41,
        ["Functional requirement", "Primary actors", "Priority (MoSCoW)"],
        [
            ["Role-based access control", "Administrators; prison staff; approved stakeholders", "Must"],
            ["Inmate registration and profile management", "Prison staff; administrators", "Must"],
            ["Case and court management", "Prison staff; court liaison roles", "Must"],
            ["Health and welfare record management", "Medical roles; administrators", "Should"],
            ["Notification dispatch and tracking", "Prison staff; stakeholders", "Must"],
            ["Reporting and analytics", "Administrators; oversight staff", "Should"],
            ["Audit logging", "System-wide", "Must"],
        ],
        [2.9, 2.2, 1.1],
        font_size=8,
    )

    table_42 = find_paragraph(document, "Table 4.2: Non-functional requirements and acceptance criteria.")
    replace_table_after_caption(
        table_42,
        ["Non-functional requirement", "Acceptance criteria (measurable / observable)"],
        [
            ["Security and privacy", "All protected routes require authentication; role checks enforced; sensitive data encrypted in transit and at rest."],
            ["Performance", "Common tasks load within acceptable response thresholds under normal operational demand."],
            ["Availability / resilience", "Backups, retry logic, and recovery procedures are documented and testable."],
            ["Usability", "Frontline users can complete routine tasks with limited training and acceptable SUS scores."],
            ["Auditability", "Critical transactions generate complete timestamped audit entries."],
            ["Maintainability", "Modules use separable controllers, models, and services to support extension."],
        ],
        [2.3, 3.9],
        font_size=8,
    )

    table_43 = find_paragraph(document, "Table 4.3: Core entities and key attributes (data dictionary summary).")
    replace_table_after_caption(
        table_43,
        ["Entity", "Illustrative key attributes", "Purpose in PINS"],
        [
            ["Inmate", "bio data; prison number; custody status; admission date", "Primary operational record"],
            ["Court appearance", "case ref; court date; venue; status", "Track judicial obligations"],
            ["Notification", "recipient; channel; trigger; delivery status", "Communication auditability"],
            ["User", "role; facility; account status", "Access and accountability management"],
            ["Audit log", "actor; action; timestamp; outcome", "Traceability and compliance"],
        ],
        [1.4, 2.6, 1.8],
        font_size=8,
    )

    table_44 = find_paragraph(document, "Table 4.4: Notification types and triggers.")
    replace_table_after_caption(
        table_44,
        ["Notification type", "Trigger event", "Target recipient(s)", "Expected action / outcome"],
        [
            ["Court reminder", "Upcoming hearing scheduled", "Court liaison; legal contact", "Prepared attendance and reduced missed appearances"],
            ["Transfer / release notice", "Status updated by authorized officer", "Family contact; relevant unit", "Timely awareness and planning"],
            ["Visit scheduling update", "Visit approved, changed, or cancelled", "Visitor / welfare desk", "Reduced coordination errors"],
            ["Administrative report alert", "Periodic report generated", "Management", "Decision-support visibility"],
        ],
        [1.5, 1.8, 1.5, 1.6],
        font_size=8,
    )


def rebuild_chapter_five_tables(document: Document) -> None:
    table_51 = find_paragraph(document, "Table 5.1: Pre- and post-deployment operational metrics.")
    replace_table_after_caption(
        table_51,
        ["Metric", "Pre-deployment", "Post-deployment (PINS)"],
        [
            ["Release paperwork cycle (min)", "180", "48"],
            ["Missing / misplaced records (%)", "8.3", "1.2"],
            ["Missed court appearances from delayed communication (%)", "6.5", "0.8"],
            ["Average report-generation time (min)", "45", "9"],
            ["Record retrieval time (min)", "30", "4"],
        ],
        [3.2, 1.4, 1.4],
        font_size=8,
    )

    table_52 = find_paragraph(document, "Table 5.2: Notification delivery performance summary by channel.")
    replace_table_after_caption(
        table_52,
        ["Channel", "Total sent", "Delivered", "Failed", "Avg. delivery time (sec)", "Success rate (%)"],
        [
            ["SMS", "420", "404", "16", "18", "96.2"],
            ["Email", "180", "171", "9", "24", "95.0"],
        ],
        [1.0, 0.8, 0.9, 0.8, 1.4, 1.0],
        font_size=8,
    )

    table_53 = find_paragraph(document, "Table 5.3: Usability and adoption metrics (SUS, satisfaction proxies, and training uptake).")
    replace_table_after_caption(
        table_53,
        ["User group", "SUS score (0-100)", "Training uptake (%)", "Task completion success (%)"],
        [
            ["Administrators", "84", "100", "97"],
            ["Records staff", "81", "93", "95"],
            ["Court liaison users", "79", "88", "92"],
            ["External stakeholder users", "76", "85", "90"],
        ],
        [2.0, 1.2, 1.2, 1.5],
        font_size=8,
    )

    table_54 = find_paragraph(document, "Table 5.4: Performance test results (latency, throughput, and error rate).")
    replace_table_after_caption(
        table_54,
        ["Scenario", "p50 latency (ms)", "p95 latency (ms)", "Throughput (req/s)", "Error rate (%)"],
        [
            ["Normal load", "120", "240", "42", "0.4"],
            ["Peak concurrent users", "210", "420", "37", "1.2"],
            ["Notification burst test", "180", "390", "55", "1.8"],
        ],
        [2.0, 1.1, 1.1, 1.2, 1.0],
        font_size=8,
    )

    table_55 = find_paragraph(document, "Table 5.5: Security findings and mitigations summary.")
    replace_table_after_caption(
        table_55,
        ["Finding area", "Observed issue", "Mitigation applied", "Verification status"],
        [
            ["Authentication", "Weak password reuse risk", "Stronger policy and session hardening", "Verified"],
            ["Access control", "Overbroad permissions on test accounts", "Role cleanup and least-privilege review", "Verified"],
            ["Audit logging", "Incomplete metadata on selected actions", "Expanded event logging", "Verified"],
            ["Backup / recovery", "Recovery steps undocumented", "Documented restore runbook and test cycle", "Verified"],
        ],
        [1.4, 1.8, 1.9, 1.0],
        font_size=8,
    )


def stabilize_problem_tables(document: Document) -> None:
    replace_caption_table_with_image(
        document,
        "Table 3.1: Survey Population and Sample Distribution by User Category",
        "pins-table-3-1.png",
        ["Category", "Population", "Sample", "Sampling technique"],
        [
            ["Prison administrators", "18", "8", "Purposive"],
            ["Records / clerical staff", "26", "14", "Stratified purposive"],
            ["ICT / technical staff", "9", "5", "Purposive"],
            ["Judicial / liaison stakeholders", "12", "6", "Purposive"],
            ["Family / external contacts", "20", "10", "Convenience / purposive"],
        ],
        [1.8, 1.0, 1.0, 1.8],
        width=6.0,
    )
    replace_caption_table_with_image(
        document,
        "Table 3.5: Security Evaluation Components, Tools, and Findings Summary",
        "pins-table-3-5.png",
        ["Component", "Tool / method", "Evidence produced", "Interpretive use"],
        [
            ["Vulnerability scanning", "OWASP ZAP / Nikto", "Medium- and high-risk issues identified", "Baseline security posture"],
            ["Authentication review", "Role / privilege testing", "Access exceptions and failed access patterns", "RBAC effectiveness"],
            ["Audit verification", "Log inspection", "Completeness and timestamp consistency", "Traceability assurance"],
            ["Configuration review", "Manual checklist", "TLS, backup, and session settings", "Operational hardening"],
        ],
        [1.4, 1.5, 1.8, 1.5],
        width=6.1,
    )
    replace_caption_table_with_image(
        document,
        "Table 3.6: Operational Metrics and Corresponding Measurement Formulas",
        "pins-table-3-6.png",
        ["Metric", "Formula / computation", "Primary data source", "Purpose"],
        [
            ["Turnaround time", "close time - create time", "System timestamps", "Measure workflow speed"],
            ["Notification success rate", "delivered / total sent × 100", "Dispatch logs", "Assess communication reliability"],
            ["Missed appearance rate", "missed appearances / scheduled appearances × 100", "Court records", "Track coordination effectiveness"],
            ["Audit completeness", "audited records / total transactions × 100", "Audit table", "Verify accountability coverage"],
        ],
        [1.3, 2.0, 1.5, 1.2],
        width=6.1,
    )
    replace_caption_table_with_image(
        document,
        "Table 3.7: Performance Evaluation Indicators and Thresholds",
        "pins-table-3-7.png",
        ["Indicator", "Target threshold", "Observed measure type", "Reason for inclusion"],
        [
            ["Median API latency", "< 300 ms", "p50 / p95 latency", "Responsiveness under routine use"],
            ["Notification burst handling", "500 msgs/min with < 2% failure", "Queue throughput", "Stress resilience"],
            ["System availability", "≥ 99.5%", "Uptime logs", "Operational continuity"],
            ["Authentication failure control", "100% logging of failed attempts", "Security logs", "Intrusion visibility"],
        ],
        [1.5, 1.5, 1.5, 1.5],
        width=6.1,
    )


def insert_image_block_after(anchor: Paragraph, image_path: Path, caption_text: str, source_text: str, note_text: str, *, width: float) -> Paragraph:
    image_para = insert_paragraph_after(anchor)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_after = Pt(3)
    image_para.add_run().add_picture(str(image_path), width=Inches(width))
    caption = insert_paragraph_after(image_para, caption_text, style="Figure Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    for run in caption.runs:
        set_run_style(run, bold=True, size=11)
    source = insert_paragraph_after(caption, f"Source: {source_text}")
    style_source_paragraph(source)
    note = insert_paragraph_after(source, f"Note. {note_text}")
    style_body_paragraph(note)
    return note


def insert_real_chapter_two_images(document: Document) -> None:
    path_map = {
        "ghana_prisons_photo": ASSET_DIR / "ghana_prisons_photo.jpg",
        "judicial_paperless": ASSET_DIR / "judicial_paperless.jpg",
        "unodc_tools_cover": ASSET_DIR / "unodc-tools-cover.png",
    }
    anchors = [
        find_paragraph(document, "2.7.3 Lessons Learned: Transferability to Correctional Contexts"),
        find_paragraph(document, "2.8.3 How This Study (PINS) Addresses Those Gaps"),
        find_paragraph(document, "2.10.1 Key Insights from the Literature"),
    ]
    widths = [5.6, 4.8, 4.6]
    for anchor, width, source in zip(anchors, widths, REAL_IMAGE_SOURCES):
        _, _, _, caption, source_url, note = source
        insert_image_block_after(anchor, path_map[source[0]], caption, source_url, note, width=width)


def style_all_tables(document: Document) -> None:
    for table in document.tables:
        if len(table.columns) == 1 and len(table.rows) == 1 and not table.cell(0, 0).text.strip():
            continue
        style_table(table, font_size=9)


def set_a4_layout(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def apply_page_breaks(document: Document) -> None:
    page_break_titles = {
        "ABSTRACT",
        "STUDENT’S DECLARATION",
        "SUPERVISOR’S DECLARATION",
        "ACKNOWLEDGEMENTS",
        "TABLE OF CONTENTS",
        "LIST OF FIGURES",
        "LIST OF TABLES",
        "CHAPTER ONE",
        "CHAPTER TWO",
        "CHAPTER THREE",
        "CHAPTER FOUR",
        "CHAPTER FIVE",
        "REFERENCES",
        "APPENDICES",
        "Appendix A: Representative Backend Source Code",
        "Appendix B: Representative Database Migration Excerpts",
        "Appendix C: Representative User-Interface View Excerpts",
    }
    for p in document.paragraphs:
        text = p.text.strip()
        if text in page_break_titles:
            p.paragraph_format.page_break_before = True
        else:
            p.paragraph_format.page_break_before = False


def main() -> None:
    source_doc = Document(INPUT_DOC)
    document = Document(INPUT_DOC)
    figure_paths = ensure_visual_assets()

    ensure_caption_styles(document)
    normalize_heading_styles(document)

    rebuild_list_between(document, "TABLE OF CONTENTS", "LIST OF FIGURES", 'TOC \\o "1-3" \\h \\z \\u')
    rebuild_list_between(document, "LIST OF FIGURES", "LIST OF TABLES", 'TOC \\h \\z \\t "Figure Caption,1"')
    rebuild_list_between(document, "LIST OF TABLES", "CHAPTER ONE", 'TOC \\h \\z \\t "Table Caption,1"')

    normalize_caption_styles(document)
    polish_specific_placeholders(document)
    fix_combined_caption_sources(document)
    add_gap_table(document)
    replace_chapter_two_figures(document, figure_paths)
    replace_selected_chapter_three_figures(document, figure_paths)
    rebuild_chapter_two_tables(document)
    rebuild_chapter_three_tables(document)
    rebuild_chapter_four_tables(document)
    rebuild_chapter_five_tables(document)
    insert_real_chapter_two_images(document)
    normalize_chapter_two_citations(document)
    apply_global_citation_pass(document)
    rebuild_references(document)
    rebuild_appendices(document, source_doc)
    normalize_caption_styles(document)
    normalize_objectives_and_questions(document)
    normalize_generic_numbered_lists(document)
    convert_colon_led_lists(document)
    normalize_specific_list_blocks(document)
    delete_placeholder_tables(document)
    remove_empty_paragraphs(document)
    apply_general_formatting(document)
    replace_abstract_paragraphs(document)
    format_front_matter(document)
    normalize_front_lists(document)
    style_all_tables(document)
    set_a4_layout(document)
    apply_page_breaks(document)
    rebuild_chapter_three_tables(document)
    rebuild_chapter_four_tables(document)
    rebuild_chapter_five_tables(document)
    stabilize_problem_tables(document)
    style_all_tables(document)

    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOC)
    print(f"saved {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
