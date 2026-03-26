from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INPUT_DOC = Path("/Users/emmanuel/Desktop/apps/gerts/output/doc/GERTS_thesis_reworked_final_v5.docx")
OUTPUT_DOC = Path("/Users/emmanuel/Desktop/apps/gerts/output/doc/GERTS_thesis_reworked_final_v6.docx")


def insert_paragraph_after(paragraph, text: str = "", style_name: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._element.getparent().remove(new_paragraph._element)
    new_p.getparent().replace(new_p, new_paragraph._element)
    if style_name:
        new_paragraph.style = style_name
    if text:
        run = new_paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
    return new_paragraph


def insert_paragraph_before(paragraph, text: str = "", style_name: str | None = None):
    new_paragraph = paragraph.insert_paragraph_before(text)
    if style_name:
        new_paragraph.style = style_name
    return new_paragraph


def format_body(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Inches(0.0)
    paragraph.paragraph_format.left_indent = Inches(0.0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        run.bold = False


def format_source(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.first_line_indent = Inches(0.0)
    paragraph.paragraph_format.left_indent = Inches(0.0)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10)
        run.italic = True
        run.bold = False


def next_nonempty(paragraphs, start_index: int):
    idx = start_index + 1
    while idx < len(paragraphs):
        text = " ".join(paragraphs[idx].text.split()).strip()
        if text:
            return idx, paragraphs[idx], text
        idx += 1
    return None, None, ""


def previous_nonempty(paragraphs, start_index: int):
    idx = start_index - 1
    while idx >= 0:
        text = " ".join(paragraphs[idx].text.split()).strip()
        if text:
            return idx, paragraphs[idx], text
        idx -= 1
    return None, None, ""


def ensure_intro_after_heading(document: Document, heading_text: str, intro_text: str) -> None:
    paragraphs = document.paragraphs
    for i, paragraph in enumerate(paragraphs):
        text = " ".join(paragraph.text.split()).strip()
        if text != heading_text:
            continue
        _, nxt, next_text = next_nonempty(paragraphs, i)
        if nxt is None:
            inserted = insert_paragraph_after(paragraph, intro_text, "Normal")
            format_body(inserted)
            return
        if nxt.style.name.startswith("Heading"):
            inserted = insert_paragraph_after(paragraph, intro_text, "Normal")
            format_body(inserted)
            return
        if next_text.startswith("Figure ") or next_text.startswith("Table "):
            inserted = insert_paragraph_after(paragraph, intro_text, "Normal")
            format_body(inserted)
            return
        return


def ensure_lead_in_before_caption(document: Document, caption_text: str, lead_text: str) -> None:
    paragraphs = document.paragraphs
    for i, paragraph in enumerate(paragraphs):
        text = " ".join(paragraph.text.split()).strip()
        if text != caption_text:
            continue
        _, prev, prev_text = previous_nonempty(paragraphs, i)
        if prev is not None and prev_text == lead_text:
            return
        inserted = insert_paragraph_before(paragraph, lead_text, "Normal")
        format_body(inserted)
        return


def guess_source(caption: str) -> str:
    if caption.startswith(("Figure 2.", "Table 2.")):
        return "Compiled by researcher from literature reviewed, 2026."
    if caption.startswith(("Figure 3.", "Table 3.")):
        return "Compiled by researcher for the GERTS study, 2026."
    if caption.startswith(("Figure 4.", "Table 4.")):
        return "Compiled by researcher from the GERTS prototype, 2026."
    if caption.startswith(("Figure 5.", "Table 5.")):
        return "Compiled by researcher from GERTS prototype evaluation, 2026."
    return "Compiled by researcher, 2026."


def add_missing_caption_sources(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    for i, paragraph in enumerate(paragraphs):
        text = " ".join(paragraph.text.split()).strip()
        if paragraph.style.name not in {"Figure Caption", "Table Caption"}:
            continue
        _, nxt, next_text = next_nonempty(paragraphs, i)
        if nxt is not None and next_text.startswith("Source:"):
            continue
        source = insert_paragraph_after(paragraph, f"Source: {guess_source(text)}", "Normal")
        format_source(source)
        paragraphs = list(document.paragraphs)


def normalize_existing_sources(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split()).strip()
        if text.startswith("Source:"):
            paragraph.text = text
            format_source(paragraph)


def main() -> None:
    document = Document(INPUT_DOC)

    ensure_intro_after_heading(
        document,
        "2.1 Current Polling-Station Result Reporting, Transmission, and Verification in Ghana",
        "This section reviews how polling-station results are currently recorded, transmitted, verified, and interpreted within Ghana's election process. The focus is on the operational path from signed polling-station forms to collation and public communication, because these stages define where delays, evidence gaps, and disputes most commonly emerge (ACE Project, n.d.; CODEO, 2020).",
    )
    ensure_intro_after_heading(
        document,
        "2.3 Decentralized Architecture for Tamper-Evident Results Reporting",
        "This section examines decentralized and tamper-evident design options that can strengthen result reporting without changing the legal authority of the Electoral Commission. The review narrows its attention to permissioned ledger models, cryptographic integrity controls, governance arrangements, and public verification mechanisms that are directly relevant to GERTS rather than to full online voting systems (Jafar et al., 2021; Capocasale et al., 2023).",
    )
    ensure_intro_after_heading(
        document,
        "2.4 Prototype Design Constraints for Realistic Election Operations",
        "This section narrows the literature to the constraints that matter most when moving from theory to a deployable prototype. The discussion emphasizes legal compliance, cybersecurity, usability, and operational realism so that the proposed system remains trustworthy under actual Ghanaian election conditions and not only in an ideal technical model (Republic of Ghana, 2012; OWASP Foundation, 2023).",
    )
    ensure_lead_in_before_caption(
        document,
        "Table 2.1: Key concepts and working definitions",
        "Table 2.1 consolidates the working definitions adopted in this study so that later discussions of transparency, verification, evidence, auditability, and trust use a consistent vocabulary.",
    )

    add_missing_caption_sources(document)
    normalize_existing_sources(document)

    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
