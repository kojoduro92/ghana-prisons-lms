from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path

from docx import Document


INPUT_DOC = Path("/Users/emmanuel/Desktop/apps/gerts/output/doc/GERTS_thesis_reworked_final_v7.docx")
OUTPUT_DOC = Path("/Users/emmanuel/Desktop/apps/gerts/output/doc/GERTS_thesis_reworked_final_v9.docx")


CAPTION_RE = re.compile(r"^(Figure|Table) (\d+)\.(\d+)(:.*)$")


def build_mapping(document: Document, caption_style: str, chapters: set[int] | None = None) -> dict[str, str]:
    by_chapter: dict[int, list[int]] = defaultdict(list)
    for paragraph in document.paragraphs:
        if paragraph.style.name != caption_style:
            continue
        text = " ".join(paragraph.text.split()).strip()
        match = CAPTION_RE.match(text)
        if not match:
            continue
        _, chapter, number, _ = match.groups()
        chapter_num = int(chapter)
        if chapters is not None and chapter_num not in chapters:
            continue
        by_chapter[chapter_num].append(int(number))

    mapping: dict[str, str] = {}
    for chapter_num, seq in by_chapter.items():
        if all(b == a + 1 for a, b in zip(seq, seq[1:])):
            continue
        for idx, old in enumerate(seq, start=1):
            mapping[f"{chapter_num}.{old}"] = f"{chapter_num}.{idx}"
    return mapping


def replace_number_refs(text: str, mapping: dict[str, str]) -> str:
    placeholders: dict[str, str] = {}
    for idx, (old, new) in enumerate(sorted(mapping.items(), key=lambda item: len(item[0]), reverse=True), start=1):
        placeholder = f"__REF_{idx}__"
        placeholders[placeholder] = new
        text = re.sub(rf"(?<!\d){re.escape(old)}(?!\d)", placeholder, text)
    for placeholder, replacement in placeholders.items():
        text = text.replace(placeholder, replacement)
    return text


def apply_caption_mapping(text: str, label: str, mapping: dict[str, str]) -> str:
    if label not in text:
        return text
    return replace_number_refs(text, mapping)


def main() -> None:
    document = Document(INPUT_DOC)

    table_mapping = build_mapping(document, "Table Caption", {2})
    figure_mapping = build_mapping(document, "Figure Caption", {2, 3})

    for paragraph in document.paragraphs:
        text = paragraph.text
        if not text:
            continue

        updated = text
        if "Table" in updated or "Tables" in updated or paragraph.style.name == "Table Caption":
            updated = apply_caption_mapping(updated, "Table", table_mapping)
        if "Figure" in updated or "Figures" in updated or paragraph.style.name == "Figure Caption":
            updated = apply_caption_mapping(updated, "Figure", figure_mapping)

        if updated != text:
            paragraph.text = updated

    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOC)
    print("table_mapping", table_mapping)
    print("figure_mapping", figure_mapping)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
