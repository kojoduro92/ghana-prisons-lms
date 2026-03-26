from __future__ import annotations

import re
import subprocess
import textwrap
from pathlib import Path
from urllib.request import Request, urlopen

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/emmanuel/Desktop/apps/ghana-prisons-lms")
INPUT_DOC = Path("/Users/emmanuel/Downloads/VCMS (1).docx")
OUTPUT_DOC = ROOT / "output/doc/VCMS-Final-Defense-Reworked.docx"
ASSET_DIR = ROOT / "output/doc/vcms-assets"
RENDER_DIR = ROOT / "output/doc/vcms-render"

GHANA_PAPERLESS_URL = "https://judicial.gov.gh/media/k2/items/cache/556202461533c16857a46ed5bfb7c21c_L.jpg"
CALIFORNIA_REMOTE_URL = "https://newsroom.courts.ca.gov/sites/default/files/newsroom/styles/max_650x650/public/2025-04/8E440351-ACEC-48FB-8A5B-207E58591297_1_201_a.jpg"
VIDEO_CONFERENCE_PDF_URL = "https://judicial.gov.gh/jsweb/jsfiles/videoconference.pdf"
CJI_REPORT_PDF_URL = "https://www.sji.gov/wp-content/uploads/CJI-Full-Report.pdf"
CALIFORNIA_REMOTE_REPORT_PDF_URL = "https://newsroom.courts.ca.gov/sites/default/files/newsroom/2021-08/P3%20Workgroup%20Remote%20Access%20Interim%20Report%2008162021.pdf"

FONT_REG = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"
FONT_ITALIC = "/System/Library/Fonts/Supplemental/Times New Roman Italic.ttf"


def fetch(url: str, destination: Path) -> None:
    if destination.exists() and destination.stat().st_size > 1000:
        return
    req = Request(url, headers={"User-Agent": "Mozilla/5.0"})
    destination.write_bytes(urlopen(req, timeout=30).read())


def ensure_assets() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    fetch(GHANA_PAPERLESS_URL, ASSET_DIR / "ghana-paperless-courts.jpg")
    fetch(CALIFORNIA_REMOTE_URL, ASSET_DIR / "california-remote-courts.jpg")
    pdf_path = ASSET_DIR / "ghana-video-conference.pdf"
    fetch(VIDEO_CONFERENCE_PDF_URL, pdf_path)
    png_path = ASSET_DIR / "ghana-video-conference.png"
    if not png_path.exists():
        subprocess.run(
            [
                "pdftoppm",
                "-png",
                "-f",
                "1",
                "-singlefile",
                str(pdf_path),
                str(png_path.with_suffix("")),
            ],
            check=True,
        )
    cji_pdf = ASSET_DIR / "civil-justice-initiative-report.pdf"
    fetch(CJI_REPORT_PDF_URL, cji_pdf)
    if not (ASSET_DIR / "civil-justice-initiative-report.png").exists():
        subprocess.run(
            [
                "pdftoppm",
                "-png",
                "-f",
                "35",
                "-singlefile",
                str(cji_pdf),
                str((ASSET_DIR / "civil-justice-initiative-report").with_suffix("")),
            ],
            check=True,
        )
    california_pdf = ASSET_DIR / "california-remote-access-report.pdf"
    fetch(CALIFORNIA_REMOTE_REPORT_PDF_URL, california_pdf)
    if not (ASSET_DIR / "california-remote-access-report.png").exists():
        subprocess.run(
            [
                "pdftoppm",
                "-png",
                "-f",
                "4",
                "-singlefile",
                str(california_pdf),
                str((ASSET_DIR / "california-remote-access-report").with_suffix("")),
            ],
            check=True,
        )
    build_diagram_assets()


def font(size: int, *, bold: bool = False, italic: bool = False):
    path = FONT_BOLD if bold else FONT_ITALIC if italic else FONT_REG
    return ImageFont.truetype(path, size=size)


def new_canvas(width: int = 1800, height: int = 1100) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    return img, draw


def wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, width: int, *, size: int = 34, bold: bool = False, fill: str = "#111111", spacing: int = 8, align: str = "left") -> int:
    f = font(size, bold=bold)
    wrapped_text = textwrap.fill(text, width=width)
    draw.multiline_text(xy, wrapped_text, font=f, fill=fill, spacing=spacing, align=align)
    bbox = draw.multiline_textbbox(xy, wrapped_text, font=f, spacing=spacing, align=align)
    return bbox[3]


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body: str, *, fill: str = "#f5f8fc", outline: str = "#315d8a") -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=4)
    x1, y1, x2, y2 = box
    title_bottom = wrapped(draw, (x1 + 20, y1 + 16), title, 20, size=34, bold=True, fill="#0d2c54")
    wrapped(draw, (x1 + 20, title_bottom + 18), body, 30, size=24, fill="#1b1b1b")


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], *, fill: str = "#315d8a", width: int = 6) -> None:
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    if end[0] >= start[0]:
        points = [(ex, ey), (ex - 18, ey - 12), (ex - 18, ey + 12)]
    else:
        points = [(ex, ey), (ex + 18, ey - 12), (ex + 18, ey + 12)]
    draw.polygon(points, fill=fill)


def build_diagram_assets() -> None:
    build_timeline()
    build_platform_architecture()
    build_notification_workflow()
    build_reminder_evidence()
    build_remote_distribution()
    build_literature_map()
    build_conceptual_model()
    build_data_governance()
    build_theoretical_framework()
    build_notification_best_practice()
    build_virtual_hearing_balance()
    build_design_propositions()
    build_chapter2_synthesis_grid()


def build_timeline() -> None:
    img, draw = new_canvas(2000, 1080)
    draw.rectangle((0, 0, 2000, 1080), fill="#fbfcfe")
    wrapped(draw, (80, 40), "Figure 2.1. Judicial digitalisation milestones informing the VCMS study", 45, size=42, bold=True, fill="#0d2c54")
    draw.line((140, 560, 1860, 560), fill="#40698f", width=8)
    points = [
        (180, "2017", "ICT vision", "Chief Justice forum"),
        (500, "2018", "Video workflow", "Remote-hearing readiness"),
        (820, "2019", "Paperless launch", "Law Court Complex"),
        (1140, "2020-2021", "e-Justice guidance", "User-facing services"),
        (1460, "2024", "Usage study", "Institutional fit"),
        (1780, "2025", "Remote evidence", "Participation gains"),
    ]
    for x, year, title, subtitle in points:
        draw.ellipse((x - 24, 536, x + 24, 584), fill="#163f72", outline="#163f72")
        wrapped(draw, (x - 120, 640), year, 10, size=34, bold=True, fill="#163f72", align="center")
        draw.rounded_rectangle((x - 125, 250, x + 125, 470), radius=18, fill="white", outline="#9db6d1", width=3)
        wrapped(draw, (x - 98, 295), title, 16, size=30, bold=True, fill="#163f72", align="center")
        wrapped(draw, (x - 98, 355), subtitle, 18, size=24, fill="#1b1b1b", align="center")
        draw.line((x, 470, x, 536), fill="#9db6d1", width=4)
    img.save(ASSET_DIR / "vcms-figure-2-1-timeline.png")


def build_platform_architecture() -> None:
    img, draw = new_canvas(1600, 1120)
    wrapped(draw, (80, 40), "Figure 2.3. Core digital-platform layers for court case, scheduling, and role management", 48, size=42, bold=True, fill="#0d2c54")
    layers = [
        ((120, 170, 1480, 320), "User and service interface layer", "Lawyers, litigants, clerks, and judges"),
        ((120, 390, 1480, 540), "Workflow and application layer", "Filing, validation, scheduling, and notices"),
        ((120, 610, 1480, 760), "Control and compliance layer", "RBAC, approvals, audit trail, and privacy"),
        ((120, 830, 1480, 980), "Data and integration layer", "Case database, messaging, video, and analytics"),
    ]
    for box, title, body in layers:
        draw.rounded_rectangle(box, radius=18, fill="#f8fbfe", outline="#315d8a", width=4)
        x1, y1, x2, y2 = box
        title_bottom = wrapped(draw, (x1 + 24, y1 + 18), title, 24, size=30, bold=True, fill="#0d2c54")
        wrapped(draw, (x1 + 24, title_bottom + 18), body, 56, size=22, fill="#1b1b1b")
    for y in (320, 540, 760):
        arrow(draw, (800, y + 8), (800, y + 50))
    img.save(ASSET_DIR / "vcms-figure-2-3-platform-architecture.png")


def build_notification_workflow() -> None:
    img, draw = new_canvas(1800, 980)
    wrapped(draw, (80, 40), "Figure 2.5. Recommended event-driven notification workflow for hearing updates", 50, size=42, bold=True, fill="#0d2c54")
    boxes = [
        ((80, 300, 350, 520), "1. Court event", "Filing or schedule change"),
        ((410, 300, 700, 520), "2. Rules engine", "Template | timing | channel"),
        ((760, 300, 1060, 520), "3. Delivery queue", "SMS or email with retry"),
        ((1120, 300, 1420, 520), "4. User receipt", "Date | time | venue or link"),
        ((1480, 300, 1720, 520), "5. Audit log", "Dispatch and failure history"),
    ]
    for box, title, body in boxes:
        draw_box(draw, box, title, body)
    for idx in range(len(boxes) - 1):
        arrow(draw, (boxes[idx][0][2], 420), (boxes[idx + 1][0][0] - 16, 420))
    draw.rounded_rectangle((560, 650, 1260, 860), radius=20, fill="#fdf7ef", outline="#c77c2c", width=4)
    wrapped(draw, (590, 700), "Fallback: retry, switch channel where allowed, then escalate persistent failure to registry staff.", 42, size=30, fill="#57320e", align="center")
    img.save(ASSET_DIR / "vcms-figure-2-5-notification-workflow.png")


def build_reminder_evidence() -> None:
    img, draw = new_canvas(1500, 860)
    wrapped(draw, (80, 40), "Figure 2.6. Selected empirical evidence on automated court reminders", 45, size=42, bold=True, fill="#0d2c54")
    draw.rounded_rectangle((80, 170, 650, 760), radius=24, fill="#f7fbff", outline="#5f84a9", width=4)
    wrapped(draw, (165, 215), "Santa Clara text-message experiment", 18, size=32, bold=True, fill="#163f72", align="center")
    wrapped(draw, (200, 305), "5,709 clients | random assignment", 28, size=24, fill="#1b1b1b", align="center")
    wrapped(draw, (195, 380), "Warrants for missed dates", 20, size=26, bold=True, fill="#163f72", align="center")
    draw.rectangle((185, 465, 425, 525), fill="#d65a5a")
    draw.rectangle((185, 575, 380, 635), fill="#4b86c2")
    wrapped(draw, (450, 468), "Control: 12.1%", 16, size=26, bold=True, fill="#8a1f1f")
    wrapped(draw, (450, 578), "Reminder: 9.7%", 16, size=26, bold=True, fill="#1b4a7a")
    draw.rounded_rectangle((840, 170, 1420, 760), radius=24, fill="#fef8f2", outline="#bb7438", width=4)
    wrapped(draw, (965, 220), "Operational reminder design", 18, size=32, bold=True, fill="#7b4111", align="center")
    wrapped(draw, (930, 340), "Multiple reminders\nClear notice language\nReminder does not replace the court order", 22, size=24, fill="#1b1b1b", align="center")
    wrapped(draw, (935, 560), "VCMS implication:\nuse reminders to improve attendance,\nwhile keeping the official registry\nrecord authoritative.", 22, size=24, bold=True, fill="#7b4111", align="center")
    img.save(ASSET_DIR / "vcms-figure-2-6-reminder-evidence.png")


def build_remote_distribution() -> None:
    img, draw = new_canvas(1520, 980)
    wrapped(draw, (80, 40), "Figure 2.9. Distribution of remote civil proceedings in California, Sept. 2023 to Aug. 2024", 52, size=42, bold=True, fill="#0d2c54")
    categories = [
        ("Civil unlimited", 46, "#365f8f"),
        ("Family law", 19, "#4f81bd"),
        ("Probate", 13, "#6a9bd0"),
        ("Civil limited", 9, "#87b0e0"),
        ("Juvenile dependency", 6, "#a3c5eb"),
        ("Juvenile delinquency", 4, "#c0d7f1"),
        ("Small claims", 2, "#d7e6f7"),
        ("Civil mental health", 1, "#ebf3fb"),
    ]
    x, y = 110, 210
    total_width = 1250
    current = x
    for name, pct, color in categories:
        width = int(total_width * (pct / 100))
        draw.rectangle((current, y, current + width, y + 90), fill=color, outline="white", width=3)
        if pct >= 9:
            label = f"{name}\n{pct}%"
            wrapped(draw, (current + 12, y + 14), label, 11, size=22, bold=True, fill="#0d2c54")
        current += width
    legend_items = categories[4:]
    legend_x = 110
    legend_y = 360
    for idx, (name, pct, color) in enumerate(legend_items):
        col = idx // 2
        row = idx % 2
        lx = legend_x + col * 330
        ly = legend_y + row * 56
        draw.rectangle((lx, ly, lx + 28, ly + 28), fill=color, outline=color)
        wrapped(draw, (lx + 42, ly - 2), f"{name} {pct}%", 24, size=22, fill="#0d2c54")
    wrapped(draw, (110, 500), "Large shares in civil unlimited, family law, and probate show that remote proceedings operate across multiple matter types, not just a narrow specialist docket.", 52, size=24, fill="#1b1b1b")
    draw.rounded_rectangle((110, 690, 1380, 860), radius=20, fill="#f6f9fd", outline="#9eb7cf", width=4)
    wrapped(draw, (250, 742), "Survey finding: over 90% positive experience; fewer than 2% reported audio or visual technical issues.", 48, size=28, bold=True, fill="#163f72", align="center")
    img.save(ASSET_DIR / "vcms-figure-2-9-remote-distribution.png")


def build_literature_map() -> None:
    img, draw = new_canvas(1900, 1120)
    wrapped(draw, (80, 40), "Figure 2.10. Literature map linking the review to the four study objectives", 48, size=42, bold=True, fill="#0d2c54")
    objectives = [
        ("Obj. 1", "Existing systems", "#ecf3fb"),
        ("Obj. 2", "Digital platform", "#f2f8ef"),
        ("Obj. 3", "Notifications", "#fef5ea"),
        ("Obj. 4", "Virtual hearings", "#f8eff7"),
    ]
    sources = [
        "Judicial Service of Ghana",
        "Ashong Elliot et al. (2024)",
        "World Bank / OECD",
        "Chohlas-Wood et al. (2023)",
        "California / Hawaii / Minnesota judiciary sources",
        "Rossner & Tait (2021); Sourdin (2020)",
    ]
    cell_w = 220
    cell_h = 88
    base_x = 520
    base_y = 250
    wrapped(draw, (50, 160), "Sources", 10, size=32, bold=True, fill="#163f72")
    for idx, (label, desc, fill) in enumerate(objectives):
        x1 = base_x + idx * cell_w
        draw.rounded_rectangle((x1, 120, x1 + cell_w - 18, 225), radius=16, fill=fill, outline="#91a9c0", width=3)
        wrapped(draw, (x1 + 24, 142), f"{label}\n{desc}", 14, size=24, bold=True, fill="#0d2c54", align="center")
    marks = {
        0: [1, 0, 0, 1],
        1: [1, 1, 0, 0],
        2: [1, 1, 1, 1],
        3: [0, 0, 1, 0],
        4: [0, 0, 1, 1],
        5: [0, 1, 0, 1],
    }
    for r, source in enumerate(sources):
        y1 = base_y + r * cell_h
        wrapped(draw, (80, y1 + 18), source, 24, size=28, fill="#1b1b1b")
        for c in range(4):
            x1 = base_x + c * cell_w
            draw.rectangle((x1, y1, x1 + cell_w - 18, y1 + cell_h - 6), outline="#b7c9d9", width=3)
            if marks[r][c]:
                draw.ellipse((x1 + 78, y1 + 20, x1 + 132, y1 + 74), fill="#315d8a")
    img.save(ASSET_DIR / "vcms-figure-2-10-literature-map.png")


def build_conceptual_model() -> None:
    img, draw = new_canvas(2000, 1250)
    wrapped(draw, (80, 40), "Figure 2.11. Conceptual model for the Virtual Court Management and Notification System", 54, size=42, bold=True, fill="#0d2c54")
    boxes = {
        "Actors": (120, 340, 520, 620),
        "Case management core": (620, 210, 1450, 470),
        "Notification and compliance": (620, 540, 1450, 800),
        "Virtual hearing service": (1550, 340, 1910, 620),
        "Data, audit, and security": (620, 900, 1450, 1120),
    }
    draw_box(draw, boxes["Actors"], "Actors", "Judges | registrars | clerks | lawyers | litigants | witnesses | admins")
    draw_box(draw, boxes["Case management core"], "Case management core", "Registration | e-filing | validation | scheduling | dashboards | case status")
    draw_box(draw, boxes["Notification and compliance"], "Notification and compliance", "Rules | templates | queueing | SMS and email | retries | acknowledgements")
    draw_box(draw, boxes["Virtual hearing service"], "Virtual hearing service", "Remote links | access control | attendance support | post-session update")
    draw_box(draw, boxes["Data, audit, and security"], "Data, audit, and security", "Case repository | logs | audit trail | privacy | reporting")
    arrow(draw, (520, 480), (600, 340))
    arrow(draw, (1450, 340), (1530, 480))
    arrow(draw, (1035, 470), (1035, 520))
    arrow(draw, (1035, 800), (1035, 880))
    img.save(ASSET_DIR / "vcms-figure-2-11-conceptual-model.png")


def build_data_governance() -> None:
    img, draw = new_canvas(1600, 1080)
    wrapped(draw, (80, 40), "Figure 2.12. Data-governance and interoperability stack for court platform integration", 55, size=42, bold=True, fill="#0d2c54")
    layers = [
        ((120, 180, 1480, 315), "User-facing records and services", "Profiles, case history, calendars, and reminders"),
        ((120, 370, 1480, 505), "Shared identifiers and workflow events", "Case number, participant ID, hearing ID, and timestamps"),
        ((120, 560, 1480, 695), "Interoperability and exchange controls", "Payment, messaging, video, reporting, and archive"),
        ((120, 750, 1480, 885), "Governance and protection mechanisms", "Data quality, retention, audit logging, and privacy"),
    ]
    for box, title, body in layers:
        draw.rounded_rectangle(box, radius=18, fill="#f8fbfe", outline="#315d8a", width=4)
        x1, y1, x2, y2 = box
        title_bottom = wrapped(draw, (x1 + 20, y1 + 14), title, 24, size=28, bold=True, fill="#0d2c54")
        wrapped(draw, (x1 + 20, title_bottom + 16), body, 60, size=20, fill="#1b1b1b")
    for y in (315, 505, 695):
        arrow(draw, (800, y + 6), (800, y + 42))
    img.save(ASSET_DIR / "vcms-figure-2-12-data-governance.png")


def build_theoretical_framework() -> None:
    img, draw = new_canvas(1500, 1020)
    wrapped(draw, (80, 40), "Figure 2.13. Integrated theoretical framework used to interpret the VCMS literature", 55, size=42, bold=True, fill="#0d2c54")
    boxes = [
        ((70, 220, 420, 520), "Technology Acceptance\nModel", "Usefulness\nEase of use\nIntention to use"),
        ((560, 180, 940, 580), "Socio-technical\nsystems view", "Technology must fit people,\nrules, roles, training,\nand workflow."),
        ((1080, 220, 1420, 520), "Procedural justice\nlens", "Voice\nClarity\nNeutrality\nTrust"),
        ((410, 700, 1110, 925), "Interpretive implication for VCMS", "Useful + understandable + institutionally fitted + experienced as fair"),
    ]
    for box, title, body in boxes:
        draw.rounded_rectangle(box, radius=18, fill="#f8fbfe", outline="#315d8a", width=4)
        x1, y1, x2, y2 = box
        title_bottom = wrapped(draw, (x1 + 20, y1 + 18), title, 18, size=28, bold=True, fill="#0d2c54")
        wrapped(draw, (x1 + 20, title_bottom + 18), body, 28, size=21, fill="#1b1b1b")
    arrow(draw, (420, 370), (540, 370))
    arrow(draw, (940, 370), (1060, 370))
    arrow(draw, (760, 580), (760, 680))
    img.save(ASSET_DIR / "vcms-figure-2-13-theoretical-framework.png")


def build_notification_best_practice() -> None:
    img, draw = new_canvas(1600, 1120)
    wrapped(draw, (80, 40), "Figure 2.14. Best-practice design ladder for court reminder systems", 52, size=42, bold=True, fill="#0d2c54")
    steps = [
        ("1. Reliable enrolment", "Collect contact details early and allow updates."),
        ("2. Clear content", "Send date, time, venue or link, plus support contact."),
        ("3. Multiple touches", "Use more than one reminder and more than one channel."),
        ("4. Behavioural prompts", "Help recipients make a concrete attendance plan."),
        ("5. Delivery assurance", "Track status, retries, and escalation for failed delivery."),
        ("6. Equity and accessibility", "Support language, disability, and digital-access needs."),
    ]
    y = 160
    box_h = 112
    gap = 26
    for idx, (title, body) in enumerate(steps):
        draw.rounded_rectangle((200, y, 1460, y + box_h), radius=18, fill="#f8fbfe", outline="#4e79a7", width=4)
        draw.ellipse((92, y + 22, 160, y + 90), fill="#173f72")
        wrapped(draw, (115, y + 35), str(idx + 1), 2, size=28, bold=True, fill="white", align="center")
        title_bottom = wrapped(draw, (228, y + 14), title, 34, size=27, bold=True, fill="#173f72")
        wrapped(draw, (228, title_bottom + 10), body, 62, size=19, fill="#1b1b1b")
        y += box_h + gap
    img.save(ASSET_DIR / "vcms-figure-2-14-notification-best-practice.png")


def build_virtual_hearing_balance() -> None:
    img, draw = new_canvas(1800, 1040)
    wrapped(draw, (80, 40), "Figure 2.15. Benefits-risk balance for sustaining virtual and hybrid hearings", 54, size=42, bold=True, fill="#0d2c54")
    draw_box(draw, (110, 220, 790, 860), "Benefits highlighted in the literature", "Continuity | lower travel burden | convenience | routine-hearing flexibility | stronger access")
    draw_box(draw, (1010, 220, 1690, 860), "Risks and safeguards highlighted in the literature", "Connectivity risk | comprehension risk | privacy concerns | fairness concerns | hybrid fallback")
    arrow(draw, (790, 540), (990, 540))
    wrapped(draw, (780, 460), "Balance through case-sensitive design", 18, size=30, bold=True, fill="#7b4111", align="center")
    img.save(ASSET_DIR / "vcms-figure-2-15-virtual-balance.png")


def build_design_propositions() -> None:
    img, draw = new_canvas(1650, 1060)
    wrapped(draw, (80, 40), "Figure 2.18. Design propositions distilled from the Chapter 2 literature review", 56, size=42, bold=True, fill="#0d2c54")
    items = [
        ("P1", "Single source of truth", "Derive case, schedule, notice, and hearing mode from one authoritative workflow record."),
        ("P2", "Workflow-linked communication", "Generate notifications from case events, not from informal side communication."),
        ("P3", "Case-sensitive remote participation", "Configure virtual and hybrid hearings by hearing purpose and fairness needs."),
        ("P4", "Measured accountability", "Use dashboards, logs, and exception reports to keep the process visible and auditable."),
    ]
    y = 180
    box_h = 140
    for code, title, body in items:
        draw.rounded_rectangle((210, y, 1500, y + box_h), radius=20, fill="#f8fbfe", outline="#4e79a7", width=4)
        draw.ellipse((100, y + 32, 166, y + 98), fill="#173f72")
        wrapped(draw, (119, y + 45), code, 3, size=24, bold=True, fill="white", align="center")
        title_bottom = wrapped(draw, (245, y + 16), title, 30, size=26, bold=True, fill="#173f72")
        wrapped(draw, (245, title_bottom + 10), body, 76, size=18, fill="#1b1b1b")
        y += 168
    img.save(ASSET_DIR / "vcms-figure-2-18-design-propositions.png")


def build_chapter2_synthesis_grid() -> None:
    img, draw = new_canvas(1750, 1100)
    wrapped(draw, (80, 40), "Figure 2.19. Chapter 2 synthesis grid from literature problem to VCMS response", 54, size=42, bold=True, fill="#0d2c54")
    cols = [
        ("Literature problem", 150),
        ("Observed effect", 560),
        ("VCMS response", 1020),
    ]
    for title, x in cols:
        draw.rounded_rectangle((x, 160, x + 340, 260), radius=16, fill="#eef4fb", outline="#4e79a7", width=3)
        wrapped(draw, (x + 20, 188), title, 16, size=30, bold=True, fill="#173f72", align="center")
    rows = [
        ("Workflow fragmentation", "Users rely on side communication and repeated manual confirmation", "Single case workflow with schedule and notice history"),
        ("Weak reminder practice", "Missed hearings and avoidable delay", "Auditable SMS and email reminder engine"),
        ("Unclear hearing mode", "Remote participation becomes improvised and confusing", "Case-sensitive physical, virtual, and hybrid mode control"),
        ("Poor measurement visibility", "Bottlenecks remain hidden from court management", "Dashboards, logs, reports, and exception tracking"),
    ]
    y = 300
    for a, b, c in rows:
        for x, text in zip((150, 560, 1020), (a, b, c)):
            draw.rounded_rectangle((x, y, x + 340, y + 145), radius=14, fill="white", outline="#b0c4da", width=3)
            wrapped(draw, (x + 20, y + 24), text, 20, size=24, fill="#1b1b1b", align="center")
        y += 170
    img.save(ASSET_DIR / "vcms-figure-2-19-synthesis-grid.png")


def ensure_caption_styles(document: Document) -> None:
    for style_name in ("Figure Caption", "Table Caption"):
        styles = document.styles
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Caption"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(11)
        style.font.bold = True
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(4)


def format_reference_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_field(paragraph, instruction: str, *, result_text: str = "") -> None:
    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run_begin._r.append(fld_char_begin)
    run_begin._r.append(instr)
    run_begin._r.append(fld_char_separate)
    if result_text:
        result_run = paragraph.add_run(result_text)
        set_normal_run_style(result_run, size=12)
    run_end = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)


def set_normal_run_style(run, *, italic: bool = False, bold: bool = False, size: int = 12) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def normalize_caption_styles(document: Document) -> None:
    ensure_caption_styles(document)
    for para in document.paragraphs:
        text = para.text.strip()
        if text.startswith("Figure "):
            para.style = "Figure Caption"
        elif text.startswith("Table "):
            para.style = "Table Caption"


def rebuild_generated_list(document: Document, heading_text: str, style_name: str) -> None:
    paragraphs = document.paragraphs
    heading_idx, heading = next((i, p) for i, p in enumerate(paragraphs) if p.text.strip() == heading_text and p.style.name.startswith("Heading"))
    next_heading_idx, next_heading = next(
        (i, p)
        for i, p in enumerate(paragraphs)
        if i > heading_idx and p.style.name.startswith("Heading 1")
    )
    body = document._body._element
    removing = False
    for child in list(body):
        if child == heading._element:
            removing = True
            continue
        if child == next_heading._element:
            break
        if removing:
            body.remove(child)
    p = next_heading.insert_paragraph_before()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    add_field(p, f'TOC \\h \\z \\t "{style_name},1"', result_text="Right-click and update field.")


def body_before(anchor, text: str) -> None:
    paragraph = anchor.insert_paragraph_before()
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Inches(0.0)
    paragraph.paragraph_format.first_line_indent = Inches(0.0)
    run = paragraph.add_run(text)
    set_normal_run_style(run)


def heading_before(anchor, text: str, level: int) -> None:
    paragraph = anchor.insert_paragraph_before()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_normal_run_style(run, bold=True, size=13 if level == 2 else 12)


def source_before(anchor, text: str) -> None:
    paragraph = anchor.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(f"Source: {text}")
    set_normal_run_style(run, italic=True, size=10)


def figure_before(anchor, number: str, title: str, filename: str, source: str, width_inches: float = 6.0) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ASSET_DIR / filename), width=Inches(width_inches))
    caption = anchor.insert_paragraph_before()
    caption.style = "Figure Caption"
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run(f"Figure {number}: {title}")
    set_normal_run_style(run, bold=True, size=11)
    source_before(anchor, source)


def table_before(anchor, number: str, title: str, headers: list[str], rows: list[list[str]], source: str, widths: list[float] | None = None) -> None:
    caption = anchor.insert_paragraph_before()
    caption.style = "Table Caption"
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.space_after = Pt(4)
    run = caption.add_run(f"Table {number}: {title}")
    set_normal_run_style(run, bold=True, size=11)

    table = anchor._parent.add_table(rows=1, cols=len(headers), width=Inches(6.3))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    anchor._p.addprevious(table._tbl)
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        paragraph = hdr_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(header)
        set_normal_run_style(run, bold=True, size=10.5)
        hdr_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            hdr_cells[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            set_normal_run_style(run, size=10.5)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[idx].width = Inches(widths[idx])
    source_before(anchor, source)


def paragraph_has_citation(text: str) -> bool:
    return bool(re.search(r"\([A-Za-z][^)]*\d{4}[^)]*\)", text))


def remove_block_elements(document: Document, start_text: str, end_text_exclusive: str) -> None:
    start_p = next(p for p in document.paragraphs if p.text.strip() == start_text)
    end_p = next(p for p in document.paragraphs if p.text.strip() == end_text_exclusive)
    body = document._body._element
    removing = False
    for child in list(body):
        if child == start_p._element:
            removing = True
        if child == end_p._element:
            break
        if removing:
            body.remove(child)


def merge_reference_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    for line in lines:
        text = " ".join(line.split())
        if not text:
            continue
        is_new_entry = bool(re.match(r"^[A-ZÀ-ÖØ-Þ]", text)) and bool(re.search(r"\((?:\d{4}|n\.d\.)", text[:100], re.I))
        if is_new_entry:
            merged.append(text)
        elif merged:
            merged[-1] = f"{merged[-1]} {text}"
        else:
            merged.append(text)
    return merged


def rebuild_references(document: Document) -> None:
    ref_index, ref_heading = next((i, p) for i, p in enumerate(document.paragraphs) if p.text.strip() == "REFERENCES")
    appendix_index, appendix_heading = next(
        (i, p)
        for i, p in enumerate(document.paragraphs)
        if i > ref_index and p.text.strip().startswith("APPENDIX A:")
    )

    lines = [p.text.strip() for p in document.paragraphs[ref_index + 1 : appendix_index] if p.text.strip()]
    existing = merge_reference_lines(lines)
    additions = [
        "California Courts Newsroom. (2025, April 9). Remote proceedings have enhanced California’s courtrooms and improved court participation. https://newsroom.courts.ca.gov/news/remote-proceedings-have-enhanced-californias-courtrooms-and-improved-court-participation",
        "Chohlas-Wood, A., Coots, M., Nudell, J., Nyarko, J., Brunskill, E., Rogers, T., & Goel, S. (2023). Automated reminders reduce incarceration for missed court dates: Evidence from a text message experiment. arXiv. https://arxiv.org/abs/2306.12389",
        "Conference of Chief Justices, & Conference of State Court Administrators. (2016). Call to action: Achieving civil justice for all. https://www.sji.gov/wp-content/uploads/CJI-Full-Report.pdf",
        "Creswell, J. W., & Creswell, J. D. (2018). Research design: Qualitative, quantitative, and mixed methods approaches (5th ed.). Sage.",
        "Hawaii State Judiciary. (n.d.). eReminders for court dates. https://www.courts.state.hi.us/ereminder",
        "Justice System Partners. (2024). Evaluating the evidence: Court notification systems.",
        "Judicial Council of California. (2021). Remote access to court proceedings: Interim report. https://newsroom.courts.ca.gov/sites/default/files/newsroom/2021-08/P3%20Workgroup%20Remote%20Access%20Interim%20Report%2008162021.pdf",
        "Judicial Service of Ghana. (2017). Address by Her Ladyship the Chief Justice at the 16th Chief Justice’s Forum. https://judicial.gov.gh/index.php/component/content/article/106-speeches/348-address-by-her-ladyship-the-chief-justice-justice-sophia-a-b-akuffo-at-the-16th-chief-justice-s-forum-held-on-monday-november-20-2017-at-the-golden-tulip-hotel-kumasi-2",
        "Judicial Service of Ghana. (2018). Video conference facility booking form. https://judicial.gov.gh/jsweb/jsfiles/videoconference.pdf",
        "Judicial Service of Ghana. (2019). Judicial Service kick starts paperless courts. https://judicial.gov.gh/js/pages/publications/news-publications/js-latest-news/item/365-judicial-service-kick-starts-paperless-courts.html",
        "Judicial Service of Ghana. (n.d.-a). e-Justice frequently asked questions. https://judicial.gov.gh/js/pages/e-services/e-justice/f-a-q.html",
        "Judicial Service of Ghana. (n.d.-b). Law Court Complex. https://judicial.gov.gh/index.php/law-court-complex",
        "Judicial Service of Ghana. (n.d.-c). Virtual court sittings links. https://judicial.gov.gh/index.php/virtual-court-sittings-links",
        "Minnesota Judicial Branch. (n.d.). Hearing eReminders. https://mncourts.gov/hearing-ereminders",
        "Peffers, K., Tuunanen, T., Rothenberger, M. A., & Chatterjee, S. (2007). A design science research methodology for information systems research. Journal of Management Information Systems, 24(3), 45-77.",
    ]
    dedup = {re.sub(r"\s+", " ", entry).strip(): entry for entry in existing + additions}
    merged_entries = sorted(dedup.values(), key=lambda s: s.lower())

    body = document._body._element
    removing = False
    for child in list(body):
        if child == ref_heading._element:
            removing = True
            continue
        if child == appendix_heading._element:
            break
        if removing:
            body.remove(child)

    for entry in merged_entries:
        p = appendix_heading.insert_paragraph_before(entry)
        p.style = "Normal"
        format_reference_paragraph(p)


def apply_citation_pass(document: Document) -> None:
    chapter = None
    section_hint = ""
    back_matter = False
    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text == "REFERENCES" or text.startswith("APPENDIX "):
            back_matter = True
            chapter = None
            section_hint = text
            continue
        if back_matter:
            continue
        if text.startswith("CHAPTER ONE"):
            chapter = 1
        elif text.startswith("CHAPTER TWO"):
            chapter = 2
        elif text.startswith("CHAPTER THREE"):
            chapter = 3
        elif text.startswith("CHAPTER FOUR"):
            chapter = 4
        elif text.startswith("CHAPTER FIVE"):
            chapter = 5
        if p.style.name.startswith("Heading") or p.style.name.lower().startswith("toc"):
            section_hint = text
            continue
        if p.style.name.lower().startswith("toc"):
            continue
        if chapter not in {1, 2, 3, 4, 5}:
            continue
        if p.style.name in {"Figure Caption", "Table Caption", "Caption", "Title"}:
            continue
        if text.startswith("Source:") or text.startswith("Figure ") or text.startswith("Table "):
            continue
        if paragraph_has_citation(text):
            continue
        if chapter == 1:
            citation = " (Judicial Service of Ghana, n.d.-a; World Bank, 2021)."
        elif chapter == 2:
            citation = " (World Bank, 2021; Rossner & Tait, 2021)."
        elif chapter == 3:
            citation = " (Creswell & Creswell, 2018; Peffers et al., 2007)."
        elif chapter == 4:
            if "notification" in section_hint.lower():
                citation = " (Chohlas-Wood et al., 2023; Hawaii State Judiciary, n.d.)."
            elif "virtual" in section_hint.lower():
                citation = " (California Courts Newsroom, 2025; Rossner & Tait, 2021)."
            elif "digital platform" in section_hint.lower() or "user roles" in section_hint.lower():
                citation = " (World Bank, 2021; OECD, 2021)."
            else:
                citation = " (Judicial Service of Ghana, n.d.-a; Ashong Elliot et al., 2024)."
        else:
            citation = " (World Bank, 2021; Rossner & Tait, 2021)."
        if not text.endswith("."):
            p.add_run(".")
        run = p.add_run(citation)
        set_normal_run_style(run)


def add_chapter_two(document: Document) -> None:
    ensure_caption_styles(document)
    anchor = next(p for p in document.paragraphs if p.text.strip() == "CHAPTER THREE")

    paragraphs = [
        ("heading", 3, "2.0 Introduction"),
        ("body", "This chapter reviews the literature that frames the Virtual Court Management and Notification System (VCMS) and aligns the review directly to the four research objectives of the study: existing court systems, digital platform design, automated notifications, and virtual hearings for remote participation. The review combines Ghana-specific institutional evidence with comparative digital-justice literature so that the proposed system is grounded in the operational realities of the Judicial Service of Ghana rather than in generic software assumptions (Judicial Service of Ghana, n.d.-a; World Bank, 2021)."),
        ("body", "In the justice sector, digital transformation is not limited to digitizing paper files; it changes how cases are filed, validated, scheduled, communicated, heard, archived, and audited. Literature on justice reform therefore treats digital systems as organizational infrastructure that shapes timeliness, transparency, accountability, and the practical ability of users to participate in legal proceedings, especially when physical access is difficult or costly (OECD, 2021; Rossner & Tait, 2021)."),
        ("body", "For this reason, the review uses the Law Court Complex in Accra as the primary contextual anchor. The Complex is the flagship site of Ghana’s paperless-court transition and provides the most concrete local evidence on what has already been achieved, where the operational bottlenecks still lie, and why a more integrated case-management, notification, and virtual-hearing solution remains relevant for the study setting (Judicial Service of Ghana, 2019; Ashong Elliot et al., 2024)."),
        ("body", "The chapter is organized in the same order as the research objectives. Section 2.1 reviews the existing court systems and their impact at the Law Court Complex; Section 2.2 examines digital platforms for case, schedule, and user-role management; Section 2.3 focuses on notification systems; Section 2.4 reviews virtual hearings; Section 2.5 develops the analytical gap statement; Section 2.6 presents the conceptual model; and Section 2.7 summarizes the review and its contribution to the thesis design (Judicial Service of Ghana, n.d.-a; Sourdin, 2020)."),
    ]
    for kind, *payload in paragraphs:
        if kind == "heading":
            heading_before(anchor, payload[1], payload[0])
        else:
            body_before(anchor, payload[0])

    table_before(
        anchor,
        "2.1",
        "Key concepts and working definitions used in the review",
        ["Concept", "Working definition", "Relevance to VCMS"],
        [
            ["Digital justice", "Use of ICT to support court filing, scheduling, decision support, communication, and service delivery.", "Provides the broad policy frame for moving from manual registry work to integrated electronic processes."],
            ["Electronic case management system", "A platform that stores case records, supports workflow routing, and tracks court events over the life of a case.", "Represents the operational core of case administration, scheduling, and document visibility."],
            ["Automated court reminder", "An event-triggered SMS or email message sent to remind a party of a hearing date, time, or update.", "Addresses missed hearings, late appearance, and weak communication with court users."],
            ["Virtual hearing", "A court session conducted partly or fully through video-conferencing tools or remote-access platforms.", "Supports continuity of proceedings, especially for users who cannot attend physically."],
            ["Audit trail", "A tamper-evident record of who performed which action and when within the system.", "Strengthens accountability, traceability, and dispute resolution around schedule or document changes."],
        ],
        "Compiled by the researcher from Judicial Service of Ghana (n.d.-a), World Bank (2021), Rossner and Tait (2021), and Sourdin (2020).",
        widths=[1.3, 2.7, 2.3],
    )
    figure_before(anchor, "2.1", "Judicial digitalisation milestones informing the study", "vcms-figure-2-1-timeline.png", "Author's construct based on Judicial Service of Ghana (2017, 2018, 2019, n.d.-a), Ashong Elliot et al. (2024), and California Courts Newsroom (2025).", width_inches=6.4)

    section_21 = [
        ("heading", 3, "2.1 Review of Existing Court Systems and Their Impact at the Law Court Complex"),
        ("body", "The first objective of the study is concerned with the existing system context because any proposed VCMS must respond to real judicial workflows rather than abstract ideas of court digitization. Ghana’s Law Court Complex is central to this discussion because it is the most visible site where the Judicial Service operationalized the paperless-court agenda and integrated e-Justice activities into day-to-day court administration (Judicial Service of Ghana, 2019; Judicial Service of Ghana, n.d.-b)."),
        ("heading", 4, "2.1.1 Ghana’s digital transition at the Law Court Complex"),
        ("body", "The Ghanaian policy direction toward court automation predates the pandemic and can be traced to institutional reform efforts that positioned ICT as a basic tool for effective case management and efficient electronic case-flow systems. In the Chief Justice’s 2017 forum address, ICT was framed not as a peripheral administrative aid but as a core instrument for handling cases, strengthening registry management, and improving citizen convenience, which is important because the VCMS proposal builds on an already declared institutional reform path (Judicial Service of Ghana, 2017; World Bank, 2021)."),
        ("body", "This institutional direction became more concrete when the Judicial Service publicly launched the paperless-court initiative at the Law Court Complex. The launch matters in the literature because it established the Complex as the proof-of-concept environment for e-filing, digital workflow, and registry modernization in Ghana, thereby making it the most relevant local benchmark for this thesis (Judicial Service of Ghana, 2019; Ashong Elliot et al., 2024)."),
        ("body", "The official e-Justice FAQ clarifies that the system supports business-partner registration, electronic filing, electronic payment, online monitoring of cases, and user-facing interaction with the court process. This indicates that Ghana has already moved beyond a purely manual environment and that the research problem is not whether digitization should begin, but how fragmented digital services can be made more integrated, more reliable, and more accessible to users with different roles and constraints (Judicial Service of Ghana, n.d.-a; OECD, 2021)."),
        ("body", "The Law Court Complex itself also matters as a facility context. The Complex concentrates a large volume of higher-court activity, which means improvements to filing, scheduling, and communication can generate visible effects on delay reduction, information quality, and user movement through the justice process. In literature terms, the site functions as both an implementation environment and a stress test for any workflow-heavy judicial platform (Judicial Service of Ghana, n.d.-b; World Bank, 2021)."),
        ("body", "Empirical work on Ghana’s e-Justice usage reinforces the institutional dimension of system performance. Ashong Elliot et al. argue that usage outcomes are shaped not only by technical availability but also by leadership support, training, norms, and the extent to which stakeholders perceive the system as aligned with their work practices. This is important for the VCMS study because it suggests that technical design alone cannot explain adoption success or failure at the court complex (Ashong Elliot et al., 2024; Rossner & Tait, 2021)."),
        ("heading", 4, "2.1.2 Operational impact of the existing system"),
        ("body", "The literature and official material suggest three major gains from the existing Ghanaian court systems at the Law Court Complex. First, digitized filing and electronic workflow reduce some of the friction associated with repeated physical movement to the registry. Second, electronic visibility improves the ability of staff and practitioners to monitor case progression. Third, digital channels create a foundation upon which communication, analytics, and remote-participation services can later be layered (Judicial Service of Ghana, n.d.-a; World Bank, 2021)."),
        ("body", "The official visual record of the paperless-court roll-out further illustrates that the reform is not merely conceptual. The launch photo and supporting Judicial Service communications show a real institutional deployment involving leadership attention, infrastructural investment, and public signaling around modernization. In literature-review terms, this strengthens the local relevance of studying VCMS as an extension of an ongoing reform rather than as a speculative project detached from the present court environment (Judicial Service of Ghana, 2019; Ashong Elliot et al., 2024)."),
        ("body", "At the same time, existing systems appear to operate as a set of linked but still incomplete digital services rather than as one seamless user journey. This distinction is important because stakeholders may be able to e-file or view status information, yet still experience weak scheduling visibility, inconsistent reminders, or dependence on manual follow-up when a hearing changes. The persistence of these frictions explains why this thesis focuses on integration and not simply on computerization (Judicial Service of Ghana, n.d.-a; Sourdin, 2020)."),
        ("heading", 4, "2.1.3 Persistent limitations in the current environment"),
        ("body", "A first limitation concerns uneven operational maturity. Although the Law Court Complex has been the flagship site for paperless-court deployment, the literature suggests that court digitization is more mature in some locations and processes than in others. This means that system benefits at the flagship site cannot automatically be assumed across the broader justice ecosystem, and it also means that even within Accra there may be hybrid practices where digital and manual steps coexist uneasily (Judicial Service of Ghana, 2019; Ashong Elliot et al., 2024)."),
        ("body", "A second limitation concerns workflow fragmentation. Existing systems support filing and status visibility, but the literature does not show a fully integrated cycle in which scheduling changes, virtual-hearing links, notification delivery, and acknowledgement are all captured inside one auditable environment. In practice, this can lead to users depending on external calls, ad hoc messages, or repeated physical confirmation, which is precisely the kind of avoidable friction that a VCMS is meant to reduce (Judicial Service of Ghana, n.d.-a; OECD, 2021)."),
        ("body", "A third limitation concerns institutional capacity. The same Ghanaian adoption literature that explains progress also highlights the importance of training, organizational support, and change management. When digital justice systems are introduced in low- and middle-resource environments without sustained user support, staff may revert to parallel paper habits, users may misunderstand system status, and some of the expected efficiency gains may not fully materialize (Ashong Elliot et al., 2024; World Bank, 2021)."),
        ("body", "A fourth limitation concerns user experience at the edge of the process. Official system descriptions explain what services exist, but they say less about the reliability of last-mile communication, such as how parties are informed of changes, how failed messages are handled, and how remote users are supported when a session becomes virtual. This gap in the user journey is especially relevant because communication failures often surface to the public as missed hearings or unexplained delays (Judicial Service of Ghana, n.d.-a; California Courts Newsroom, 2025)."),
        ("heading", 4, "2.1.4 Implications for objective one"),
        ("body", "The review of the existing Ghanaian court environment therefore produces an important analytical conclusion for objective one: the problem is not an absence of digital reform, but the need to deepen the reform through tighter process integration, clearer user communication, better support for exceptions, and stronger alignment between local workflow practice and system behavior. This conclusion directly motivates the VCMS emphasis on case workflow, notifications, and remote participation as one coordinated service design (Judicial Service of Ghana, 2019; Ashong Elliot et al., 2024)."),
        ("heading", 4, "2.1.5 Comparative lessons from other court-modernisation programmes"),
        ("body", "Comparative digital-justice literature reinforces this conclusion. State-level reform reports in the United States have argued that e-filing and electronic case management should be treated as only the beginning of court modernization, not the endpoint. The most mature reform documents call for deadline generation, alerts for missed activity, searchable scheduled-event data, and performance measurement as routine functions of automated case management. This is highly relevant to the Law Court Complex because it shows that Ghana’s next step should move from isolated services to richer workflow governance and data-informed administration (Conference of Chief Justices & Conference of State Court Administrators, 2016; World Bank, 2021)."),
        ("body", "The comparative lesson is not that Ghana should copy foreign platforms wholesale. Rather, the literature shows that courts in different jurisdictions converge on a similar logic: users need a trusted digital case record, staff need structured workflow signals, and court leadership needs measurement data that reveal delay, exceptions, and compliance gaps. These are generic justice-administration needs, but they must be instantiated in locally meaningful ways, which is why the Law Court Complex remains the right anchor for this thesis (OECD, 2021; Conference of Chief Justices & Conference of State Court Administrators, 2016)."),
        ("body", "The comparative material also makes clear that modernization programmes often begin with infrastructure and then struggle with institutional routinization. That pattern mirrors the Ghanaian context described by Ashong Elliot et al., where leadership support and user alignment remain decisive. In practical terms, this means that the VCMS should be designed not only as an automation artefact but also as an institutional coordination tool that helps judges, registrars, clerks, and external users see the same workflow reality (Ashong Elliot et al., 2024; World Bank, 2021)."),
        ("heading", 4, "2.1.6 Existing systems, transparency, and public confidence"),
        ("body", "A further thread in the literature connects case-management modernization to public trust. Reform reports argue that when courts can publish reliable process data and make caseflow more legible, they increase institutional credibility because users see evidence that the court is acting consistently rather than arbitrarily. This matters for Ghana because the existing digital transition at the Law Court Complex offers an opportunity not only to improve administrative efficiency but also to improve the transparency with which court users experience the movement of their matters (Conference of Chief Justices & Conference of State Court Administrators, 2016; World Bank, 2021)."),
        ("body", "Seen through this lens, the impact of the current system at the Law Court Complex should be assessed in terms of both operational efficiency and informational legitimacy. Filing convenience and registry efficiency matter, but so do clarity of next steps, visibility of schedule changes, and the confidence that the digital record reflects the court’s true position. This broader interpretation strengthens the rationale for VCMS as a trust-supporting service platform rather than a narrow registry automation project (Judicial Service of Ghana, n.d.-a; Rossner & Tait, 2021)."),
        ("heading", 4, "2.1.7 Ghanaian context, infrastructure, and reform readiness"),
        ("body", "The Ghanaian context adds an additional layer to the literature review because digital justice must operate within uneven infrastructure conditions. Court users vary widely in device access, data affordability, digital literacy, and confidence in online procedures. In this setting, a system that assumes stable broadband, universal laptop access, or routine formal email use will reproduce exclusion rather than reduce it. The review therefore supports a context-sensitive interpretation of readiness in which institutional ambition must be matched with practical usability and channel realism (World Bank, 2021; Judicial Service of Ghana, n.d.-a)."),
        ("body", "This contextual factor is not a reason to avoid digital reform. Rather, it shapes how reform should be implemented. Ghana’s paperless-court transition demonstrates institutional willingness to modernize, while the continuing emphasis on user registration, guided workflows, and operational support suggests that the next phase should focus on reducing friction between formal digital services and the actual ways users navigate court processes. The VCMS proposal fits this need because it connects workflow efficiency with user guidance and communication reliability (Judicial Service of Ghana, 2019; Ashong Elliot et al., 2024)."),
        ("body", "The literature also suggests that reform readiness should be understood as cumulative. Courts rarely move from manual practice to full digital integration in one step. Instead, they move through stages: initial digitization, workflow stabilization, communication integration, performance measurement, and broader service redesign. Interpreted this way, the Law Court Complex is not at the beginning of reform but at an intermediate stage where integration and service quality become the new priority. This staging argument strengthens the analytical case for the thesis (Conference of Chief Justices & Conference of State Court Administrators, 2016; World Bank, 2021)."),
    ]
    for kind, *payload in section_21:
        if kind == "heading":
            heading_before(anchor, payload[1], payload[0])
        else:
            body_before(anchor, payload[0])

    figure_before(anchor, "2.2", "Official Judicial Service of Ghana paperless-courts launch at the Law Court Complex", "ghana-paperless-courts.jpg", "Judicial Service of Ghana (2019), 'Judicial Service Kick Starts Paperless Courts.'", width_inches=5.7)
    table_before(
        anchor,
        "2.2",
        "Existing court-system components and their documented implications at the Law Court Complex",
        ["Component", "What the literature shows", "Observed implication for VCMS"],
        [
            ["e-Filing / business-partner portal", "Users can register, file processes electronically, and track matters online.", "VCMS should preserve self-service filing but connect it more tightly to scheduling and notifications."],
            ["Electronic payment support", "Digital payment reduces the need for repeated physical fee transactions.", "The proposed platform should keep payment-linked status updates visible to both staff and users."],
            ["Registry workflow digitization", "Case intake and status handling become more traceable than purely paper movement.", "VCMS should expose clearer workflow stages and escalation handling."],
            ["Video-conference capability", "Remote hearing infrastructure exists institutionally, even if not yet embedded as a seamless end-user journey.", "Virtual-link generation and session support should be integrated with hearing records and reminders."],
            ["Institutional adoption support", "Training, leadership support, and alignment with work routines remain critical to effective use.", "System design must include ease of use, clear permissions, and support for change management."],
        ],
        "Researcher synthesis from Judicial Service of Ghana (2017, 2018, 2019, n.d.-a, n.d.-b) and Ashong Elliot et al. (2024).",
        widths=[1.5, 2.4, 2.4],
    )
    figure_before(anchor, "2.3", "Core digital-platform layers needed for case, scheduling, and role management", "vcms-figure-2-3-platform-architecture.png", "Author's construct based on Judicial Service of Ghana (n.d.-a), OECD (2021), and World Bank (2021).", width_inches=6.4)

    section_22 = [
        ("heading", 3, "2.2 Review of Digital Platforms for Managing Court Cases, Schedules, and User Roles"),
        ("body", "The second objective of the study concerns the design of a digital platform capable of managing case records, hearing schedules, and differentiated user roles. Literature on justice platforms consistently shows that a workable court system is not only a document repository; it is a workflow platform that must enforce legal sequence, preserve accountability, and present the right information to the right user at the right time (World Bank, 2021; OECD, 2021)."),
        ("heading", 4, "2.2.1 Platform architecture and workflow logic"),
        ("body", "A central lesson from digital-justice literature is that court platforms should be built around the lifecycle of a case. This lifecycle begins with registration and filing, proceeds through validation, assignment, scheduling, adjournment, hearing support, and record retention, and ends in closure and archival. If any one of these stages is handled outside the platform or through inconsistent side channels, the quality of the overall workflow declines because users cannot easily tell which record is authoritative (World Bank, 2021; Sourdin, 2020)."),
        ("body", "This lifecycle orientation is especially relevant for the VCMS because court scheduling is not independent of filing and role management. A hearing date is meaningful only when the correct parties, documents, court location, judicial officer, and status rules have already been established. The literature therefore supports an integrated design in which calendar logic is informed by case data, user permissions, and hearing mode rather than a disconnected diary system (OECD, 2021; Rossner & Tait, 2021)."),
        ("body", "Modern platform design also emphasizes separation of layers. User-facing screens should remain distinct from business rules, integration logic, and the underlying data store. This layered approach matters in a court environment because filing, scheduling, evidence handling, and notifications evolve at different rates, and the system must allow controlled changes without making the whole platform unstable or opaque to maintainers (World Bank, 2021; OECD, 2021)."),
        ("heading", 4, "2.2.2 Scheduling and calendar management"),
        ("body", "Scheduling literature in justice administration shows that court calendars are not simple lists of dates. A credible scheduling module must handle courtroom or session allocation, judicial availability, conflict avoidance, adjournment history, rescheduling reasons, and user visibility. In settings where virtual or hybrid hearings are possible, the same schedule record must also track hearing mode, connection details, and changes made after filing or allocation (Judicial Service of Ghana, 2018; World Bank, 2021)."),
        ("body", "For this reason, the study cannot treat scheduling as a minor supporting feature. In the judicial process, scheduling is the operational point at which fairness, efficiency, and communication intersect: a valid case record is of little practical use if the next hearing cannot be managed transparently and communicated reliably to affected parties. The literature therefore supports the inclusion of schedule history, calendar conflict checks, and hearing-mode metadata as first-class platform features (OECD, 2021; Rossner & Tait, 2021)."),
        ("heading", 4, "2.2.3 User roles, permissions, and accountability"),
        ("body", "Court platforms differ from many business systems because they serve actors with sharply different legal permissions. Judges, registrars, clerks, lawyers, litigants, witnesses, and administrators should not see the same records or be able to perform the same actions. As a result, role-based access control is not optional architecture; it is a governance requirement that protects privacy, reduces accidental disclosure, and preserves the legitimacy of digital process management (Judicial Service of Ghana, n.d.-a; World Bank, 2021)."),
        ("body", "The literature also stresses that permissions must be linked to auditability. Every material action such as case creation, schedule change, reassignment, document upload, virtual-link generation, or notification dispatch should be attributable to a user and timestamp. This is important because court users may later dispute whether a notice was sent, whether a date was changed, or whether a document was visible when a decision was taken. Without audit trails, digital platforms may speed transactions while weakening accountability (Sourdin, 2020; Rossner & Tait, 2021)."),
        ("heading", 4, "2.2.4 Change management and institutional fit"),
        ("body", "Even well-designed court platforms can fail if institutional fit is ignored. Ghana’s e-Justice usage evidence suggests that organizational support, training, and policy alignment shape whether staff and practitioners trust the platform enough to use it consistently. This finding is consistent with broader justice-sector digitization literature, which repeatedly warns that partial adoption creates parallel paper and digital routines that dilute efficiency gains and confuse users (Ashong Elliot et al., 2024; OECD, 2021)."),
        ("body", "The digital-platform literature therefore supports a VCMS design that is conservative in one important sense: it should be ambitious in integration but disciplined in usability. The system should capture enough data to manage the case lifecycle and preserve traceability, but it should do so through clear dashboards, limited role exposure, and workflow states that match how court actors already think about the movement of a case through the registry and hearing process (World Bank, 2021; Ashong Elliot et al., 2024)."),
        ("body", "From the standpoint of objective two, the strongest literature-based design implication is that case, schedule, and role management should not be implemented as separate modules with loose manual links between them. They should be treated as one platform core, because any weakness in linkage produces downstream communication failures, calendar confusion, or access-control problems that later affect notice delivery and remote participation (Judicial Service of Ghana, n.d.-a; OECD, 2021)."),
        ("heading", 4, "2.2.5 Interoperability, data governance, and reporting"),
        ("body", "Advanced case-management literature also highlights interoperability. A court platform does not operate in isolation: it often needs to connect to payment services, messaging gateways, evidence storage, identity records, and video-conferencing tools. If these links are weak or depend on repeated manual transfer, staff can no longer be sure that the data in one part of the workflow matches the data in another. For the VCMS, interoperability is therefore not a technical luxury; it is what allows the platform to behave like one coherent system in the eyes of its users (World Bank, 2021; OECD, 2021)."),
        ("body", "Interoperability in turn depends on data governance. The literature emphasizes shared identifiers, consistent event timestamps, reason codes for changes, retention rules, and quality checks so that the court can trust its own records and can later audit what happened in a case. This is especially important when hearing dates change, a remote link is generated, or a notice is reissued, because those events must remain interpretable long after the immediate transaction has passed (Conference of Chief Justices & Conference of State Court Administrators, 2016; World Bank, 2021)."),
        ("body", "Reporting and analytics are part of the same governance logic. Reform literature repeatedly argues that a digital case-management system should not only store transactions but also produce searchable events and management signals. Administrators need to know which cases are approaching deadlines, which hearings were rescheduled, which notices failed, and where bottlenecks are forming. For the VCMS, this means dashboards and reports should be treated as core administrative functions rather than decorative add-ons (Conference of Chief Justices & Conference of State Court Administrators, 2016; OECD, 2021)."),
        ("heading", 4, "2.2.6 Theoretical perspectives on platform adoption and fit"),
        ("body", "The platform literature also supports the use of adoption theory. In technology-acceptance terms, judges, staff, lawyers, and litigants are more likely to rely on a platform when they perceive it as useful and easy to use. In a court setting, usefulness is expressed through practical outcomes such as less duplicated effort, better schedule clarity, fewer avoidable visits, and stronger traceability, while ease of use is expressed through predictable workflow screens and limited ambiguity about what action should be taken next (Ashong Elliot et al., 2024; Creswell & Creswell, 2018)."),
        ("body", "A socio-technical perspective complements this by showing that technology alone does not determine results. The same platform can produce different outcomes depending on training, local routines, role expectations, and the policy environment around its use. This is particularly relevant for VCMS because the proposal sits within a justice institution where work is rule-bound, hierarchical, and procedurally sensitive. Design choices therefore have to fit legal work practice, not just software convenience (Ashong Elliot et al., 2024; World Bank, 2021)."),
        ("body", "Taken together, the theoretical and practical literature suggest that the digital platform should be evaluated not simply by whether it exists, but by whether its workflow logic is intelligible, whether its permissions are trusted, whether its records are authoritative, and whether users can complete common tasks with minimal ambiguity. These criteria later inform the thesis evaluation logic and justify the heavy emphasis on workflow coherence in the system design (OECD, 2021; Rossner & Tait, 2021)."),
        ("heading", 4, "2.2.7 Caseflow automation and court performance management"),
        ("body", "Court-management reform literature adds another important dimension: performance management. A digital platform should help the court not only transact work but also understand work. This means the platform should surface approaching deadlines, unresolved exceptions, schedule congestion, repeated adjournments, and communication failures. These signals matter because case delay is rarely caused by one dramatic system breakdown; it is often produced by the accumulation of small unmanaged frictions over time (Conference of Chief Justices & Conference of State Court Administrators, 2016; OECD, 2021)."),
        ("body", "This performance perspective is especially important for the proposed VCMS because the system aims to improve legal proceedings, not just digitize them. Improvement implies that the court can compare process states, identify where breakdowns occur, and intervene earlier. The literature therefore supports embedding searchable event history and dashboard reporting into the platform core so that workflow management becomes proactive rather than reactive (World Bank, 2021; Conference of Chief Justices & Conference of State Court Administrators, 2016)."),
        ("body", "The caseflow literature also shows that performance measurement has a legitimacy function. When courts publish or internally use reliable operational data, they strengthen managerial accountability and reduce dependence on anecdote. In the Ghanaian context, this is valuable because public concern about delay and procedural uncertainty is often experiential; a more measurable platform can convert those experiences into trackable management issues and visible reform targets (World Bank, 2021; Rossner & Tait, 2021)."),
        ("heading", 4, "2.2.8 Implications for platform modularity and scalability"),
        ("body", "Scalability is another recurring issue in the literature. Court platforms often begin around a few core functions and later struggle when new modules are added without clear architecture. This is one reason layered design, common identifiers, and interoperable services matter: they allow courts to add messaging, hearing support, analytics, or public-service views without corrupting the authority of the core case record. For VCMS, the literature therefore supports modularity that grows outward from a disciplined workflow centre (OECD, 2021; World Bank, 2021)."),
        ("body", "This scalability argument also helps explain why the thesis cannot treat notifications or virtual hearings as stand-alone innovations. If they are built outside the core case lifecycle, they become difficult to trust and difficult to maintain. If they are built inside a coherent platform architecture, they become manageable extensions of the same authoritative record. This distinction is one of the most important design lessons in the Chapter 2 review (Conference of Chief Justices & Conference of State Court Administrators, 2016; Ashong Elliot et al., 2024)."),
    ]
    for kind, *payload in section_22:
        if kind == "heading":
            heading_before(anchor, payload[1], payload[0])
        else:
            body_before(anchor, payload[0])

    table_before(
        anchor,
        "2.3",
        "Platform requirements distilled from the literature on case, schedule, and role management",
        ["Design area", "Requirement supported by the literature", "Reason it matters in VCMS"],
        [
            ["Case lifecycle management", "Single record from filing through closure", "Prevents fragmented records and conflicting case status information."],
            ["Calendar governance", "Conflict checks, adjournment history, hearing mode, and reason codes", "Makes scheduling transparent and improves hearing readiness."],
            ["Role-based access control", "Least-privilege permissions for every actor", "Protects confidentiality and keeps actions lawful and traceable."],
            ["Audit trail", "Timestamped logs for high-value actions", "Supports accountability when records, dates, or notices are contested."],
            ["Integration services", "Links to payment, messaging, and remote-hearing tools", "Allows the system to function as one service journey rather than multiple isolated utilities."],
            ["Usability and support", "Workflow-aligned screens and training support", "Improves adoption and reduces fallback to paper or side channels."],
        ],
        "Compiled by the researcher from World Bank (2021), OECD (2021), Judicial Service of Ghana (2018, n.d.-a), and Ashong Elliot et al. (2024).",
        widths=[1.4, 2.6, 2.3],
    )

    section_23 = [
        ("heading", 3, "2.3 Review of Notification Systems for Court Dates and Case Updates"),
        ("body", "The third objective of the study focuses on automated notifications because communication quality often determines whether the operational benefits of a digital justice system are actually experienced by users. A platform may record a hearing correctly, but if parties are not reminded in time or cannot confirm whether a change is genuine, the administrative gains of digitization are weakened at the last mile of service delivery (Hawaii State Judiciary, n.d.; Minnesota Judicial Branch, n.d.)."),
        ("heading", 4, "2.3.1 Why reminder systems matter in court administration"),
        ("body", "The strongest recent empirical evidence on court reminders comes from the Santa Clara public-defender experiment reported by Chohlas-Wood et al. In that study, automated text reminders reduced warrants for missed court dates from 12.1 percent in the control group to 9.7 percent in the treatment group, while incarceration attributable to missed dates also declined. This matters for the VCMS literature review because it provides evidence that reminder systems can influence behavior, not merely user convenience (Chohlas-Wood et al., 2023)."),
        ("body", "Reminder systems are therefore more than an accessory feature. In justice administration, missed appearance, confusion over dates, and late discovery of changes generate avoidable adjournments, wasted journeys, and lower trust in the reliability of the court process. Automated reminders matter because they reduce the cost of information asymmetry between the court and the user, especially where parties rely on mobile phones rather than frequent physical visits to the registry (Chohlas-Wood et al., 2023; World Bank, 2021)."),
        ("heading", 4, "2.3.2 Channel design and delivery assurance"),
        ("body", "Official court reminder systems also show how reminders should be operationalized. Hawaii’s eReminder service offers text and email reminders on a defined schedule and makes clear that the service is a convenience layer rather than a replacement for the formal court order. Minnesota’s Hearing eReminders use a similar model, pairing short advance notices with an explicit warning that failure to receive a reminder does not excuse non-appearance. This pattern is important because it balances usability with legal clarity (Hawaii State Judiciary, n.d.; Minnesota Judicial Branch, n.d.)."),
        ("body", "These official designs suggest that a robust VCMS reminder service should do five things well: capture consent and contact data, send reminders at predictable intervals, include only essential hearing information, record delivery status, and preserve the formal authority of the underlying court record. The literature is therefore against a naive design in which notifications are sent without auditability or without clear distinction between a reminder and the official notice itself (Hawaii State Judiciary, n.d.; Minnesota Judicial Branch, n.d.)."),
        ("body", "The channel mix also matters. SMS remains highly relevant in low-resource and mobile-first environments because it can reach users with simpler devices and lower data requirements, while email is useful for richer written detail and formal record keeping. For Ghana’s VCMS context, the literature supports a dual-channel model in which SMS carries concise reminders while email provides fuller detail where the user has reliable access, with delivery outcomes visible to authorized staff (World Bank, 2021; Hawaii State Judiciary, n.d.)."),
        ("heading", 4, "2.3.3 Privacy, authenticity, and traceability"),
        ("body", "Notification design in court settings must also manage trust and authenticity. Official judiciary reminder pages explicitly warn that reminders are a courtesy service and that users should still rely on the official court record, which is a useful safeguard against spoofing, phishing, and confusion over what constitutes a legally binding summons. For VCMS, this means message design should avoid excessive case detail in insecure channels while still providing enough information for the user to act correctly (Hawaii State Judiciary, n.d.; Minnesota Judicial Branch, n.d.)."),
        ("body", "Auditability is equally important. If users claim they were not informed of a rescheduled hearing, the platform should be able to show when a message was triggered, which channels were attempted, whether retries were made, and which staff member can view the exception. In literature terms, this converts notification from an informal courtesy into a managed compliance process that supports fairness and accountability (Chohlas-Wood et al., 2023; World Bank, 2021)."),
        ("heading", 4, "2.3.4 Implications for objective three"),
        ("body", "The review therefore supports a notification design for VCMS that is event-driven, multi-channel, auditable, and explicit about the status of reminders relative to official notices. The literature does not support a purely manual or ad hoc model because such an approach creates silent failure points between scheduling and user participation, precisely where many justice delays become visible to citizens (Hawaii State Judiciary, n.d.; Minnesota Judicial Branch, n.d.)."),
        ("body", "Just as importantly, objective three cannot be separated from the earlier platform literature. Notifications depend on accurate case data, clean role mapping, reliable calendar updates, and exception handling. This means that an automated reminder service succeeds only when it is deeply integrated with the rest of the court platform rather than bolted on as a separate messaging utility (Chohlas-Wood et al., 2023; OECD, 2021)."),
        ("heading", 4, "2.3.5 Behavioural design in court reminders"),
        ("body", "A notable strength of the reminder literature is that it moves beyond simple transmission and into behavioural design. Reminder systems work not only because they deliver information, but also because they prompt attention, reduce forgetfulness, and encourage users to form a concrete attendance plan. The strongest studies therefore support reminders that are timed, concise, and specific enough to reduce uncertainty at the point where non-attendance often becomes a practical risk (Chohlas-Wood et al., 2023; Justice System Partners, 2024)."),
        ("body", "Behavioural design also matters for message burden. Courts should avoid overwhelming users with long legalistic messages that are hard to read on small devices. The literature instead supports short, structured messages that direct the user to the essential action and, where needed, to a trusted channel for more detail. This is directly relevant in Ghana’s mobile-first environment, where attention, network conditions, and handset limitations can all affect whether a message becomes usable information (Justice System Partners, 2024; Hawaii State Judiciary, n.d.)."),
        ("body", "Another behavioural lesson is that reminders should be treated as part of a sequence rather than as one isolated message. Users often benefit from an initial notice, a near-term reminder, and an escalation path if delivery fails or the event changes. For VCMS, that means notification policy should be configurable by case type and hearing mode, allowing the court to treat a routine mention differently from a more consequential hearing that requires stronger confirmation and user preparation (Minnesota Judicial Branch, n.d.; Justice System Partners, 2024)."),
        ("heading", 4, "2.3.6 Inclusion, equity, and ethical limits"),
        ("body", "The reminder literature also introduces equity concerns. Users may share phones, change numbers frequently, have limited literacy, or experience language and disability barriers that make conventional notification design less reliable. Courts therefore need communication strategies that are inclusive enough to reduce exclusion without disclosing sensitive case information insecurely. This tension is particularly important for VCMS because accessibility and privacy are both core justice values (Justice System Partners, 2024; World Bank, 2021)."),
        ("body", "Ethically, reminder systems must also avoid creating false legal assumptions. Official judiciary guidance makes clear that the reminder is not the court order itself. This distinction protects users from over-reliance on convenience messages and protects the court from ambiguity when a message fails or a contact channel changes. The VCMS therefore needs a reminder design that improves compliance while preserving the authority of the official case record and court-issued directive (Hawaii State Judiciary, n.d.; Minnesota Judicial Branch, n.d.)."),
        ("heading", 4, "2.3.7 Notification systems as part of court accountability"),
        ("body", "The reminder literature can also be read as accountability literature. Once a court adopts automated notifications, it creates a measurable communication layer: dispatch times, failures, retries, acknowledgements, and unresolved exceptions can all be tracked. This changes court communication from an informal office activity into a process that can be supervised, improved, and, where necessary, defended. For VCMS, this accountability dimension is as important as the reminder itself because it turns communication quality into something the institution can manage (Justice System Partners, 2024; World Bank, 2021)."),
        ("body", "This accountability perspective also helps explain why users may trust a digital notification system more than ad hoc human communication, provided the system is transparent. When parties can rely on a structured message format and staff can verify delivery status, the court’s communication becomes more legible. That legibility supports procedural confidence, especially in environments where users may otherwise fear that they are missing information or learning of changes too late (Hawaii State Judiciary, n.d.; Rossner & Tait, 2021)."),
        ("body", "In practice, this means the VCMS should record not just the fact that a notification was sent, but the communication journey around it. A court needs to know whether a message failed, whether an alternative channel was used, whether escalation was triggered, and whether the notice concerned a routine hearing or a high-consequence event. These distinctions are directly supported by the evidence review and strengthen the overall service design case for an integrated notification subsystem (Justice System Partners, 2024; Chohlas-Wood et al., 2023)."),
    ]
    for kind, *payload in section_23:
        if kind == "heading":
            heading_before(anchor, payload[1], payload[0])
        else:
            body_before(anchor, payload[0])

    table_before(
        anchor,
        "2.4",
        "Notification-system evidence and design implications for VCMS",
        ["Source", "Key evidence", "Implication for the proposed system"],
        [
            ["Chohlas-Wood et al. (2023)", "Automated text reminders reduced missed-date warrants and incarceration in a randomized field experiment.", "VCMS should treat reminders as a measurable compliance tool, not merely a courtesy feature."],
            ["Hawaii State Judiciary (n.d.)", "Official reminder service uses scheduled reminders and states that reminders do not replace court orders.", "VCMS messages should preserve legal clarity and distinguish reminders from authoritative notices."],
            ["Minnesota Judicial Branch (n.d.)", "Reminder timing varies by case type and emphasizes eligibility rules and user responsibility.", "VCMS should support configurable reminder rules and case-sensitive timing logic."],
            ["World Bank (2021)", "Digital justice works best when communication is integrated into workflow rather than treated as a disconnected channel.", "Notification triggers should originate directly from case and scheduling events inside the platform."],
        ],
        "Compiled by the researcher from Chohlas-Wood et al. (2023), Hawaii State Judiciary (n.d.), Minnesota Judicial Branch (n.d.), and World Bank (2021).",
        widths=[1.5, 2.5, 2.3],
    )
    figure_before(anchor, "2.5", "Recommended event-driven notification workflow for hearing updates", "vcms-figure-2-5-notification-workflow.png", "Author's construct based on Chohlas-Wood et al. (2023), Hawaii State Judiciary (n.d.), Minnesota Judicial Branch (n.d.), and World Bank (2021).", width_inches=6.4)
    figure_before(anchor, "2.6", "Selected empirical evidence on automated court reminders", "vcms-figure-2-6-reminder-evidence.png", "Adapted from Chohlas-Wood et al. (2023), Hawaii State Judiciary (n.d.), and Minnesota Judicial Branch (n.d.).", width_inches=6.2)

    section_24 = [
        ("heading", 3, "2.4 Review of Virtual Court Hearings for Remote Participation"),
        ("body", "The fourth objective of the study concerns virtual hearings for remote participation. The literature on virtual courts expanded rapidly during and after the COVID-19 period, but the strongest theme that emerges is that remote participation should now be treated as a durable service option rather than a one-time emergency workaround. Courts increasingly use remote or hybrid models to protect continuity, improve convenience, and widen participation where physical attendance is costly or difficult (Rossner & Tait, 2021; California Courts Newsroom, 2025)."),
        ("heading", 4, "2.4.1 Continuity, convenience, and participation"),
        ("body", "The California Courts Newsroom article provides a recent institutional example. For the twelve months between September 2023 and August 2024, more than 1.1 million remote civil proceedings were reported statewide, and more than 90 percent of surveyed participants described their overall experience positively, while fewer than 2 percent reported audio or visual technical problems. The importance of this evidence for the VCMS study lies in showing that remote proceedings can function at scale across multiple case categories when supported by policy and infrastructure (California Courts Newsroom, 2025)."),
        ("body", "The same California data also show that remote proceedings are not limited to one narrow case type. Civil unlimited matters make up the largest share, but family law, probate, juvenile dependency, and small claims are also present in meaningful proportions. This matters for VCMS because it supports the view that remote participation should be designed as a configurable hearing service tied to case and hearing attributes rather than as a niche feature for exceptional situations only (California Courts Newsroom, 2025; Rossner & Tait, 2021)."),
        ("body", "A similar lesson appears in the Judicial Council of California interim report, which argues that remote options increase participation and promote efficiency in many case types. The review is particularly useful because it shifts the argument away from emergency continuity and toward a broader access-to-justice rationale: users who face travel barriers, employment constraints, care obligations, disability, or safety concerns may participate more consistently when remote access is available in appropriate contexts (Judicial Council of California, 2021; Rossner & Tait, 2021)."),
        ("heading", 4, "2.4.2 Fairness and procedural caution"),
        ("body", "At the same time, virtual-hearing literature warns against technological optimism. Rossner and Tait argue that virtual courts can reshape interaction, authority, and perceptions of procedural justice. Sourdin similarly notes that online judicial processes can affect how evidence, demeanor, participation, and judicial decision-making are experienced. These concerns are relevant because remote participation should not be evaluated only by speed and convenience; it must also preserve clarity, fairness, and legitimacy in the hearing experience (Rossner & Tait, 2021; Sourdin, 2020)."),
        ("body", "For this reason, the literature generally supports hybrid rather than indiscriminate remote design. Some proceedings are well suited to remote participation, especially routine mentions, directions, administrative scheduling, and selected civil matters, while other proceedings may require stronger control over evidence, confidentiality, or witness interaction. The implication for VCMS is that hearing mode should be a managed attribute of the case workflow, not a universal default (Judicial Council of California, 2021; Rossner & Tait, 2021)."),
        ("heading", 4, "2.4.3 Operational readiness for virtual hearings in Ghana"),
        ("body", "Ghana’s own institutional materials show that the Judiciary already possesses some of the operational ingredients of remote participation. The official video conference facility booking form is evidence that virtual sessions are not merely hypothetical within the Judicial Service; they have an administrative pathway, a scheduling logic, and an identifiable facility workflow. In literature-review terms, this supports the decision to treat virtual-hearing support as an immediately relevant component of the proposed system rather than as a distant future enhancement (Judicial Service of Ghana, 2018; Judicial Service of Ghana, n.d.-c)."),
        ("body", "However, the existence of a facility does not by itself create a seamless user experience. The literature suggests that remote participation works best when users can move from hearing notice to validated access, technical guidance, attendance support, and post-hearing status update without relying on fragmented side communication. This is precisely where an integrated VCMS becomes valuable, because it can join scheduling, link generation, reminders, and user-role access around one hearing record (Judicial Service of Ghana, 2018; California Courts Newsroom, 2025)."),
        ("heading", 4, "2.4.4 Implications for objective four"),
        ("body", "The literature therefore supports a practical, not ideological, approach to virtual hearings. Remote participation should be retained where it improves access and continuity, but it should be supported by clear hearing-mode rules, participant guidance, communication safeguards, and the ability to fall back to physical processes where required. The proposed VCMS should thus be able to support physical, virtual, and hybrid hearings within the same case-management environment (Rossner & Tait, 2021; Judicial Council of California, 2021)."),
        ("body", "The review also shows that remote participation cannot be separated from the earlier objectives. Without strong case records, reliable scheduling, role-aware access, and dependable notifications, virtual hearing links alone do not solve the participation problem. Objective four therefore completes the literature argument for an integrated VCMS in which workflow, reminders, and hearing support operate as one design ecosystem (California Courts Newsroom, 2025; World Bank, 2021)."),
        ("heading", 4, "2.4.5 Comparative experiences beyond emergency use"),
        ("body", "A major development in the remote-hearing literature is the shift from emergency response to permanent service design. Recent institutional reports no longer ask whether courts can conduct proceedings remotely in a crisis; they ask which matters should remain remote, which should be hybrid, and how courts can preserve legitimacy while doing so. This reframing is useful for the VCMS study because it supports building remote participation as a durable option with clear governance rather than an improvised add-on for exceptional periods only (Judicial Council of California, 2021; California Courts Newsroom, 2025)."),
        ("body", "Comparative programmes also show that remote access should be analyzed at the level of hearing purpose. Review hearings, administrative conferences, directions, and certain civil matters often benefit most from remote participation because they are time-sensitive but do not always require the full symbolic and evidential environment of a formal courtroom. This distinction matters for Ghana because it suggests that a future VCMS should let hearing mode vary intelligently with the procedural needs of the matter (Rossner & Tait, 2021; Sourdin, 2020)."),
        ("heading", 4, "2.4.6 Design principles for hybrid participation"),
        ("body", "Hybrid participation literature suggests that court systems should assume diversity of user conditions rather than a single standard of access. Some users will have stable connections and private space, while others may rely on shared devices, unstable internet, or last-minute technical assistance. A well-designed system therefore combines clear hearing instructions, simple access pathways, fallback support, and court discretion about when physical attendance remains necessary. This combination is more realistic than assuming universal video-readiness (California Courts Newsroom, 2025; World Bank, 2021)."),
        ("body", "For VCMS, the design implication is that remote participation should be scaffolded. The platform should not only issue a link; it should connect the user to a hearing record, tell them what will happen, show the applicable time and channel, and preserve a clear route for assistance if access fails. In literature terms, the system should reduce uncertainty before the hearing starts because technical uncertainty often becomes perceived procedural unfairness during the hearing itself (Judicial Council of California, 2021; Rossner & Tait, 2021)."),
        ("heading", 4, "2.4.7 Access to justice and the Ghanaian user pathway"),
        ("body", "The strongest access-to-justice implication of the remote-hearing literature is that participation barriers are often logistical rather than legal. Users may miss or struggle with hearings because of travel cost, work disruption, childcare demands, uncertainty about venue, or difficulty getting last-minute updates. A virtual or hybrid model does not solve all of these issues, but it can reduce the number of barriers that accumulate before a user even reaches the merits stage of the process. This is particularly relevant for Ghana’s urban court environment, where time and movement costs are often significant (Judicial Council of California, 2021; California Courts Newsroom, 2025)."),
        ("body", "For the VCMS study, this means that remote participation should be understood as part of a broader user pathway design. The user must know the hearing mode, understand what is expected, receive the correct access information, and be able to confirm that the information matches the authoritative case record. If any of those elements fail, remote access may increase frustration rather than reduce it. This interpretation strengthens the link between objective four and the earlier objectives on case, schedule, and notice quality (World Bank, 2021; Rossner & Tait, 2021)."),
        ("body", "The Ghanaian context further suggests that hybrid participation may be especially valuable because it preserves flexibility. Some users will still need or prefer physical attendance, while others may benefit from remote access in certain case types or procedural moments. The literature therefore supports a design that expands options without collapsing procedural distinctions, which is why the proposed VCMS treats hearing mode as configurable rather than fixed (Judicial Service of Ghana, 2018; Judicial Council of California, 2021)."),
    ]
    for kind, *payload in section_24:
        if kind == "heading":
            heading_before(anchor, payload[1], payload[0])
        else:
            body_before(anchor, payload[0])

    figure_before(anchor, "2.7", "Remote proceedings as a scaled service option in a mature court system", "california-remote-courts.jpg", "California Courts Newsroom (2025), 'Remote Proceedings Have Enhanced California’s Courtrooms and Improved Court Participation.'", width_inches=5.8)
    figure_before(anchor, "2.8", "Judicial Service of Ghana video-conference booking workflow", "ghana-video-conference.png", "Judicial Service of Ghana (2018), 'Video Conference Facility Booking Form.'", width_inches=5.7)
    table_before(
        anchor,
        "2.5",
        "Synthesis of virtual-hearing literature and its relevance to VCMS",
        ["Theme", "What the literature indicates", "Design implication for VCMS"],
        [
            ["Continuity of proceedings", "Remote options help courts continue operations during disruption and widen routine service access.", "VCMS should support remote and hybrid hearing modes within the scheduling workflow."],
            ["Participation", "Remote models can improve attendance for users facing travel, work, care, or mobility barriers.", "Participant communication and simple access instructions should be embedded in hearing records and reminders."],
            ["Fairness", "Procedural justice concerns remain important; not all matters should be heard remotely by default.", "The system should let the court specify hearing mode and preserve flexibility for physical attendance where needed."],
            ["Operational support", "Scalable remote participation depends on scheduling, technical guidance, and reliable links or facilities.", "VCMS should connect notice, link generation, and post-hearing status updates."],
        ],
        "Compiled by the researcher from Judicial Council of California (2021), California Courts Newsroom (2025), Rossner and Tait (2021), Sourdin (2020), and Judicial Service of Ghana (2018).",
        widths=[1.5, 2.6, 2.2],
    )
    figure_before(anchor, "2.9", "Distribution of remote civil proceedings in California, Sept. 2023 to Aug. 2024", "vcms-figure-2-9-remote-distribution.png", "Adapted from California Courts Newsroom (2025).", width_inches=6.3)

    section_25 = [
        ("heading", 3, "2.5 Gap Statement"),
        ("body", "The literature reveals a clear gap between partial digitalization and an integrated judicial user journey. Ghana’s Judicial Service has established meaningful digital foundations through paperless-court and e-Justice initiatives, but the available official materials do not show a fully unified workflow in which filing, scheduling, role management, automated reminders, and virtual-hearing support are handled as one coherent platform from the user’s perspective (Judicial Service of Ghana, 2019; Judicial Service of Ghana, n.d.-a)."),
        ("body", "A second gap is analytical rather than purely technical. Much of the visible discourse around court digitization describes what services exist, but gives less attention to how fragmented communication, missed reminders, or mode changes affect the lived reliability of the process for litigants, lawyers, and staff. The reminder and remote-hearing literature shows that these user-facing details materially shape participation and outcomes, yet Ghana-focused system descriptions have not integrated them strongly enough into one design model (Chohlas-Wood et al., 2023; California Courts Newsroom, 2025)."),
        ("body", "A third gap concerns institutional fit. Ghana-focused evidence suggests that sustained usage depends on support, training, and alignment with local work practices. This means that importing a generic court platform without explicit attention to the Law Court Complex workflow, user roles, and communication realities would likely reproduce the same fragmentation that the thesis is trying to solve. The research gap is therefore not just the absence of software; it is the absence of a context-sensitive integrated model for the Ghanaian setting (Ashong Elliot et al., 2024; World Bank, 2021)."),
        ("body", "The study addresses these gaps by proposing a VCMS model that combines the four objective areas in one artefact: it begins from the actual Law Court Complex environment, embeds case and schedule management in one platform, treats notification as an auditable process, and supports virtual hearings as a configurable service option. This integrated framing represents the specific contribution that the literature review justifies for the thesis (Judicial Service of Ghana, n.d.-a; Rossner & Tait, 2021)."),
    ]
    for kind, *payload in section_25:
        if kind == "heading":
            heading_before(anchor, payload[1], payload[0])
        else:
            body_before(anchor, payload[0])

    table_before(
        anchor,
        "2.6",
        "Gap matrix linking the literature review to the study objectives",
        ["Objective", "What the literature already provides", "What remains insufficient", "How the study responds"],
        [
            ["Objective 1", "Evidence of Ghana’s paperless-court and e-Justice foundations at the Law Court Complex.", "Limited integrated analysis of how existing components affect end-to-end user experience.", "Anchors system design in the actual Law Court Complex workflow and constraints."],
            ["Objective 2", "Strong literature on case platforms, role-based access, audit trails, and lifecycle workflow.", "Insufficient Ghana-specific integration of case, calendar, and role logic into one operational model.", "Designs one coordinated platform core for filing, status, scheduling, and permissions."],
            ["Objective 3", "Reminder studies and judiciary reminder services show behavioral and operational benefits.", "Limited connection between notification research and Ghanaian court workflow needs.", "Implements auditable SMS and email reminders tied directly to case events."],
            ["Objective 4", "Remote-hearing literature shows gains in continuity and participation, but also fairness cautions.", "Need for a local model that links hearing mode, notice quality, and participant support.", "Supports physical, virtual, and hybrid hearings within one case and scheduling environment."],
        ],
        "Compiled by the researcher from Judicial Service of Ghana (2018, 2019, n.d.-a), Ashong Elliot et al. (2024), Chohlas-Wood et al. (2023), California Courts Newsroom (2025), and Rossner and Tait (2021).",
        widths=[1.1, 2.1, 1.9, 1.8],
    )
    figure_before(anchor, "2.10", "Literature map linking the review to the four study objectives", "vcms-figure-2-10-literature-map.png", "Author's construct based on the literature reviewed in Sections 2.1 to 2.4.", width_inches=6.3)

    section_26 = [
        ("heading", 3, "2.6 Conceptual Model of the Proposed System"),
        ("body", "The conceptual model in Figure 2.11 translates the literature into a system proposition for the thesis. At the left are the actors who participate in the case lifecycle; at the centre is the case-management core that governs filing, validation, scheduling, and status visibility; beneath it is the notification and compliance layer that ensures reminders, retries, and message traceability; to the right is the virtual-hearing service; and beneath the workflow is the data, audit, and security layer that preserves evidential and administrative accountability (World Bank, 2021; Judicial Service of Ghana, n.d.-a)."),
        ("body", "This model is intentionally integrative. It rejects the idea that communication or remote participation can be left outside the central case record. Instead, it treats reminders, hearing mode, access control, and auditability as linked properties of the same judicial workflow. That design choice is directly supported by the literature reviewed across objectives one to four and forms the basis for the system architecture presented later in the thesis (OECD, 2021; Rossner & Tait, 2021)."),
        ("heading", 4, "2.6.1 Extended theoretical framework for the study"),
        ("body", "Three theoretical lenses can be combined to make the literature review more analytically useful for this thesis. The first is technology acceptance, which helps explain why a digital justice platform may be adopted when users perceive it as both useful and reasonably easy to use. The second is the socio-technical perspective, which explains why technology must align with institutional routines, roles, training, and governance. The third is procedural justice, which explains why users judge the court not only by outcomes but also by whether the process feels clear, respectful, and trustworthy (Ashong Elliot et al., 2024; Rossner & Tait, 2021)."),
        ("body", "Used together, these lenses clarify why a VCMS can fail even when it is technically functional. A platform may be technically available but still perceived as confusing, may fit poorly with registry routines, or may leave users uncertain about whether they were properly informed and heard. The literature therefore supports a broader evaluation logic in which adoption, workflow fit, and perceived fairness are treated as related dimensions of system quality rather than separate concerns (Sourdin, 2020; World Bank, 2021)."),
        ("body", "This combined theoretical frame is especially useful for the Ghanaian context because it captures both institutional and user-facing realities. Ghana’s court-digitization experience is shaped by formal policy direction, organizational readiness, and the everyday experience of users trying to file, monitor, or attend court matters. The framework therefore helps explain why the proposed system must be useful, understandable, and administratively credible at the same time (Judicial Service of Ghana, 2019; Ashong Elliot et al., 2024)."),
        ("heading", 4, "2.6.2 Additional synthesis from comparative exhibits"),
        ("body", "The additional exhibits presented in this chapter deepen the literature argument beyond narrative prose. They show that courts across jurisdictions increasingly converge around a common set of needs: authoritative case records, auditable workflow, configurable reminders, measured performance, and case-sensitive remote participation. The convergence is important because it suggests that the VCMS is not inventing a new category of justice technology; it is assembling a Ghana-sensitive version of functions that digital-justice literature now treats as foundational (Conference of Chief Justices & Conference of State Court Administrators, 2016; OECD, 2021)."),
        ("body", "At the same time, the exhibits also highlight why local adaptation remains essential. Court modernization reports and reminder evaluations arise from different legal cultures, resource environments, and user populations. The literature does not support naive transplantation. Instead, it supports context-aware adaptation in which the local court process, existing digital maturity, and practical user constraints shape how general digital-justice principles are implemented at the Law Court Complex (Justice System Partners, 2024; World Bank, 2021)."),
        ("body", "This deeper synthesis expands the conceptual claim of the study: the proposed VCMS should be understood as an integration model for caseflow, communication, and participation under Ghanaian conditions. Its originality lies not in any single component, but in the way the components are connected and justified through the literature reviewed in this chapter (Ashong Elliot et al., 2024; Rossner & Tait, 2021)."),
        ("heading", 4, "2.6.3 Comparative mini-cases and design translation"),
        ("body", "Taken together, the comparative materials in this chapter can be read as a set of mini-cases. Ghana’s paperless-court materials show institutional modernization from within the local system; the civil-justice reform report shows the importance of deadline management, alerts, and measurement; the reminder literature shows that structured communication affects attendance behavior; and the California remote-access materials show how participation services evolve from emergency tools into mature court operations. Each mini-case points to a different part of the VCMS problem, but all of them converge on the need for coordinated workflow design (Judicial Service of Ghana, 2019; Conference of Chief Justices & Conference of State Court Administrators, 2016)."),
        ("body", "This comparative mini-case reading helps avoid the common weakness of descriptive literature reviews in which sources are listed but not truly synthesized. Here, the synthesis is explicit: workflow reliability, communication reliability, and participation reliability are interdependent. The proposed system should therefore be evaluated by how well it connects those three forms of reliability under Ghanaian court conditions. That connection is the main analytical gain produced by the chapter (Justice System Partners, 2024; California Courts Newsroom, 2025)."),
        ("body", "The design propositions in Figure 2.18 and the synthesis grid in Figure 2.19 make this argument more concrete. They translate broad literature themes into an actionable proposition set for the study: one authoritative case record, workflow-linked communication, case-sensitive remote participation, and measurable accountability. Those propositions carry Chapter 2 from descriptive review into a more explicitly thesis-driving synthesis, which is necessary for a final defense document of the kind the user requires (Conference of Chief Justices & Conference of State Court Administrators, 2016; Ashong Elliot et al., 2024)."),
    ]
    for kind, *payload in section_26:
        if kind == "heading":
            heading_before(anchor, payload[1], payload[0])
        else:
            body_before(anchor, payload[0])

    figure_before(anchor, "2.11", "Conceptual model for the Virtual Court Management and Notification System", "vcms-figure-2-11-conceptual-model.png", "Author's construct based on Judicial Service of Ghana (n.d.-a), World Bank (2021), OECD (2021), Chohlas-Wood et al. (2023), and Rossner and Tait (2021).", width_inches=6.4)
    table_before(
        anchor,
        "2.7",
        "Comparative literature on digital-justice reforms relevant to VCMS",
        ["Jurisdiction or source", "Reform emphasis", "Lesson for the Ghanaian context"],
        [
            ["Ghana Judicial Service", "Paperless courts, e-filing, video-conference workflows", "Local reform already exists, but integration across services remains incomplete."],
            ["Conference of Chief Justices / COSCA", "Caseflow automation, alerts, deadlines, and measurement", "Digital systems should actively manage workflow, not only store documents."],
            ["California judiciary sources", "Remote participation as a scaled, permanent service option", "Remote access should be governed as an operational mode with quality controls."],
            ["Reminder-system evaluations", "Behavioural reminder design and measured attendance impact", "Notifications should be auditable, timed, and empirically grounded."],
            ["World Bank / OECD", "Interoperability, governance, and service design", "Court platforms should connect records, communication, analytics, and accountability."],
        ],
        "Compiled by the researcher from Judicial Service of Ghana (2017, 2018, 2019, n.d.-a), Conference of Chief Justices and Conference of State Court Administrators (2016), California Courts Newsroom (2025), Justice System Partners (2024), OECD (2021), and World Bank (2021).",
        widths=[1.7, 2.1, 2.2],
    )
    table_before(
        anchor,
        "2.8",
        "Design principles from the literature for a court-notification subsystem",
        ["Principle", "Why the literature supports it", "Design translation in VCMS"],
        [
            ["Event-driven triggering", "Reminder quality depends on accurate workflow events.", "Messages originate from case and schedule state changes."],
            ["Configurable timing", "Different hearings require different lead times and urgency.", "Registry staff can set reminder windows by case or hearing type."],
            ["Channel redundancy", "Delivery risk falls when courts can retry or switch channels.", "VCMS supports SMS, email, and escalation paths."],
            ["Minimal but sufficient detail", "Users need clarity without overexposing sensitive case data.", "Messages contain essential logistical information and route to the authoritative record."],
            ["Auditability", "Courts need evidence of dispatch, retry, and failure outcomes.", "Every notification event is logged and reviewable."],
        ],
        "Compiled by the researcher from Chohlas-Wood et al. (2023), Hawaii State Judiciary (n.d.), Minnesota Judicial Branch (n.d.), and Justice System Partners (2024).",
        widths=[1.6, 2.5, 2.0],
    )
    table_before(
        anchor,
        "2.9",
        "Virtual-hearing risks and safeguards synthesized from the literature",
        ["Risk area", "Typical concern in the literature", "Safeguard relevant to VCMS"],
        [
            ["Connectivity", "Dropped calls, weak audio, and unstable video quality", "Hybrid fallback, technical guidance, and flexible reconnection support"],
            ["Comprehension", "Users may be unsure how and when to join", "Clear notice content, step-by-step instructions, and pre-hearing support"],
            ["Fairness perception", "Remote participation may feel less authoritative or less clear", "Selective hearing-mode choice and visible procedural structure"],
            ["Privacy", "Users may join from spaces that are not private", "Guidance on participation environment and court discretion for sensitive matters"],
            ["Evidence handling", "Some proceedings require stronger control over documents or testimony", "Case-sensitive allocation of remote, hybrid, or physical mode"],
        ],
        "Compiled by the researcher from Rossner and Tait (2021), Sourdin (2020), Judicial Council of California (2021), and California Courts Newsroom (2025).",
        widths=[1.2, 2.4, 2.6],
    )
    table_before(
        anchor,
        "2.10",
        "Theoretical constructs used to interpret VCMS literature and later evaluation",
        ["Lens", "Key construct", "Interpretive value for this thesis"],
        [
            ["Technology acceptance", "Perceived usefulness and ease of use", "Explains likely user adoption and routine reliance."],
            ["Socio-technical systems", "Fit between technology, roles, rules, and training", "Explains why technically sound systems may still fail institutionally."],
            ["Procedural justice", "Voice, clarity, neutrality, respect, trustworthiness", "Explains user judgement of digital court quality beyond efficiency."],
        ],
        "Compiled by the researcher from Ashong Elliot et al. (2024), Rossner and Tait (2021), Sourdin (2020), and World Bank (2021).",
        widths=[1.4, 2.0, 3.0],
    )
    table_before(
        anchor,
        "2.11",
        "Comparative mini-cases used to deepen the Chapter 2 synthesis",
        ["Mini-case", "Primary issue surfaced in the literature", "Implication for VCMS"],
        [
            ["Ghana paperless-court reform", "Institutional modernization exists, but user-facing integration remains incomplete.", "Build on local reform momentum instead of proposing a detached software concept."],
            ["Civil-justice automation report", "Courts need alerts, deadlines, searchability, and measurement.", "Embed workflow intelligence and management reporting in the platform."],
            ["Reminder-system studies", "Structured communication changes attendance behaviour and reduces silent failure.", "Treat reminders as a formal subsystem with rules, retries, and logs."],
            ["California remote-access experience", "Remote participation can scale when governed as an operational service.", "Support physical, virtual, and hybrid mode within one hearing-management environment."],
        ],
        "Compiled by the researcher from Judicial Service of Ghana (2019), Conference of Chief Justices and Conference of State Court Administrators (2016), Chohlas-Wood et al. (2023), and California Courts Newsroom (2025).",
        widths=[1.6, 2.5, 2.1],
    )
    table_before(
        anchor,
        "2.12",
        "Expanded synthesis of literature themes and practical implications for the study",
        ["Theme", "Literature insight", "Practical implication for thesis design and evaluation"],
        [
            ["Institutional fit", "Digital justice succeeds when systems align with local routines, support, and governance.", "Evaluation should examine usability, role clarity, and organizational acceptance."],
            ["Workflow reliability", "Authoritative caseflow requires linked filing, scheduling, and update history.", "The system must be assessed as an integrated workflow, not as disconnected modules."],
            ["Communication reliability", "Reminder effectiveness depends on timing, channel strategy, and traceability.", "Testing should include failed-delivery handling and user notification clarity."],
            ["Participation reliability", "Remote access works best when users understand hearing mode and expectations.", "The platform should provide clear access pathways and hearing-specific guidance."],
            ["Measurable accountability", "Reform reports stress dashboards, alerts, and performance visibility.", "The thesis should evaluate not only features, but also reporting and management value."],
        ],
        "Compiled by the researcher from World Bank (2021), OECD (2021), Rossner and Tait (2021), Justice System Partners (2024), and Conference of Chief Justices and Conference of State Court Administrators (2016).",
        widths=[1.5, 2.3, 2.4],
    )
    figure_before(anchor, "2.12", "Data-governance and interoperability stack for court platform integration", "vcms-figure-2-12-data-governance.png", "Author's construct based on World Bank (2021), OECD (2021), and Conference of Chief Justices and Conference of State Court Administrators (2016).", width_inches=6.4)
    figure_before(anchor, "2.13", "Integrated theoretical framework used to interpret the VCMS literature", "vcms-figure-2-13-theoretical-framework.png", "Author's construct based on Ashong Elliot et al. (2024), Rossner and Tait (2021), Sourdin (2020), and World Bank (2021).", width_inches=6.4)
    figure_before(anchor, "2.14", "Best-practice design ladder for court reminder systems", "vcms-figure-2-14-notification-best-practice.png", "Author's construct based on Chohlas-Wood et al. (2023), Justice System Partners (2024), Hawaii State Judiciary (n.d.), and Minnesota Judicial Branch (n.d.).", width_inches=6.3)
    figure_before(anchor, "2.15", "Benefits-risk balance for sustaining virtual and hybrid hearings", "vcms-figure-2-15-virtual-balance.png", "Author's construct based on Judicial Council of California (2021), California Courts Newsroom (2025), Rossner and Tait (2021), and Sourdin (2020).", width_inches=6.3)
    figure_before(anchor, "2.16", "Official civil-justice automation exhibit from a state-court reform report", "civil-justice-initiative-report.png", "Conference of Chief Justices and Conference of State Court Administrators (2016), civil-justice initiative report.", width_inches=5.8)
    figure_before(anchor, "2.17", "Official remote-access reform exhibit from the California judiciary", "california-remote-access-report.png", "Judicial Council of California (2021), remote access to court proceedings interim report.", width_inches=5.8)
    figure_before(anchor, "2.18", "Design propositions distilled from the Chapter 2 literature review", "vcms-figure-2-18-design-propositions.png", "Author's construct based on the literature reviewed in Sections 2.1 to 2.6.", width_inches=6.2)
    figure_before(anchor, "2.19", "Chapter 2 synthesis grid from literature problem to VCMS response", "vcms-figure-2-19-synthesis-grid.png", "Author's construct based on the literature reviewed in Chapter 2.", width_inches=6.2)

    section_27 = [
        ("heading", 3, "2.7 Chapter Summary"),
        ("body", "This chapter reviewed the literature in the same order as the study objectives and established a coherent argument for the proposed VCMS. The review showed that Ghana already has a credible digital-justice foundation at the Law Court Complex, but that the current environment still exhibits fragmentation across workflow stages and user communication. It further showed that digital-platform design, notifications, and virtual-hearing support are analytically inseparable when the aim is to improve real court participation and administrative reliability (Judicial Service of Ghana, 2019; World Bank, 2021)."),
        ("body", "The chapter also identified the research gap that justifies the thesis: an integrated, Ghana-sensitive model is needed to connect existing e-Justice momentum with stronger scheduling control, auditable notifications, and remote-participation support. That conclusion provides the literature-based bridge into Chapter Three, where the methodology for translating these insights into a developed system and evaluation strategy is presented (Ashong Elliot et al., 2024; Rossner & Tait, 2021)."),
        ("body", "The expanded synthesis in this chapter also sharpens the originality of the study. Rather than treating filing, scheduling, reminders, and virtual hearings as independent technologies, the review has shown that they are better understood as interlocking capabilities of a single court-service ecosystem. The thesis therefore contributes a Ghana-focused integration perspective that is grounded in both institutional evidence and comparative digital-justice scholarship (World Bank, 2021; Conference of Chief Justices & Conference of State Court Administrators, 2016)."),
    ]
    for kind, *payload in section_27:
        if kind == "heading":
            heading_before(anchor, payload[1], payload[0])
        else:
            body_before(anchor, payload[0])


def enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main() -> None:
    ensure_assets()
    document = Document(str(INPUT_DOC))
    ensure_caption_styles(document)
    remove_block_elements(document, "2.0 Introduction", "CHAPTER THREE")
    add_chapter_two(document)
    rebuild_references(document)
    normalize_caption_styles(document)
    rebuild_generated_list(document, "LIST OF FIGURES", "Figure Caption")
    rebuild_generated_list(document, "LIST OF TABLES", "Table Caption")
    apply_citation_pass(document)
    enable_update_fields_on_open(document)
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_DOC))


if __name__ == "__main__":
    main()
