from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


INPUT_DOC = Path("/Users/emmanuel/Desktop/apps/gerts/output/doc/GERTS_thesis_reworked_final_v6.docx")
OUTPUT_DOC = Path("/Users/emmanuel/Desktop/apps/gerts/output/doc/GERTS_thesis_reworked_final_v7.docx")


def insert_paragraph_after_element(element, parent, text: str = "", style_name: str | None = None):
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    paragraph = Paragraph(new_p, parent)
    if style_name:
        paragraph.style = style_name
    if text:
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
    return paragraph


def format_source(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.first_line_indent = Inches(0.0)
    paragraph.paragraph_format.left_indent = Inches(0.0)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10)
        run.italic = True
        run.bold = False


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def style_table(table) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if r_idx else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(11)
                    run.bold = r_idx == 0


def main() -> None:
    document = Document(INPUT_DOC)
    paragraphs = document.paragraphs
    caption_idx = next(i for i, p in enumerate(paragraphs) if " ".join(p.text.split()).strip() == "Table 2.1: Key concepts and working definitions")
    caption_paragraph = paragraphs[caption_idx]
    source_paragraph = paragraphs[caption_idx + 1]

    source_text = " ".join(source_paragraph.text.split()).strip()
    if source_text.startswith("Source:"):
        delete_paragraph(source_paragraph)

    table = document.add_table(rows=1, cols=3)
    style_table(table)
    headers = ["Concept", "Working definition", "Design relevance to GERTS"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text

    rows = [
        [
            "Election results management",
            "The set of processes and tools used to record, transmit, collate, verify, and publish election results from polling stations to the final declaration level.",
            "Defines the operational workflow that GERTS must support from field capture to public publication.",
        ],
        [
            "Transparency in results reporting",
            "The extent to which results and supporting evidence are accessible, understandable, and open to verification by stakeholders.",
            "Requires low-level publication, evidence access, and clear status information in the system interface.",
        ],
        [
            "Verifiability and auditability",
            "The ability of stakeholders to check that reported results correspond to recorded evidence and that any significant manipulation is detectable.",
            "Motivates signed submissions, immutable audit trails, and public verification mechanisms.",
        ],
        [
            "Trust in election technology",
            "Stakeholder confidence in the election process, including the devices, software, procedures, and people operating them.",
            "Shows that usability, governance, and transparency are as important as technical controls.",
        ],
        [
            "End-to-end verifiability (E2E-V)",
            "Approaches that make recorded outcomes evidence-visible and checkable by observers while preserving required privacy protections.",
            "Informs the evidence-first design logic behind publication, verification, and audit support in GERTS.",
        ],
        [
            "Risk-limiting audit (RLA)",
            "A post-election audit that provides statistical confidence in the reported outcome by examining a sample of voter-verifiable records.",
            "Frames how GERTS should preserve evidence quality for future audit and dispute review.",
        ],
        [
            "Immutable audit log",
            "A log designed so that tampering is detectable, often implemented through cryptographic hashing and append-only storage.",
            "Supports non-repudiation, traceability, and forensic readiness in the proposed architecture.",
        ],
        [
            "Permissioned blockchain",
            "A blockchain network where participation is restricted to approved entities and governed by agreed membership rules.",
            "Provides the tamper-evident and governed ledger model considered appropriate for the GERTS context.",
        ],
    ]
    for row_values in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_values):
            cell.text = text

    style_table(table)
    tbl = table._tbl
    caption_paragraph._element.addnext(tbl)

    source = insert_paragraph_after_element(tbl, document._body, "Source: Compiled by researcher from literature reviewed, 2026.", "Normal")
    format_source(source)

    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
